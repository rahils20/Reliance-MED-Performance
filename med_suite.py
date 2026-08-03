import streamlit as st
import datetime
import pandas as pd
import numpy as np
import io
import os
import time
import math
import joblib
import base64
import altair as alt
from io import BytesIO
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def standardize_dates(date_series):
    """Robust master parser for multi-format date registries.
    Intercepts any format (1-Apr-26, 2026-04-01, 01/04/2026, 01-04-2026) and aligns them.
    All ambiguous numeric formats (slash or hyphen) are treated as DAY-FIRST (DD-MM-YYYY),
    matching the plant's standard convention. Never falls back to pandas' default
    month-first interpretation, which was silently swapping day/month for any
    day-of-month <= 12 (e.g. 09-07-2026 read as 7 September instead of 9 July)."""
    parsed = pd.to_datetime(date_series, format='%d-%b-%y', errors='coerce')
    parsed = parsed.fillna(pd.to_datetime(date_series, format='%d-%b-%Y', errors='coerce'))
    parsed = parsed.fillna(pd.to_datetime(date_series, format='%Y-%m-%d', errors='coerce'))
    parsed = parsed.fillna(pd.to_datetime(date_series, format='%d/%m/%Y', errors='coerce'))
    parsed = parsed.fillna(pd.to_datetime(date_series, format='%d-%m-%Y', errors='coerce'))
    # Final catch-all: force dayfirst=True instead of pandas' default month-first
    # inference, so any leftover ambiguous numeric date is still read as DD-MM-YYYY.
    parsed = parsed.fillna(pd.to_datetime(date_series, errors='coerce', dayfirst=True))
    return parsed

def upsert_daily_logs(existing_df, new_df):
    """Merge new_df into existing_df by Date, updating ONLY the columns new_df actually provides for
    matching dates, and leaving every other existing column for that date untouched. New rows are
    appended for dates that don't exist yet.

    This replaces the old concat + drop_duplicates(keep='last') pattern, which does a destructive
    whole-row replace: uploading a narrower file (e.g. HTC-only data) for a date that already has
    operational data would silently wipe that operational data out, since the "new" row for that
    date wouldn't have those columns at all. With separate Operational / HTC / Water Quality bulk
    uploads now sharing the same master registry, that whole-row-replace behavior would destroy data
    on every single upload - this function is what makes running them independently safe.
    """
    new_df = new_df.copy()
    new_df['Date'] = standardize_dates(new_df['Date']).dt.strftime('%Y-%m-%d')
    new_df = new_df.dropna(subset=['Date']).drop_duplicates(subset=['Date'], keep='last').set_index('Date')

    if existing_df is None or existing_df.empty or 'Date' not in existing_df.columns:
        return new_df.reset_index()

    existing_df = existing_df.copy()
    existing_df['Date'] = standardize_dates(existing_df['Date']).dt.strftime('%Y-%m-%d')
    existing_df = existing_df.dropna(subset=['Date']).drop_duplicates(subset=['Date'], keep='last').set_index('Date')

    # Make sure both frames share the same columns before combining, so combine_first has a clean
    # NaN to fall back on rather than silently dropping a column one side doesn't have.
    for col in new_df.columns:
        if col not in existing_df.columns:
            existing_df[col] = np.nan
    for col in existing_df.columns:
        if col not in new_df.columns:
            new_df[col] = np.nan

    # new_df's non-null values always win (it's the fresh upload); anything new_df leaves null
    # (including entire columns it doesn't cover) falls back to whatever existing_df already had.
    merged = new_df.combine_first(existing_df)
    return merged.reset_index()

# MED GLOBAL CONSTANTS
# ---------------------------------------------------------------------------------------------
# MRA PREDICTOR SPECIFICATION (revised for Reliance, replacing the original 2014 7-input model)
#
# One row per predictor, so the live prediction, the variance table, the training pipeline and the
# bulk template all read from a SINGLE definition. Adding or removing a predictor is now a one-line
# change here rather than eight parallel edits that could silently drift out of alignment.
#
#   coef_key      key used in the saved calibration config
#   db_column     column name in the master registry / training CSV
#   live_var      session-state variable feeding the live daily prediction
#   label         human-readable name shown in the UI
#   fallback_base provisional reference value, used ONLY until a calibration writes the real
#                 post-cleaning mean for this plant (see BASE_ prefixed keys in the config)
#
# Anti_PPM (antiscalant residual) was REMOVED at Reliance's instruction - the residual measurement
# was judged unreliable and was degrading the fit. Its contribution at the old baseline dosing
# (-7.0301 x 4.82 ppm = -33.885) has been folded into the default intercept below, so a provisional
# prediction made at typical dosing lands in the same place it used to rather than jumping ~34 m3/h.
#
# LLP steam flow/temperature, 1st effect flow, desuperheating water flow and FFC tube side pressure
# were requested but are NOT included: the first three could not be mapped to an existing tag with
# confidence, and the last two have no column or history in the registry. Columns for the latter two
# are reserved in EXACT_DB_COLUMNS so capture can begin now and they can join the model later.
MED_MRA_PARAMS = [
    # coef_key,           db_column,                    live_var,      label,                       fallback_base
    ("Press_1st",         "1st effect vapour pressure", "mra_press",   "1st effect vapour pressure",  231.76),
    ("Temp_1st",          "1st Effect Vapour Temp",     "mra_t1",      "1st effect vapour temp",       68.47),
    ("Brine_Temp_1st",    "1st effect brine temp",      "mra_bt1",     "1st effect brine temp",        65.46),
    ("SW_Upper",          "Sea Water Upper",            "sw_upper",    "Sea water upper",             553.63),
    ("SW_Feed_Total",     "Sea Water Feed",             "sw_total",    "Seawater feed to MED",       2062.00),
    ("Brine_Flow",        "Brine Water Return",         "brine_ret",   "Brine water return",         1275.50),
    ("LP_Steam",          "LP Steam consumption",       "steam",       "LP steam consumption",         71.75),
    ("Cond_Flow",         "Condensate Return",          "cond_flow",   "Condensate flow",              71.75),
    ("Brine_Disch_Temp",  "Brine Discharge Temp",       "brine_out_t", "Brine discharge temp",         40.00),
    ("Brine_11th_Temp",   "11th Effect Brine Temp",     "brine_11",    "11th effect brine temp",       40.00),
]

MRA_COEF_KEYS   = [p[0] for p in MED_MRA_PARAMS]
MRA_DB_COLUMNS  = [p[1] for p in MED_MRA_PARAMS]
MRA_LIVE_VARS   = [p[2] for p in MED_MRA_PARAMS]
MRA_LABELS      = [p[3] for p in MED_MRA_PARAMS]
N_MRA_PREDICTORS = len(MED_MRA_PARAMS)

# Provisional references. These are overwritten per-plant by the calibration run, which stores the
# actual mean of each predictor over the clean post-cleaning window under "BASE_<coef_key>". That
# is what the deviation column in the variance table is supposed to measure against - this plant
# when clean, not a generic design figure.
MRA_BASELINE = {p[0]: p[4] for p in MED_MRA_PARAMS}

# Default coefficients. The six predictors carried over from the previous model keep their 2014
# values; the four newly added inputs start at 0.0 because no fit exists for them yet. "calibrated"
# stays 0 until a real calibration is committed, which drives the warning banner on the dashboard.
# This is deliberately a PROVISIONAL state, not a pretend-fitted model.
MED_MRA_COEF_DEFAULT = {
    "model_type": "OLS",
    "calibrated": 0,
    "Intercept": -195.4489,
    "Press_1st": 0.6136,
    "Temp_1st": 3.6392,
    "Brine_Temp_1st": -7.6638,
    "SW_Upper": 0.8111,
    "SW_Feed_Total": 0.0,
    "Brine_Flow": -0.2329,
    "LP_Steam": 8.2539,
    "Cond_Flow": 0.0,
    "Brine_Disch_Temp": 0.0,
    "Brine_11th_Temp": 0.0,
}

# Kept as an alias so nothing that still imports the old name breaks.
MRA_COEF_2014 = MED_MRA_COEF_DEFAULT

BASE_EFFECTS = pd.DataFrame({
    "Effect ID": [f"Effect {i}" for i in range(1, 12)],
    "Base Vapor (°C)": np.round(np.linspace(69.0, 42.0, 11), 1),
    "Base Brine (°C)": np.round(np.linspace(66.3, 40.0, 11), 1),
    "Base HTC": np.round(np.linspace(2800.0, 1500.0, 11), 1) 
})

WATER_SPECS = {
    "Feed": {
        "pH": {"lim": (7.5, 9.2), "var": "f_ph", "db_col": "Feed_pH", "avg": 8.14},
        "Turbidity (NTU)": {"lim": (0.0, 5.0), "var": "f_turb", "db_col": "Feed_Turbidity", "avg": 3.2},
        "TSS (ppm)": {"lim": (0.0, 10.0), "var": "f_tss", "db_col": "Feed_TSS", "avg": 6.5},
        "TDS (ppm)": {"lim": (0.0, 42000.0), "var": "f_tds", "db_col": "Feed_TDS", "avg": 41000.0},
        "Total Alkalinity": {"lim": (160.0, 190.0), "var": "f_alk", "db_col": "Feed_Alkalinity", "avg": 170.0},
        "Calcium Hardness": {"lim": (950.0, 1100.0), "var": "f_ca", "db_col": "Feed_Calcium", "avg": 1040.0},
        "Mg Hardness": {"lim": (5400.0, 5700.0), "var": "f_mg", "db_col": "Feed_MgHardness", "avg": 5550.0},
        "Total Hardness": {"lim": (0.0, 7000.0), "var": "f_hard", "db_col": "Feed_TotalHardness", "avg": 6640.0},
        "Conductivity (μs/cm)": {"lim": (0.0, 70000.0), "var": "f_cond", "db_col": "Feed_Cond", "avg": 57000.0},
        "Silica": {"lim": (0.0, 0.67), "var": "f_sio2", "db_col": "Feed_Silica", "avg": 0.3},
        "Chlorides": {"lim": (21000.0, 22000.0), "var": "f_cl", "db_col": "Feed_Chlorides", "avg": 21500.0},
        "Sulphate": {"lim": (3050.0, 3250.0), "var": "f_so4", "db_col": "Feed_Sulphate", "avg": 3150.0},
        "Sulphide": {"lim": (0.0, 1.0), "var": "f_sulfide", "db_col": "Feed_Sulphide", "avg": 0.0}
    },
    "Product": {
        "pH": {"lim": (5.5, 7.0), "var": "p_ph", "db_col": "Product_pH", "avg": 6.5},
        "Turbidity (NTU)": {"lim": (0.0, 1.0), "var": "p_turb", "db_col": "Product_Turbidity", "avg": 0.1},
        "Conductivity (μs/cm)": {"lim": (0.0, 15.0), "var": "p_cond", "db_col": "Product_Cond", "avg": 4.6},
        "TDS (ppm)": {"lim": (0.0, 10.0), "var": "p_tds", "db_col": "Product_TDS", "avg": 2.5},
        "Total Alkalinity": {"lim": (0.0, 10.0), "var": "p_alk", "db_col": "Product_Alkalinity", "avg": 2.0},
        "Calcium Hardness": {"lim": (0.0, 1.0), "var": "p_ca", "db_col": "Product_Calcium", "avg": 0.0},
        "Mg Hardness": {"lim": (0.0, 1.0), "var": "p_mg", "db_col": "Product_MgHardness", "avg": 0.0},
        "Total Hardness": {"lim": (0.0, 0.1), "var": "p_hard", "db_col": "Product_TotalHardness", "avg": 0.0},
        "Total Iron": {"lim": (0.0, 0.1), "var": "p_iron", "db_col": "Product_Iron", "avg": 0.05},
        "Silica": {"lim": (0.0, 0.02), "var": "p_sio2", "db_col": "Product_Silica", "avg": 0.0},
        "Chlorides": {"lim": (0.0, 5.0), "var": "p_cl", "db_col": "Product_Chlorides", "avg": 0.0},
        "Sulphate": {"lim": (0.0, 1.0), "var": "p_so4", "db_col": "Product_Sulphate", "avg": 0.0}
    }
}

# Brine water analysis - right-hand block of the 'Feed & Brine Water Analysis' sheet.
# The sheet lists no specified limits for brine, so these are tracked/trended, not pass-fail graded.
BRINE_SPECS = {
    "pH": {"var": "b_ph", "db_col": "Brine_pH", "avg": 8.4},
    "Turbidity (NTU)": {"var": "b_turb", "db_col": "Brine_Turbidity", "avg": 14.0},
    "Conductivity (μs/cm)": {"var": "b_cond", "db_col": "Brine_Cond", "avg": 80500.0},
    "TDS (ppm)": {"var": "b_tds", "db_col": "Brine_TDS", "avg": 52325.0},
    "Total Alkalinity": {"var": "b_alk", "db_col": "Brine_Alkalinity", "avg": 218.0},
    "Calcium Hardness": {"var": "b_ca", "db_col": "Brine_Calcium", "avg": 1790.0},
    "Mg Hardness": {"var": "b_mg", "db_col": "Brine_MgHardness", "avg": 10710.0},
    "Total Hardness": {"var": "b_hard", "db_col": "Brine_TotalHardness", "avg": 12500.0},
    "Silica": {"var": "b_sio2", "db_col": "Brine_Silica", "avg": 0.0},
    "Chlorides": {"var": "b_cl", "db_col": "Brine_Chlorides", "avg": 31200.0},
}

EXACT_DB_COLUMNS = [
    "Date", "Sea Water Upper", "Sea Water Lower", "Sea Water Feed", "Sea Water Pressure",
    "Brine Water Return", "Desal production", "LP Steam consumption", "LP Steam Pressure",
    "Condensate Return", "condensate temp", "Condensate Conductivity",
    "1st Effect Vapour Temp", "1st effect brine temp", "11th Effect Brine Temp", "Feed Temp to Cold Group",
    "Intermediate Effects Avg Brine Temp", "Delta T", "1st effect vapour pressure", "Brine Discharge Temp", "Brine Discharge Pressure",
    "Sea Water cond I/L temp", "Sea Water Condenser O/L Temp", 
    "CW supply", "CW Return", "CW Flow", "Gross production", "GOR", "STEC", "Overall HTC", "1st Effect HTC", 
    "Residual", "Antiscalant (kg)", "Antifoam (kg)", "Anti_PPM", "Foam_PPM", "Area_1st", "Area_Overall", "Remarks",
    # --- Operational sheet extras ---
    "Steam Inlet Temp", "Recovery", "Conversion", "Steam Economy", "Overall Delta T",
    "Anti_PPM_Hot", "Anti_PPM_Brine",
    # --- Reserved for the two MRA tags Reliance requested that have no history yet. Capturing them
    #     from now on means they can be added to the model once enough rows accumulate.
    "Desuperheating Water Flow", "FFC Tube Side Pressure",
    # --- 1st Effect HTC sheet: its OWN inputs. "Feed flow" here is flow to the 1st effect (~514 m3/hr)
    #     and "Feed Temp" here is the AVG BRINE TEMP OF EFFECTS 4,5,6,7 (~49C) - both are physically
    #     different measurements from the identically-named columns on the Overall-HTC sheet.
    "HTC1_Feed_Flow", "HTC1_Product_Flow", "HTC1_Cond_Flow", "HTC1_Steam_TPH",
    "HTC1_Feed_Temp_Eff4to7", "HTC1_Brine_Temp", "HTC1_Vapor_Temp", "HTC1_Cond_Temp",
    "HTC1_dT1", "HTC1_dT2", "HTC1_LMTD", "HTC1_Q_Steam", "HTC1_Fouling", "HTC1_Rf",
    # --- Overall HTC sheet: its OWN inputs. "Feed flow" here is TOTAL seawater feed (~2062 m3/hr)
    #     and "Feed Temp" here is the FEED TEMP TO COLD GROUP (~40C). Area is 11x12950x1.15.
    "HTCO_Feed_Flow", "HTCO_Product_Flow", "HTCO_Cond_Flow", "HTCO_Steam_TPH",
    "HTCO_Feed_Temp_ColdGrp", "HTCO_Brine_Disch_Temp", "HTCO_Vapor_Temp", "HTCO_Cond_Temp",
    "HTCO_dT1", "HTCO_dT2", "HTCO_LMTD", "HTCO_Q_Steam", "HTCO_Fouling", "HTCO_Rf",
    # --- Chemicals doses sheet: antiscalant (Kem Watreat r 3687) + antifoam (Kem Antifoam 1795).
    #     Dosing derived from tank level drop: (Initial + Top-up - Final) over N hrs -> LPH -> Kg/hr -> PPM.
    "AS_Initial", "AS_Topup", "AS_Final", "AS_LevelDrop", "AS_Hours", "AS_LPH", "AS_KgHr", "AS_PPM",
    "AF_Initial", "AF_Topup", "AF_Final", "AF_LevelDrop", "AF_Hours", "AF_LPH", "AF_KgHr", "AF_PPM",
    # MMC stock (KG): opening / received / consumed / closing for each chemical.
    "AS_Stock_Open", "AS_Stock_Recd", "AS_Stock_Consumed", "AS_Stock_Close",
    "AF_Stock_Open", "AF_Stock_Recd", "AF_Stock_Consumed", "AF_Stock_Close",
]
for cat in ['Feed', 'Product']:
    for param, details in WATER_SPECS[cat].items(): 
        EXACT_DB_COLUMNS.append(details['db_col'])
for param, details in BRINE_SPECS.items():
    EXACT_DB_COLUMNS.append(details['db_col'])

RIL_EXCEL_HEADERS = [
    'Parameter', 'Sea water Upper', 'Sea water Lower', 'Sea water feed', 'Brine return', 
    ' Desal Production', 'LP Steam Consumption', 'Condensate return', 'Condensate Temp', 
    "1'st effect vapour Temp", '1st Effect Brine Temp', '(1st effect vapour-1st effect brine) Delta Temp', 
    '1st Effect Vapour pres', 'Steam Inlet Temp', 'Brine DischargeTemp', 'Sea water cond (FFC) I/L temp', 
    'Sea water cond (FFC) o/L temp', 'CW (FCC) supply', 'CW (FCC) return', 
    'Gross desal water production', 'Recovery', 'Conversion (Product to Feed)', 'Gain Output Ratio', 
    '11 effect brine Temp', 'Overall delta T(1st eff brine temp - 11th eff brine temp)', 
    'Steam Economy (Steam/Desal)', 'Antiscalant residual (Cold group)', 'Antiscalant residual (Hot group)', 
    'Antiscalant residual (Brine)', 'Feed Temp to Cold Group', 'Intermediate Effects Avg Brine Temp (4,5,6,7)', 'Remarks'
]

# --- HTC reference constants, read straight off rows 5 (Design) and 6 (SOR/Baseline) of the two HTC sheets.
# Rf (fouling resistance) = 1/U_actual - 1/U_clean, where the sheets use the SOR baseline as "clean".
HTC_1ST_AREA = 12950.013120000001      # 1st effect-HTC!K  = pi * 5.5m * 31244 tubes * 0.024m OD
HTC_OVERALL_AREA = 163818.0            # Overall-HTC!K     = 11 effects * 12950 * 1.15
HTC_1ST_U_SOR = 415.31060504252554     # 1st effect-HTC!AA6 (steam condensation basis, SOR baseline)
HTC_OVERALL_U_SOR = 17.726796070321715 # Overall-HTC!AA6   (steam condensation basis, SOR baseline)
CP_WATER_KJ_KGC = 4.186                # specific heat, both sheets col P

# --- 1st Effect HTC bulk template: mirrors the '1st effect-HTC' sheet's INPUT columns (A-K) exactly.
# Everything from dT1 onward (L..AC) is recomputed by the calculator, not read from the file.
HTC_1ST_BULK_HEADERS = [
    'Date', 'Feed flow', 'Product flow', 'Condensate Flow', 'Steam consumption rate',
    'Feed Temp', 'Brine Temp', '1st effect vapor temp', 'Condensate temperature', 'Heat Transfer Area'
]

# --- Overall HTC bulk template: mirrors the 'Overall-HTC' sheet's INPUT columns (A-K) exactly.
# NOTE: 'Feed flow', 'Feed Temp' and the brine column here are DIFFERENT physical measurements from the
# same-named columns on the 1st-effect sheet - which is exactly why these need to be two separate uploads.
HTC_OVERALL_BULK_HEADERS = [
    'Date', 'Feed flow', 'Product flow', 'Condensate Flow', 'Steam consumption rate',
    'Feed Temp', 'Brine discharge Temp', '1st effect vapor temp', 'Condensate temperature', 'Heat Transfer Area'
]

# --- Feed & Brine Water Analysis template: mirrors that sheet's columns A-X.
FEEDBRINE_BULK_HEADERS = [
    'Date', 'pH', 'Turbidity', 'TSS', 'Conductivity', 'TDS', 'Total Alkalinity', 'Calcium Hardness',
    'Mg Hardness', 'Total Hardness', 'Silica', 'Chloride', 'Sulphate', 'Sulphide',
    'Brine pH', 'Brine Turbidity', 'Brine Conductivity', 'Brine TDS', 'Brine Total Alkalinity',
    'Brine Calcium Hardness', 'Brine Mg Hardness', 'Brine Total Hardness', 'Brine Silica', 'Brine Chloride',
    'REMARKS'
]

# --- Desal (product) Analysis template: mirrors that sheet's columns A-N.
DESAL_BULK_HEADERS = [
    'Date', 'pH', 'Turbidity', 'Conductivity', 'TDS', 'Total Alkalinity', 'Calcium Hardness',
    'Mg Hardness', 'Total Hardness', 'Chloride', 'Total Iron', 'Silica', 'Sulphate', 'REMARKS'
]

# --- Chemical Doses template: mirrors the 'Chemicals doses' sheet. Only the raw INPUTS are uploaded
# (tank initial/top-up/final levels, hours, and MMC stock movements). LPH, Kg/hr and PPM are recomputed.
CHEM_BULK_HEADERS = [
    'Date',
    'AS Initial', 'AS Top-up', 'AS Final', 'AS Nos of Hrs',
    'AF Initial', 'AF Top-up', 'AF Final', 'AF Nos of Hrs',
    'AS Stock Opening', 'AS Stock Received', 'AS Stock Consumed', 'AS Stock Closing',
    'AF Stock Opening', 'AF Stock Received', 'AF Stock Consumed', 'AF Stock Closing',
    'Remarks'
]
# Physical constants, reverse-engineered exactly from the 'Chemicals doses' sheet's own ratios:
#   LPH        = (level drop / hours) x 23      (tank litres per unit of level drop)
#   AS kg/hr   = LPH x 1.20                     (Kem Watreat r 3687 density)
#   AF kg/hr   = LPH x 0.02                     (Kem Antifoam 1795 effective density factor)
#   PPM        = kg/hr x 1000 / feed(m3/hr)     (dose relative to seawater feed)
LITRES_PER_LEVEL_UNIT = 23.0
AS_DENSITY = 1.20
AF_DENSITY = 0.02

# --- Operational Data bulk template: throughput/production/chemicals only. Matches your existing
# 'Operational data' sheet / DCS export format exactly, so your existing file works unmodified.
# Computes GOR, STEC and MRA Residual - never touches any HTC field.
OPERATIONAL_BULK_HEADERS = [
    'Parameter', 'Sea water Upper', 'Sea water Lower', 'Sea water feed', 'Brine return',
    ' Desal Production', 'LP Steam Consumption', 'Condensate return', 'Condensate Temp',
    "1'st effect vapour Temp", '1st Effect Brine Temp', '1st Effect Vapour pres',
    'Steam Inlet Temp', 'Brine DischargeTemp',
    'Sea water cond (FFC) I/L temp', 'Sea water cond (FFC) o/L temp',
    'CW (FCC) supply', 'CW (FCC) return', 'Gross desal water production', '11 effect brine Temp',
    'Antiscalant residual (Cold group)', 'Antiscalant residual (Hot group)', 'Antiscalant residual (Brine)',
    'Remarks'
]
# Derived on the sheet, deliberately NOT in the template: Delta Temp, Recovery, Conversion,
# Gain Output Ratio, Overall delta T, Steam Economy. The calculator recomputes all of these.

# --- HTC Data bulk template: mirrors your '1st effect-HTC' and 'Overall-HTC' calculation sheets exactly.
# Column names are unambiguous about which effect they belong to (unlike the source sheets, which reuse
# generic names like "Feed Temp" for two physically different measurements - that ambiguity is exactly
# what's been causing confusion). Computes Overall HTC and 1st Effect HTC via LMTD - never touches any
# operational/production field.
HTC_BULK_HEADERS = [
    'Date', 'LP Steam Consumption (TPH)', '1st Effect Vapour Temp (C)', 'Condensate Temp (C)',
    '1st Effect Brine Temp (C)', 'Intermediate Effects Avg Brine Temp 4-5-6-7 (C)',
    'Brine Discharge Temp (C)', 'Feed Temp to Cold Group (C)',
    '1st Effect Vapour Pressure (optional)', '11th Effect Brine Temp (optional)', 'Remarks'
]

DEFAULTS = {
    'steam': 71.75, 'stm_press': 4.3, 'desal': 800.0, 'gross': 801.4, 'sw_upper': 553.63, 'sw_total': 2100.0, 'sw_press': 1.7, 
    'brine_ret': 1275.5, 'brine_press': 1.3,
    'sw_in_t': 30.0, 'brine_out_t': 41.0, 'vap_out_t': 70.0, 'mra_press': 231.76, 'mra_t1': 68.47, 'mra_bt1': 65.46,
    'brine_11': 40.17, 'feed_cold': 40.0, 'mid_effects_temp': 49.14, 'htc1_feed_flow': 514.0,
    'steam_in_t': 172.34,
    'f_ph': 8.14, 'f_turb': 3.2, 'f_tss': 6.5, 'f_tds': 41000.0, 'f_alk': 170.0, 'f_ca': 1040.0, 'f_mg': 5550.0,
    'f_hard': 6640.0, 'f_cond': 57000.0, 'f_sio2': 0.3, 'f_cl': 21500.0, 'f_so4': 3150.0, 'f_sulfide': 0.0,
    'p_ph': 6.5, 'p_turb': 0.1, 'p_cond': 4.6, 'p_tds': 2.5, 'p_alk': 2.0, 'p_ca': 0.0, 'p_mg': 0.0,
    'p_hard': 0.0, 'p_iron': 0.05, 'p_sio2': 0.0, 'p_cl': 0.0, 'p_so4': 0.0,
    'b_ph': 8.4, 'b_turb': 14.0, 'b_cond': 80500.0, 'b_tds': 52325.0, 'b_alk': 218.0, 'b_ca': 1790.0,
    'b_mg': 10710.0, 'b_hard': 12500.0, 'b_sio2': 0.0, 'b_cl': 31200.0,
    'chem_anti_ppm': 4.82, 'chem_anti_cons': 13.5, 'chem_foam_ppm': 0.0, 'chem_foam_cons': 0.0,
    # Area_1st = pi * tube_length(5.5m) * tube_count(31244) * tube_OD(0.024m); Area_Overall = 11 effects * Area_1st * 1.15
    # (correction factor). Previous defaults (1757.49 / 19332.0) were roughly 7-8x too small, which alone made HTC
    # numbers wrong by close to an order of magnitude regardless of anything else. See tube-geometry calc sheet.
    'skip_eff': False, 'skip_wq': False, 'remarks': "", 'area_1st': HTC_1ST_AREA, 'area_overall': HTC_OVERALL_AREA,
    'sw_lower': 0.0, 'cond_flow': 0.0, 'cond_temp': 0.0, 'htc1_cond_temp': 0.0, 'htco_cond_temp': 0.0, 'cond_cond': 3.0, 'sw_out_t': 0.0, 'cw_supply': 0.0, 'cw_return': 0.0, 'cw_flow': 2726.0
}

SYNC_MAP = {
    'steam': ['in_steam', 't5_steam'], 'stm_press': ['in_stm_press'], 'desal': ['in_desal'], 'gross': ['in_gross'],
    'sw_upper': ['in_sw_up', 't5_sw_up', 't2_sw_up'], 'sw_total': ['in_sw_tot', 't4_sw_tot', 't2_sw_tot'], 'sw_press': ['in_sw_press'],
    'brine_ret': ['in_brine', 't5_bflow'], 'brine_press': ['in_brine_press'], 
    'sw_in_t': ['in_sw_in', 't2_sw_in'], 'brine_out_t': ['in_brine_out', 't2_brine_out'], 
    'vap_out_t': ['in_vap_out', 't2_vap_out'], 'mra_press': ['in_press', 't5_press'], 
    'mra_t1': ['in_t1', 't5_t1', 't2_t1'], 'mra_bt1': ['in_bt1', 't5_bt1', 't2_bt1'], 
    'brine_11': ['in_brine_11'], 'feed_cold': ['in_feed_cold', 't2_feed_cold'],
    'mid_effects_temp': ['in_mid_effects_temp', 't2_mid_effects_temp'],
    'htc1_feed_flow': ['in_htc1_feed_flow', 't2_htc1_feed_flow'], 'steam_in_t': ['in_steam_in_t'],
    'f_ph': ['in_f_ph', 't3_f_ph'], 
    'f_turb': ['in_f_turb', 't3_f_turb'], 'f_tss': ['in_f_tss', 't3_f_tss'], 'f_tds': ['in_f_tds', 't3_f_tds'],
    'f_alk': ['in_f_alk', 't3_f_alk'], 'f_ca': ['in_f_ca', 't3_f_ca'], 'f_mg': ['in_f_mg', 't3_f_mg'],
    'f_hard': ['in_f_hard', 't3_f_hard'], 'f_cond': ['in_f_cond', 't3_f_cond'],
    'f_sio2': ['in_f_sio2', 't3_f_sio2'], 'f_cl': ['in_f_cl', 't3_f_cl'], 'f_so4': ['in_f_so4', 't3_f_so4'],
    'f_sulfide': ['in_f_sulfide', 't3_f_sulfide'],
    'p_ph': ['in_p_ph', 't3_p_ph'], 'p_turb': ['in_p_turb', 't3_p_turb'], 'p_cond': ['in_p_cond', 't3_p_cond'],
    'p_tds': ['in_p_tds', 't3_p_tds'], 'p_alk': ['in_p_alk', 't3_p_alk'], 'p_ca': ['in_p_ca', 't3_p_ca'],
    'p_mg': ['in_p_mg', 't3_p_mg'], 'p_hard': ['in_p_hard', 't3_p_hard'], 'p_iron': ['in_p_iron', 't3_p_iron'],
    'p_sio2': ['in_p_sio2', 't3_p_sio2'], 'p_cl': ['in_p_cl', 't3_p_cl'], 'p_so4': ['in_p_so4', 't3_p_so4'],
    'b_ph': ['in_b_ph'], 'b_turb': ['in_b_turb'], 'b_cond': ['in_b_cond'], 'b_tds': ['in_b_tds'],
    'b_alk': ['in_b_alk'], 'b_ca': ['in_b_ca'], 'b_mg': ['in_b_mg'], 'b_hard': ['in_b_hard'],
    'b_sio2': ['in_b_sio2'], 'b_cl': ['in_b_cl'],
    'chem_anti_ppm': ['in_anti_ppm', 't4_anti_ppm', 't5_anti'], 'chem_anti_cons': ['in_anti_cons', 't4_anti_cons'],
    'chem_foam_ppm': ['in_foam_ppm', 't4_foam_ppm'], 'chem_foam_cons': ['in_foam_cons', 't4_foam_cons'],
    'remarks': ['in_remarks', 't0_remarks'],
    'area_1st': ['in_area_1st', 't2_area_1st'], 'area_overall': ['in_area_overall', 't2_area_overall'],
    'sw_lower': ['in_sw_low'], 'cond_flow': ['in_cond_flow'], 'cond_temp': ['in_cond_temp'], 'htc1_cond_temp': [], 'htco_cond_temp': [], 'cond_cond': ['in_cond_cond'],
    'sw_out_t': ['in_sw_out'], 'cw_supply': ['in_cw_supply'], 'cw_return': ['in_cw_return'], 'cw_flow': ['in_cw_flow']
}

LATENT_HEAT_STEAM_KJ_KG = 2330.0

def generate_daily_csv(date, ops, w_data, chem_data, mra, extra_tags):
    data_dict = {
        "Date": date.strftime('%d-%m-%Y'),
        "Sea Water Upper": ops['SW_Feed_1st'], "Sea Water Lower": extra_tags['sw_lower'],
        "Sea Water Feed": ops['SW Total'], "Sea Water Pressure": extra_tags['sw_press'], 
        "Brine Water Return": ops['Brine Return'], "Desal production": ops['Desal'], 
        "LP Steam consumption": ops['Steam'], "LP Steam Pressure": extra_tags['stm_press'],
        "Condensate Return": extra_tags['cond_flow'], "condensate temp": extra_tags['cond_temp'], "Condensate Conductivity": extra_tags['cond_cond'],
        "1st Effect Vapour Temp": ops['Stm In_1st'], "1st effect brine temp": ops['Brine_1st'],
        "11th Effect Brine Temp": extra_tags['brine_11'], "Feed Temp to Cold Group": extra_tags['feed_cold'],
        "Intermediate Effects Avg Brine Temp": extra_tags['mid_effects_temp'],
        "Delta T": ops['dt_1st'], "1st effect vapour pressure": ops['Press_1st'],
        "Brine Discharge Temp": ops['Brine Out_overall'], "Brine Discharge Pressure": extra_tags['brine_press'],
        "Sea Water cond I/L temp": ops['SW In_overall'], "Sea Water Condenser O/L Temp": extra_tags['sw_out_t'],
        "CW supply": extra_tags['cw_supply'], "CW Return": extra_tags['cw_return'], "CW Flow": extra_tags['cw_flow'],
        "Gross production": ops['Gross Prod'], "Recovery (%)": round(ops['Recovery'], 2),
        "GOR": round(ops['GOR'], 2), "STEC": round(ops['STEC'], 2), "Overall HTC": round(ops['htc_overall'], 2),
        "1st Effect HTC": round(ops['htc_1st'], 2), "Residual": round(mra['Residual'], 2),
        "Antiscalant Dosing (PPM)": chem_data['anti_ppm'], "Antiscalant (kg)": chem_data['anti_cons'],
        "Antifoam Dosing (PPM)": chem_data['foam_ppm'], "Antifoam (kg)": chem_data['foam_cons'],
        "Remarks": extra_tags['remarks']
    }
    for cat in ['Feed', 'Product']:
        for param, details in w_data[cat].items(): data_dict[details['db_col']] = details['val']
        
    df = pd.DataFrame([data_dict])
    return df.to_csv(index=False).encode('utf-8')

# =====================================================================================
# REPORTING ENGINE
# Reports are the deliverable Reliance actually reads, so they must explain WHY the plant
# performed as it did - not just dump numbers. Everything below is built around that:
# consistent rounding, deliberate page breaks, and a narrative layer that interprets
# variance from SOR, assesses fouling, and recommends action.
# =====================================================================================

# Stamped into every generated report so it is immediately obvious which build produced a file.
# If a downloaded report does not show this version in its footer, the app is running older code.
REPORT_VERSION = "v2.0 (analysis + charts)"

# SOR reference values used for variance analysis and interpretation.
SOR_REF = {
    'GOR': 11.4,
    'Steam_TPH': 76.94,
    'Steam_Press': 4.3,
    'Steam_Inlet_Temp': 176.0,
    'Vapour_Temp_1st': 68.47,
    'Anti_PPM': 10.5,
    'Foam_PPM': 0.16,
}


def _num(v, dp=1, dash_if_zero=False):
    """Format a number for a report. Raw floats like 795.3299999999999 are unreadable in a
    client document, so every printed value goes through here."""
    try:
        f = float(v)
        if pd.isna(f):
            return "-"
        if dash_if_zero and f == 0:
            return "-"
        return f"{f:,.{dp}f}"
    except (ValueError, TypeError):
        s = str(v).strip()
        return s if s else "-"


def _safe_pct(actual, ref):
    """Percentage deviation of actual from a reference, or None if not computable."""
    try:
        a, r = float(actual), float(ref)
        if r == 0 or pd.isna(a) or pd.isna(r):
            return None
        return (a - r) / r * 100.0
    except (ValueError, TypeError):
        return None


def _para(doc, text, bold=False, italic=False, size=None, color=None, align=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if align is not None:
        p.alignment = align
    return p


def _table(doc, headers, rows, widths=None):
    """Build a Table Grid with a bolded header row."""
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        r = cell.paragraphs[0].add_run(str(h))
        r.bold = True
    for row in rows:
        rc = t.add_row().cells
        for i, val in enumerate(row):
            rc[i].text = str(val)
    if widths:
        for r_ in t.rows:
            for i, w in enumerate(widths):
                if i < len(r_.cells):
                    r_.cells[i].width = Inches(w)
    return t


def _chart(plot_fn, figsize=(7.2, 2.9)):
    """Render a matplotlib chart to PNG bytes for embedding. Returns None on any failure so a
    charting problem can never prevent the report itself from being produced."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
        fig, ax = plt.subplots(figsize=figsize, dpi=150)
        plot_fn(ax, plt)
        fig.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format="png", bbox_inches="tight")
        plt.close(fig)
        buf.seek(0)
        return buf
    except Exception:
        return None


# -------------------------------------------------------------------------------------
# INTERPRETATION ENGINE
# -------------------------------------------------------------------------------------
def build_interpretation(ops, mra, chem_ppm=None):
    """Turn the day's numbers into engineering findings, variance explanations and actions.

    Returns a dict with: status, headline, variance (list), fouling (dict), actions (list).
    Kept separate from document formatting so the daily and monthly reports, and the on-screen
    view, all describe the plant the same way.
    """
    out = {'variance': [], 'actions': [], 'fouling': {}, 'status': 'normal', 'headline': ''}

    gor = float(ops.get('GOR') or 0)
    steam = float(ops.get('Steam') or 0)
    gross = float(ops.get('Gross Prod') or 0)
    htc_o = float(ops.get('htc_overall') or 0)
    htc_1 = float(ops.get('htc_1st') or 0)
    recovery = float(ops.get('Recovery') or 0)

    # ---- Efficiency vs SOR ----
    gor_dev = _safe_pct(gor, SOR_REF['GOR'])
    steam_dev = _safe_pct(steam, SOR_REF['Steam_TPH'])

    if gor > 0 and gor_dev is not None:
        if gor_dev < -5:
            out['variance'].append(
                f"Gain Output Ratio was {gor:.2f} against an SOR baseline of {SOR_REF['GOR']:.2f} "
                f"({gor_dev:+.1f}%). GOR is distillate produced per unit of steam, so a shortfall here "
                f"is a thermal efficiency loss rather than a throughput limitation - the unit consumed "
                f"more steam than it should have for the water it made."
            )
        elif gor_dev > 2:
            out['variance'].append(
                f"Gain Output Ratio was {gor:.2f} against an SOR baseline of {SOR_REF['GOR']:.2f} "
                f"({gor_dev:+.1f}%), indicating the unit converted steam to distillate more efficiently "
                f"than the reference condition."
            )
        else:
            out['variance'].append(
                f"Gain Output Ratio was {gor:.2f} against an SOR baseline of {SOR_REF['GOR']:.2f} "
                f"({gor_dev:+.1f}%), essentially in line with the reference condition."
            )

    # ---- Throughput: separate a steam-supply constraint from an efficiency loss ----
    if steam > 0 and steam_dev is not None and steam_dev < -5:
        implied = abs(steam_dev) / 100.0 * SOR_REF['Steam_TPH'] * (gor if gor > 0 else SOR_REF['GOR'])
        out['variance'].append(
            f"LP steam consumption was {steam:.1f} TPH against an SOR of {SOR_REF['Steam_TPH']:.1f} TPH "
            f"({steam_dev:+.1f}%). At the GOR achieved, this reduced steam input alone accounts for "
            f"approximately {implied:.0f} m³/h of the production shortfall. Production below SOR on this "
            f"day is therefore primarily a steam availability constraint, not a fault in the unit."
        )
    elif steam > 0 and steam_dev is not None and steam_dev > 5:
        out['variance'].append(
            f"LP steam consumption was {steam:.1f} TPH against an SOR of {SOR_REF['Steam_TPH']:.1f} TPH "
            f"({steam_dev:+.1f}%). Steam input above the reference with production at or below SOR points "
            f"to reduced heat transfer efficiency rather than a supply limitation."
        )

    if recovery > 0:
        out['variance'].append(
            f"Recovery was {recovery:.1f}% of seawater feed converted to product."
        )

    # ---- Fouling assessment: HTC drop against the SOR-clean baseline ----
    d_o = _safe_pct(htc_o, HTC_OVERALL_U_SOR) if htc_o > 0 else None
    d_1 = _safe_pct(htc_1, HTC_1ST_U_SOR) if htc_1 > 0 else None
    out['fouling'] = {'overall_dev': d_o, 'first_dev': d_1, 'htc_overall': htc_o, 'htc_1st': htc_1}

    worst = min([x for x in (d_o, d_1) if x is not None], default=None)
    if worst is None:
        out['fouling']['verdict'] = "Heat transfer coefficients were not calculable for this day, so no fouling assessment can be made."
        out['fouling']['tier'] = 'unknown'
    elif worst >= -5:
        out['fouling']['verdict'] = (
            "Heat transfer coefficients are within 5% of the post-clean SOR baseline. The tube surfaces "
            "are effectively clean and no chemical or mechanical intervention is indicated."
        )
        out['fouling']['tier'] = 'clean'
    elif worst >= -15:
        out['fouling']['verdict'] = (
            "Heat transfer coefficients have declined between 5% and 15% from the post-clean SOR baseline. "
            "This is the signature of early-stage scale formation. It is recoverable at this stage through "
            "dosing control and does not yet warrant taking the unit offline."
        )
        out['fouling']['tier'] = 'early'
    else:
        out['fouling']['verdict'] = (
            "Heat transfer coefficients have declined more than 15% from the post-clean SOR baseline. This "
            "represents established scale on the heat transfer surfaces, which will continue to depress "
            "output and raise specific energy consumption until it is removed."
        )
        out['fouling']['tier'] = 'significant'

    # ---- MRA: measured output against what the calibrated model expects ----
    pred = float(mra.get('Predicted') or 0)
    resid = float(mra.get('Residual') or 0)
    diff_pct = (resid / pred * 100.0) if pred > 0 else 0.0
    out['mra_diff_pct'] = diff_pct
    if pred > 0:
        out['mra_text'] = (
            f"The Multiple Regression Analysis model, calibrated on this unit's own historical operation, "
            f"predicted {pred:.1f} m³/h under the day's operating conditions against an actual "
            f"{float(mra.get('Actual') or 0):.1f} m³/h - a deviation of {diff_pct:+.1f}%. Because the model "
            f"already accounts for steam rate, seawater temperature, pressure and dosing, a persistent "
            f"negative deviation isolates performance loss that operating conditions do not explain, which "
            f"in a thermal desalination unit is characteristically scale."
        )
    else:
        out['mra_text'] = "The MRA model did not return a valid prediction for this day, so no model-based fouling assessment is available."

    # ---- Overall status ----
    if diff_pct <= -5.0 or out['fouling'].get('tier') == 'significant':
        out['status'] = 'action'
        out['headline'] = "Cleaning intervention recommended"
    elif diff_pct <= -4.0 or out['fouling'].get('tier') == 'early':
        out['status'] = 'watch'
        out['headline'] = "Early performance deviation - corrective dosing advised"
    else:
        out['status'] = 'normal'
        out['headline'] = "Operating within expected performance envelope"

    # ---- Recommendations ----
    if out['status'] == 'action':
        out['actions'].append(
            "Schedule a chemical (acid) clean of the heat transfer surfaces at the next available "
            "production window. The combined HTC decline and MRA deviation indicate scale that dosing "
            "alone will not reverse."
        )
        out['actions'].append(
            "Until cleaning is carried out, expect elevated specific energy consumption; steam demand per "
            "m³ of product will remain above the SOR reference."
        )
    elif out['status'] == 'watch':
        out['actions'].append(
            "Increase antiscalant dosing toward the SOR target and hold it there while monitoring HTC daily. "
            "Early-stage scale is normally reversible at this point without taking the unit offline."
        )
        out['actions'].append(
            "Re-assess after seven days. If HTC continues to decline, plan a cleaning window rather than "
            "continuing to increase chemical dosage."
        )
    else:
        out['actions'].append(
            "Continue current operating and dosing regime. No intervention is indicated."
        )

    if chem_ppm is not None:
        try:
            ppm = float(chem_ppm)
            if 0 < ppm < SOR_REF['Anti_PPM'] * 0.8:
                out['actions'].append(
                    f"Antiscalant residual measured {ppm:.2f} ppm against an SOR target of "
                    f"{SOR_REF['Anti_PPM']:.2f} ppm. Under-dosing at this level materially raises scaling "
                    f"risk on the first effect and should be corrected irrespective of current HTC."
                )
            elif ppm > SOR_REF['Anti_PPM'] * 1.3:
                out['actions'].append(
                    f"Antiscalant residual measured {ppm:.2f} ppm against an SOR target of "
                    f"{SOR_REF['Anti_PPM']:.2f} ppm. Dosing above requirement adds chemical cost without "
                    f"proportional scale protection and can be trimmed."
                )
        except (ValueError, TypeError):
            pass

    return out


# -------------------------------------------------------------------------------------
# DAILY REPORT
# -------------------------------------------------------------------------------------
def generate_comprehensive_report(date, ops, sor_dfs, w_data, chem_data, mra, skip_wq, remarks):
    doc = Document()
    for s in doc.sections:
        s.left_margin = Inches(0.8); s.right_margin = Inches(0.8)
        s.top_margin = Inches(0.7); s.bottom_margin = Inches(0.7)

    interp = build_interpretation(ops, mra, chem_data.get('anti_ppm') if isinstance(chem_data, dict) else None)

    # ---- Cover block ----
    doc.add_heading('MED-4 Daily Performance Report', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para(doc, 'Reliance Industries Limited  |  Multi-Effect Distillation Unit 4',
          size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, f"Reporting date: {date.strftime('%d %B %Y')}", bold=True, size=11,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, 'Prepared by Chembond Water Technologies Limited', italic=True, size=10,
          align=WD_ALIGN_PARAGRAPH.CENTER)

    color = {'normal': RGBColor(0, 128, 0), 'watch': RGBColor(216, 130, 43), 'action': RGBColor(200, 40, 40)}[interp['status']]
    _para(doc, f"Assessment: {interp['headline']}", bold=True, size=12, color=color,
          align=WD_ALIGN_PARAGRAPH.CENTER)

    # ---- 1. Executive summary ----
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        f"On {date.strftime('%d %B %Y')} the MED-4 unit produced {_num(ops.get('Gross Prod'))} m³/h gross, "
        f"at a Gain Output Ratio of {_num(ops.get('GOR'), 2)} and a specific thermal energy consumption of "
        f"{_num(ops.get('STEC'))} kWh per tonne of distillate. Recovery from seawater feed was "
        f"{_num(ops.get('Recovery'))}%. Overall plant heat transfer coefficient was "
        f"{_num(ops.get('htc_overall'), 2)} W/m²K and first effect heat transfer coefficient "
        f"{_num(ops.get('htc_1st'))} W/m²K."
    )
    doc.add_paragraph(interp['fouling'].get('verdict', ''))

    doc.add_heading('Key indicators', level=2)
    _table(doc,
           ['Indicator', 'Value', 'SOR Reference', 'Deviation'],
           [
               ['Gross production (m³/h)', _num(ops.get('Gross Prod')), '-', '-'],
               ['Gain Output Ratio', _num(ops.get('GOR'), 2), f"{SOR_REF['GOR']:.2f}",
                f"{_safe_pct(ops.get('GOR'), SOR_REF['GOR']):+.1f}%" if _safe_pct(ops.get('GOR'), SOR_REF['GOR']) is not None else '-'],
               ['STEC (kWh/tonne)', _num(ops.get('STEC')), '-', '-'],
               ['Overall HTC (W/m²K)', _num(ops.get('htc_overall'), 2), f"{HTC_OVERALL_U_SOR:.2f}",
                f"{interp['fouling']['overall_dev']:+.1f}%" if interp['fouling'].get('overall_dev') is not None else '-'],
               ['1st Effect HTC (W/m²K)', _num(ops.get('htc_1st')), f"{HTC_1ST_U_SOR:.1f}",
                f"{interp['fouling']['first_dev']:+.1f}%" if interp['fouling'].get('first_dev') is not None else '-'],
               ['Recovery (%)', _num(ops.get('Recovery')), '-', '-'],
           ],
           widths=[2.6, 1.3, 1.3, 1.2])

    # ---- 2. Performance interpretation (the part that was missing entirely) ----
    doc.add_heading('2. Performance Analysis', level=1)
    doc.add_paragraph(
        "The following explains how the unit performed relative to its System Operating Reference (SOR) "
        "baseline, and what accounts for any deviation."
    )
    for line in interp['variance']:
        doc.add_paragraph(line, style='List Bullet')

    doc.add_heading('Thermal integrity and fouling', level=2)
    doc.add_paragraph(interp['fouling'].get('verdict', ''))
    _table(doc,
           ['Heat exchanger', 'Measured (W/m²K)', 'Post-clean SOR', 'Deviation'],
           [
               ['First effect', _num(ops.get('htc_1st')), f"{HTC_1ST_U_SOR:.1f}",
                f"{interp['fouling']['first_dev']:+.1f}%" if interp['fouling'].get('first_dev') is not None else '-'],
               ['Overall plant', _num(ops.get('htc_overall'), 2), f"{HTC_OVERALL_U_SOR:.2f}",
                f"{interp['fouling']['overall_dev']:+.1f}%" if interp['fouling'].get('overall_dev') is not None else '-'],
           ],
           widths=[1.8, 1.7, 1.5, 1.4])
    doc.add_paragraph(
        "Heat transfer coefficient is calculated on the steam condensation basis, U = Q / (A x LMTD), using "
        "the log mean temperature difference across each exchanger. A falling coefficient at constant duty "
        "indicates an increasing resistance on the tube surface, which is the direct physical signature of scale."
    )

    doc.add_heading('Model-based assessment (MRA)', level=2)
    doc.add_paragraph(interp['mra_text'])

    doc.add_heading('3. Recommendations', level=1)
    for i, a in enumerate(interp['actions'], 1):
        doc.add_paragraph(f"{i}. {a}")

    # ---- 4. SOR matrix (reference detail, moved onto its own page) ----
    doc.add_page_break()
    doc.add_heading('4. SOR Performance Matrix', level=1)
    doc.add_paragraph(
        "Full parameter-by-parameter comparison against the System Operating Reference. Positive deviations "
        "indicate operation above the reference condition."
    )
    for section_name, df in sor_dfs.items():
        doc.add_heading(str(section_name), level=2)
        rows = []
        for _, row in df.iterrows():
            rows.append([
                str(row.get('Parameter', '')),
                str(row.get('UOM', '')),
                str(row.get('Design', '')),
                _num(row.get('SOR Base'), 2),
                _num(row.get('Actual'), 2),
                _num(row.get('Difference'), 2),
            ])
        _table(doc, ['Parameter', 'UOM', 'Design', 'SOR Base', 'Actual', 'Deviation'], rows,
               widths=[2.2, 1.0, 0.9, 0.9, 0.9, 0.9])

    # ---- 5. Water quality ----
    doc.add_page_break()
    doc.add_heading('5. Water Quality', level=1)
    if skip_wq:
        doc.add_paragraph("Laboratory water quality parameters were not recorded for this operational day.")
    else:
        rows, exceedances = [], []
        for stream_label, key in (('Sea Water Feed', 'Feed'), ('Desal Product', 'Product')):
            for param, data in w_data.get(key, {}).items():
                val = data.get('val')
                lo, hi = data.get('min'), data.get('max')
                ok = True
                try:
                    if val is not None and float(val) != 0:
                        ok = float(lo) <= float(val) <= float(hi)
                except (ValueError, TypeError):
                    ok = True
                if not ok:
                    exceedances.append(f"{stream_label} {param} ({_num(val, 2)} against {_num(lo, 2)}-{_num(hi, 2)})")
                rows.append([str(param), stream_label, f"{_num(lo, 2)} - {_num(hi, 2)}",
                             _num(val, 2, dash_if_zero=True), 'Within spec' if ok else 'Out of spec'])
        _table(doc, ['Parameter', 'Stream', 'Specification', 'Result', 'Status'], rows,
               widths=[1.8, 1.4, 1.5, 1.0, 1.1])
        if exceedances:
            doc.add_paragraph(
                "The following parameters fell outside specification: " + "; ".join(exceedances) +
                ". Feed side exceedances raise scaling risk and should be reviewed against pretreatment "
                "performance; product side exceedances affect downstream water quality directly."
            )
        else:
            doc.add_paragraph("All recorded parameters were within specification.")

    if remarks and str(remarks).strip():
        doc.add_heading('6. Operator Remarks', level=1)
        doc.add_paragraph(str(remarks))

    _para(doc, f"Report generated by the Chembond MED Performance Monitoring System  |  Report engine {REPORT_VERSION}",
          italic=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


# -------------------------------------------------------------------------------------
# MONTHLY REPORT
# -------------------------------------------------------------------------------------
def generate_period_report(df_period, period_label, period_kind="Period"):
    """Build the aggregate performance report for ANY span of days.

    Previously this was hard-wired to a calendar month (it took a month name and a year and printed
    "Monthly" throughout). Reliance asked to be able to pull a report for an arbitrary start/end date,
    so the period is now passed in as a ready-made label and the wording adapts: 'Monthly' when a whole
    calendar month was selected, otherwise 'Period'. Nothing about the underlying maths changed - it
    always operated on whatever rows it was handed.
    """
    doc = Document()
    for s in doc.sections:
        s.left_margin = Inches(0.8); s.right_margin = Inches(0.8)
        s.top_margin = Inches(0.7); s.bottom_margin = Inches(0.7)

    d = df_period.copy()
    for c in ['Gross production', 'GOR', 'STEC', 'Overall HTC', '1st Effect HTC', 'Desal production',
              'LP Steam consumption', 'Recovery', 'Anti_PPM']:
        if c in d.columns:
            d[c] = pd.to_numeric(d[c], errors='coerce')
        else:
            d[c] = np.nan
    if 'Date' in d.columns:
        d['_d'] = standardize_dates(d['Date'])
        d = d.dropna(subset=['_d']).sort_values('_d')

    # Only days the unit actually ran should shape the averages.
    run = d[d['Gross production'].fillna(0) > 0]
    n_days = len(run)
    period = ""
    if '_d' in d.columns and not d.empty:
        period = f"{d['_d'].min().strftime('%d %B %Y')} to {d['_d'].max().strftime('%d %B %Y')}"

    # ---- Cover: the month is now unmistakable ----
    doc.add_heading(f'MED-4 {period_kind} Performance Report', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para(doc, period_label, bold=True, size=20, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, 'Reliance Industries Limited  |  Multi-Effect Distillation Unit 4', size=11,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    if period:
        _para(doc, f"Period covered: {period}  ({n_days} operating days)", size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    _para(doc, 'Prepared by Chembond Water Technologies Limited', italic=True, size=10,
          align=WD_ALIGN_PARAGRAPH.CENTER)

    if n_days == 0:
        doc.add_paragraph(f"No operating days with recorded production were found in {period_label}.")
        bio = io.BytesIO(); doc.save(bio); return bio.getvalue()

    avg = lambda c: run[c].mean()
    gor_m, htc_o_m, htc_1_m = avg('GOR'), avg('Overall HTC'), avg('1st Effect HTC')

    # Trend across the month: compare first third against last third.
    trend_txt = ""
    if n_days >= 6:
        k = max(2, n_days // 3)
        for label, col, ref in (('Overall HTC', 'Overall HTC', HTC_OVERALL_U_SOR),
                                ('First effect HTC', '1st Effect HTC', HTC_1ST_U_SOR)):
            a, b = run[col].head(k).mean(), run[col].tail(k).mean()
            if pd.notna(a) and pd.notna(b) and a > 0:
                ch = (b - a) / a * 100
                direction = "declined" if ch < -1 else "improved" if ch > 1 else "held steady"
                trend_txt += (f"{label} {direction} across the period, averaging {a:.2f} in the opening days "
                              f"against {b:.2f} in the closing days ({ch:+.1f}%). ")

    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        f"Over {n_days} operating days across {period_label}, MED-4 averaged "
        f"{_num(avg('Gross production'))} m³/h gross production at a Gain Output Ratio of "
        f"{_num(gor_m, 2)} against an SOR reference of {SOR_REF['GOR']:.2f}. Average specific thermal "
        f"energy consumption was {_num(avg('STEC'))} kWh per tonne. Heat transfer coefficients averaged "
        f"{_num(htc_o_m, 2)} W/m²K overall and {_num(htc_1_m)} W/m²K on the first effect."
    )
    if trend_txt:
        doc.add_paragraph(trend_txt.strip())

    interp_m = build_interpretation(
        {'GOR': gor_m, 'Steam': avg('LP Steam consumption'), 'Gross Prod': avg('Gross production'),
         'htc_overall': htc_o_m, 'htc_1st': htc_1_m, 'Recovery': avg('Recovery'), 'STEC': avg('STEC')},
        {'Predicted': 0, 'Residual': 0, 'Actual': 0},
        avg('Anti_PPM') if pd.notna(avg('Anti_PPM')) else None
    )
    doc.add_paragraph(interp_m['fouling'].get('verdict', ''))

    # ---- 2. Aggregate table ----
    doc.add_heading(f'2. {period_kind} Aggregates', level=1)
    rows = []
    for name, col, dp in (
        ('Gross production (m³/h)', 'Gross production', 1),
        ('Desal production (m³/h)', 'Desal production', 1),
        ('LP steam consumption (TPH)', 'LP Steam consumption', 1),
        ('Gain Output Ratio', 'GOR', 2),
        ('STEC (kWh/tonne)', 'STEC', 1),
        ('Overall HTC (W/m²K)', 'Overall HTC', 2),
        ('1st Effect HTC (W/m²K)', '1st Effect HTC', 1),
        ('Recovery (%)', 'Recovery', 1),
    ):
        s = run[col].dropna()
        if s.empty:
            rows.append([name, '-', '-', '-', '-'])
        else:
            rows.append([name, _num(s.min(), dp), _num(s.mean(), dp), _num(s.max(), dp), _num(s.std(), dp)])
    _table(doc, ['Metric', 'Minimum', 'Average', 'Maximum', 'Std Dev'], rows,
           widths=[2.4, 1.1, 1.1, 1.1, 1.1])

    # ---- 3. Charts ----
    doc.add_page_break()
    doc.add_heading('3. Performance Trends', level=1)
    doc.add_paragraph(
        "The charts below show how the unit behaved across the reporting period. Dashed red lines mark the System "
        "Operating Reference, which represents the unit's expected performance in clean condition."
    )
    x = run['_d'] if '_d' in run.columns else range(len(run))

    def _prod(ax, plt):
        ax.plot(x, run['Gross production'], marker='o', ms=3, lw=1.4, color='#0072FF', label='Gross production')
        ax.set_ylabel('m³/h'); ax.set_title('Gross Production', fontsize=10)
        ax.grid(alpha=.3); ax.tick_params(labelsize=7)
        plt.setp(ax.get_xticklabels(), rotation=30, ha='right')
    def _gor(ax, plt):
        ax.plot(x, run['GOR'], marker='o', ms=3, lw=1.4, color='#00A06A', label='GOR')
        ax.axhline(SOR_REF['GOR'], ls='--', color='red', lw=1, label=f"SOR {SOR_REF['GOR']}")
        ax.set_ylabel('GOR'); ax.set_title('Gain Output Ratio vs SOR', fontsize=10)
        ax.grid(alpha=.3); ax.legend(fontsize=7); ax.tick_params(labelsize=7)
        plt.setp(ax.get_xticklabels(), rotation=30, ha='right')
    def _htco(ax, plt):
        ax.plot(x, run['Overall HTC'], marker='o', ms=3, lw=1.4, color='#0072FF')
        ax.axhline(HTC_OVERALL_U_SOR, ls='--', color='red', lw=1, label=f"SOR {HTC_OVERALL_U_SOR:.2f}")
        ax.set_ylabel('W/m²K'); ax.set_title('Overall HTC vs SOR (fouling indicator)', fontsize=10)
        ax.grid(alpha=.3); ax.legend(fontsize=7); ax.tick_params(labelsize=7)
        plt.setp(ax.get_xticklabels(), rotation=30, ha='right')
    def _htc1(ax, plt):
        ax.plot(x, run['1st Effect HTC'], marker='o', ms=3, lw=1.4, color='#D9822B')
        ax.axhline(HTC_1ST_U_SOR, ls='--', color='red', lw=1, label=f"SOR {HTC_1ST_U_SOR:.0f}")
        ax.set_ylabel('W/m²K'); ax.set_title('First Effect HTC vs SOR (fouling indicator)', fontsize=10)
        ax.grid(alpha=.3); ax.legend(fontsize=7); ax.tick_params(labelsize=7)
        plt.setp(ax.get_xticklabels(), rotation=30, ha='right')

    any_chart = False
    for fn, caption in ((_prod, "Daily gross production."),
                        (_gor, "Thermal efficiency against the SOR reference."),
                        (_htco, "Overall heat transfer coefficient. A sustained downward slope indicates scale."),
                        (_htc1, "First effect heat transfer coefficient, the earliest indicator of scaling.")):
        buf = _chart(fn)
        if buf:
            doc.add_picture(buf, width=Inches(6.4))
            _para(doc, caption, italic=True, size=8)
            any_chart = True
    if not any_chart:
        doc.add_paragraph("Charts could not be rendered in this environment; the tabulated values above carry the same information.")

    # ---- 4. Assessment and recommendations ----
    doc.add_page_break()
    doc.add_heading('4. Assessment and Recommendations', level=1)
    for line in interp_m['variance']:
        doc.add_paragraph(line, style='List Bullet')
    doc.add_heading('Recommended actions', level=2)
    for i, a in enumerate(interp_m['actions'], 1):
        doc.add_paragraph(f"{i}. {a}")

    doc.add_heading('Chembond scope', level=2)
    doc.add_paragraph(
        "Chembond Water Technologies maintains continuous performance monitoring of MED-4, covering thermal "
        "efficiency, heat transfer integrity, water chemistry and antiscalant programme effectiveness. The "
        "analysis above is generated from the unit's own operating data and calibrated against its "
        "post-cleaning reference condition, enabling scale formation to be identified and corrected before "
        "it materially affects output."
    )

    _para(doc, f"Report generated by the Chembond MED Performance Monitoring System  |  Report engine {REPORT_VERSION}",
          italic=True, size=8, align=WD_ALIGN_PARAGRAPH.CENTER)

    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def generate_monthly_report(df_month, month_str, year_str):
    """Thin wrapper kept so the existing calendar-month button keeps its exact previous behaviour."""
    return generate_period_report(df_month, f"{month_str} {year_str}", period_kind="Monthly")


def render_med_suite(db_conn, LOCAL_DB_FILE, LOCAL_CONFIG_FILE, AI_MODEL_FILE, save_database, save_config,
                     render_chatbot, SKLEARN_INSTALLED, XGB_INSTALLED, PIL_INSTALLED,
                     save_model_blob=None, load_model_blob=None):
    
    # MED Internal State Setup
    if 'vars' not in st.session_state: st.session_state.vars = DEFAULTS.copy()
    for k, v in DEFAULTS.items():
        if k not in st.session_state.vars: st.session_state.vars[k] = v

    def sync_var(var_name, source_key):
        st.session_state.vars[var_name] = st.session_state[source_key]
        for target_key in SYNC_MAP.get(var_name, []):
            if target_key != source_key: st.session_state[target_key] = st.session_state[source_key]

    def get_v(var_name): return st.session_state.vars[var_name]

    if 'sync_initialized' not in st.session_state:
        for var_name, keys in SYNC_MAP.items():
            for k in keys: 
                if k not in st.session_state: st.session_state[k] = st.session_state.vars[var_name]
        st.session_state.sync_initialized = True

    if 'shared_effect_df' not in st.session_state or 'Live Vapor (°C)' not in st.session_state.shared_effect_df.columns:
        st.session_state.shared_effect_df = pd.DataFrame({
            "Effect ID": [f"Effect {i}" for i in range(1, 12)], 
            "Live Vapor (°C)": [np.nan] * 11, 
            "Live Brine (°C)": [np.nan] * 11
        })

    med_unit_choice = st.sidebar.selectbox("Select Active Unit Train", [f"MED-{unit_idx}" for unit_idx in range(1, 12)], index=3)
    if med_unit_choice != "MED-4":
        st.title(f"{med_unit_choice}")
        st.info(f"System data hooks for {med_unit_choice} are under configuration. Diagnostic dashboard metrics will become available upon plant startup.")
        render_chatbot()
        return
            
    st.sidebar.divider()
    # Default to the most recent date that actually has data, not today. Defaulting to today meant
    # that whenever the registry lagged the calendar (e.g. data ends 6 Jul, today is 14 Jul) the app
    # opened on a date with no record, reset every field to 0, and showed HTC/KPIs as 0 - even though
    # the registry and the trend charts were perfectly fine.
    _default_date = datetime.date.today()
    _logs0 = st.session_state.get('daily_logs')
    if _logs0 is not None and not _logs0.empty and 'Date' in _logs0.columns:
        _d = standardize_dates(_logs0['Date']).dropna()
        if not _d.empty:
            _default_date = _d.max().date()
    log_date = st.sidebar.date_input("Date", _default_date, format="DD/MM/YYYY")
    log_date_str = log_date.strftime('%Y-%m-%d')
    if _default_date != datetime.date.today() and log_date == _default_date:
        st.sidebar.caption(f"Showing latest record ({_default_date.strftime('%d-%m-%Y')}).")
    
    if 'last_selected_date' not in st.session_state: 
        st.session_state.last_selected_date = None

    if log_date_str != st.session_state.last_selected_date:
        st.session_state.last_selected_date = log_date_str
        date_found = False

        # ALWAYS clear every measured field first, then overlay whatever this date's record actually
        # contains. Resetting only when no row exists was unsafe: a row that existed but was blank or
        # only partly filled set date_found=True, skipped the reset, loaded nothing, and left the
        # PREVIOUS date's readings on screen looking like real data for the selected day.
        # Plant CONSTANTS are excluded - the heat transfer areas are fixed equipment geometry, not
        # daily readings, and zeroing them would force HTC to 0 even once valid data is entered.
        PLANT_CONSTANTS = ('area_1st', 'area_overall')
        for var_key, default_val in DEFAULTS.items():
            if var_key in PLANT_CONSTANTS:
                continue
            zero_val = 0.0 if isinstance(default_val, (int, float)) and not isinstance(default_val, bool) else default_val
            if var_key in ('remarks',): zero_val = ""
            if var_key in ('skip_eff', 'skip_wq'): zero_val = False
            st.session_state.vars[var_key] = zero_val
            for tk in SYNC_MAP.get(var_key, []):
                st.session_state[tk] = zero_val

        if not st.session_state.daily_logs.empty and 'Date' in st.session_state.daily_logs.columns:
            # CORE FIX: Standardize all registry dates right now, extract as safe strings
            db_dates_parsed = standardize_dates(st.session_state.daily_logs['Date'])
            db_dates = db_dates_parsed.dt.strftime('%Y-%m-%d').values
            
            if log_date_str in db_dates:
                date_found = True
                row_idx = np.where(db_dates == log_date_str)[0][-1]
                row = st.session_state.daily_logs.iloc[row_idx]
                
                db_to_var_mapping = {
                    'gross': ['Gross production'], 
                    'stm_press': ['LP Steam Pressure'],
                    'sw_press': ['Sea Water Pressure'],
                    'sw_upper': ['Sea Water Upper'], 'sw_lower': ['Sea Water Lower'],
                    'cond_cond': ['Condensate Conductivity'],
                    'sw_out_t': ['Sea Water Condenser O/L Temp'], 
                    'cw_supply': ['CW supply'], 'cw_return': ['CW Return'], 'cw_flow': ['CW Flow'],
                    'chem_anti_cons': ['Antiscalant (kg)'], 'chem_foam_cons': ['Antifoam (kg)'], 
                    'mra_press': ['1st effect vapour pressure'], 
                    'brine_11': ['11th Effect Brine Temp'],
                    'brine_ret': ['Brine Water Return'], 'brine_press': ['Brine Discharge Pressure'],
                    'chem_anti_ppm': ['Anti_PPM'], 'chem_foam_ppm': ['Foam_PPM'],
                    'sw_in_t': ['Sea Water cond I/L temp'], 
                    'vap_out_t': ['Vap_Out_Temp'], 
                    'remarks': ['Remarks'], 'area_1st': ['Area_1st'], 'area_overall': ['Area_Overall'],

                    # --- HTC-critical inputs: list EVERY column that can carry the value, in priority
                    # order, because the same physical reading is stored under different names depending
                    # on which uploader wrote it. Previously these pointed at a single Operational column:
                    #   - 'Feed Temp to Cold Group' is no longer written by ANY uploader (the Overall HTC
                    #     uploader writes HTCO_Feed_Temp_ColdGrp), so feed_cold never loaded.
                    #   - 'Brine Discharge Temp' is '-' (blank) for every row of the Operational sheet,
                    #     so brine_out_t never loaded either.
                    # Both silently stayed 0, which collapsed the Overall HTC driving forces to 0.
                    'feed_cold': ['HTCO_Feed_Temp_ColdGrp', 'Feed Temp to Cold Group'],
                    'brine_out_t': ['HTCO_Brine_Disch_Temp', 'Brine Discharge Temp'],
                    'cond_temp': ['condensate temp', 'HTC1_Cond_Temp', 'HTCO_Cond_Temp'],
                    # Each HTC calc gets its OWN condensate temp, preferring the value from that calc's own
                    # HTC sheet, and only falling back to the shared Operational reading if the HTC sheet
                    # didn't supply one. This is what keeps the live calculator in agreement with the stored
                    # (graphed) HTC: the two source sheets can legitimately disagree on condensate temp
                    # (e.g. the Overall-HTC sheet used 50C from Jun 11 while Operational still read ~76C),
                    # and for HTC purposes the HTC sheet's own value is authoritative.
                    'htc1_cond_temp': ['HTC1_Cond_Temp', 'condensate temp'],
                    'htco_cond_temp': ['HTCO_Cond_Temp', 'condensate temp'],
                    'mra_t1': ['1st Effect Vapour Temp', 'HTC1_Vapor_Temp', 'HTCO_Vapor_Temp'],
                    'mra_bt1': ['1st effect brine temp', 'HTC1_Brine_Temp'],
                    'steam': ['LP Steam consumption', 'HTC1_Steam_TPH', 'HTCO_Steam_TPH'],
                    'sw_total': ['Sea Water Feed', 'HTCO_Feed_Flow'],
                    'desal': ['Desal production', 'HTC1_Product_Flow', 'HTCO_Product_Flow'],
                    'cond_flow': ['Condensate Return', 'HTC1_Cond_Flow', 'HTCO_Cond_Flow'],
                }
                
                for cat in ['Feed', 'Product']:
                    for param, d in WATER_SPECS[cat].items(): 
                        db_to_var_mapping[d['var']] = [d['db_col']]
                for param, d in BRINE_SPECS.items():
                    db_to_var_mapping[d['var']] = [d['db_col']]
                db_to_var_mapping['mid_effects_temp'] = ['HTC1_Feed_Temp_Eff4to7', 'Intermediate Effects Avg Brine Temp']
                db_to_var_mapping['htc1_feed_flow'] = ['HTC1_Feed_Flow']
                db_to_var_mapping['steam_in_t'] = ['Steam Inlet Temp']
                db_to_var_mapping['chem_anti_ppm'] = ['Anti_PPM', 'AS_PPM']
                db_to_var_mapping['chem_foam_ppm'] = ['Foam_PPM', 'AF_PPM']
                db_to_var_mapping['chem_anti_cons'] = ['Antiscalant (kg)', 'AS_KgHr']
                db_to_var_mapping['chem_foam_cons'] = ['Antifoam (kg)', 'AF_KgHr']

                loaded_vars = False
                n_loaded = 0
                for var_key, col_names in db_to_var_mapping.items():
                    for col_name in col_names:
                        if col_name in row.index and pd.notna(row[col_name]):
                            try:
                                val_str = str(row[col_name]).strip()
                                if val_str and val_str.lower() not in ['nan', 'none', 'null', 'na']:
                                    if var_key == 'remarks': 
                                        val = val_str
                                    else: 
                                        val = float(val_str.replace(',', ''))
                                    # Heat transfer areas are fixed equipment geometry. An old row that
                                    # stored 0/blank must not overwrite the real constant, or every HTC
                                    # on this date silently reads 0.
                                    if var_key in ('area_1st', 'area_overall') and (not val or val <= 0):
                                        break
                                    st.session_state.vars[var_key] = val
                                    for tk in SYNC_MAP.get(var_key, []): 
                                        st.session_state[tk] = val
                                    loaded_vars = True
                                    n_loaded += 1
                                break
                            except: 
                                pass 
                if loaded_vars: 
                    st.session_state.date_status = ('partial' if n_loaded < 12 else 'full')
                    st.session_state.date_status_n = n_loaded
                else:
                    # The date row exists but every mapped field is blank - treat it as no data at all,
                    # which is exactly the case that used to leave another date's readings on screen.
                    st.session_state.date_status = 'blank'
                    st.session_state.date_status_n = 0
                st.rerun()

        if not date_found:
            st.session_state.date_status = 'none'
            st.session_state.date_status_n = 0
            st.rerun()

    # Display MED-4 Title
    st.title("MED-4 Management Suite")

    # State plainly what the selected date actually holds. Every field was cleared before loading, so
    # anything not supplied by this date's record reads 0 rather than carrying over from another day.
    _ds = st.session_state.get('date_status', 'none')
    _dn = st.session_state.get('date_status_n', 0)
    _dstr = log_date.strftime('%d-%m-%Y')
    if _ds == 'full':
        st.success(f"Showing logged data for {_dstr} ({_dn} fields).")
    elif _ds == 'partial':
        st.warning(f"Only partial data was logged for {_dstr} ({_dn} fields). Everything not logged that day is shown as 0, not carried over from another date.")
    elif _ds == 'blank':
        st.error(f"A record exists for {_dstr} but it contains no readings. All values are shown as 0 - nothing here is measured data.")
    else:
        st.error(f"No data was logged for {_dstr}. All values are shown as 0 - nothing here is measured data.")

    tabs = st.tabs([
        "Inputs", "Performance", "Heat Transfer", "Water Quality", 
        "Chemicals", "Prediction", "Reports", 
        "Model", "Bulk Upload"
    ])

    ops_data = {
        'Steam': get_v('steam'), 
        'Desal': get_v('desal'), 
        'Gross Prod': get_v('gross'), 
        'SW_Feed_1st': get_v('sw_upper'), 
        'SW Total': get_v('sw_total'), 
        'Brine Return': get_v('brine_ret'),
        'SW In_overall': get_v('sw_in_t'), 
        'Brine Out_overall': get_v('brine_out_t'), 
        'Stm In_1st': get_v('mra_t1'), 
        'Brine_1st': get_v('mra_bt1'), 
        'Press_1st': get_v('mra_press')
    }
    
    ops_data['GOR'] = ops_data['Gross Prod'] / ops_data['Steam'] if ops_data['Steam'] > 0 else 0
    ops_data['STEC'] = (((ops_data['Steam'] * 1000) / 3600) * LATENT_HEAT_STEAM_KJ_KG) / ops_data['Desal'] if ops_data['Desal'] > 0 else 0
    ops_data['Recovery'] = (ops_data['Gross Prod'] / ops_data['SW Total']) * 100 if ops_data['SW Total'] > 0 else 0
    ops_data['Conversion'] = ops_data['Desal'] / ops_data['SW Total'] if ops_data['SW Total'] > 0 else 0
    ops_data['Economy'] = ops_data['Steam'] / ops_data['Desal'] if ops_data['Desal'] > 0 else 0

    display_effect_df = pd.merge(BASE_EFFECTS, st.session_state.shared_effect_df, on="Effect ID")
    for col in ["Base Vapor (°C)", "Live Vapor (°C)", "Base Brine (°C)", "Live Brine (°C)", "Base HTC"]:
        if col not in display_effect_df.columns:
            display_effect_df[col] = np.nan
            
    display_effect_df = display_effect_df[["Effect ID", "Base Vapor (°C)", "Live Vapor (°C)", "Base Brine (°C)", "Live Brine (°C)", "Base HTC"]]

    # ---- HEAT DUTY (steam condensation basis) --------------------------------------------------
    # Mirrors cols V/W/X of BOTH HTC sheets:
    #   ms (kg/hr) = Steam(TPH) x 1000
    #   W  (kJ/hr) = ms x latent heat
    #   Q  (W)     = W x 1000 / 3600
    ops_data['q_1st'] = (ops_data['Steam'] * 1000 * LATENT_HEAT_STEAM_KJ_KG * 1000) / 3600
    ops_data['q_overall'] = ops_data['q_1st']

    def _lmtd_scalar(dt1, dt2):
        """LMTD = (dT1 - dT2) / ln(dT1/dT2), col N of both sheets. Returns 0 when either driving
        force is missing or non-positive, so the HTC downstream honestly reports 0 rather than a
        fabricated number. Note dT2 > dT1 in this plant's data - the formula handles that fine."""
        try:
            if dt1 is None or dt2 is None or pd.isna(dt1) or pd.isna(dt2):
                return 0.0
            if dt1 <= 0 or dt2 <= 0:
                return 0.0
            if dt1 == dt2:
                return float(dt1)
            return (dt1 - dt2) / math.log(dt1 / dt2)
        except Exception:
            return 0.0

    # ---- 1st EFFECT HTC  (sheet: '1st effect-HTC') ----------------------------------------------
    # dT1 = 1st effect vapour temp - 1st effect brine temp        (col L)
    # dT2 = condensate temp - AVG BRINE TEMP OF EFFECTS 4,5,6,7   (col M)
    #       NB: the sheet labels that column "Feed Temp", but its tag row reads "Avg of effects of
    #       7,6,5,4". It is NOT a seawater temperature.
    ops_data['dt_1st'] = get_v('mra_t1') - get_v('mra_bt1')
    _cond_1st = get_v('htc1_cond_temp') or get_v('cond_temp')
    ops_data['dt2_1st'] = _cond_1st - get_v('mid_effects_temp')
    ops_data['lmtd_1st'] = _lmtd_scalar(ops_data['dt_1st'], ops_data['dt2_1st'])
    _a1 = get_v('area_1st')
    ops_data['htc_1st'] = (
        ops_data['q_1st'] / (_a1 * ops_data['lmtd_1st'])
        if ops_data['lmtd_1st'] > 0 and _a1 > 0 else 0
    )
    ops_data['fouling_1st'] = 1 / ops_data['htc_1st'] if ops_data['htc_1st'] > 0 else 0
    # Rf = 1/U_actual - 1/U_SOR_baseline   (col AC)
    ops_data['rf_1st'] = (
        (1 / ops_data['htc_1st']) - (1 / HTC_1ST_U_SOR) if ops_data['htc_1st'] > 0 else 0
    )

    # ---- OVERALL HTC  (sheet: 'Overall-HTC') ----------------------------------------------------
    # dT1 = 1st effect vapour temp - brine DISCHARGE temp   (col L)
    # dT2 = condensate temp - FEED TEMP TO COLD GROUP       (col M)
    #       NB: this sheet ALSO labels its column "Feed Temp", but here it means the cold-group feed
    #       temp (~40 C) - a different measurement from the 1st-effect sheet's "Feed Temp" (~49 C).
    ops_data['dt1_overall'] = get_v('mra_t1') - get_v('brine_out_t')
    _cond_ov = get_v('htco_cond_temp') or get_v('cond_temp')
    ops_data['dt2_overall'] = _cond_ov - get_v('feed_cold')
    ops_data['lmtd_overall'] = _lmtd_scalar(ops_data['dt1_overall'], ops_data['dt2_overall'])
    _ao = get_v('area_overall')
    ops_data['htc_overall'] = (
        ops_data['q_overall'] / (_ao * ops_data['lmtd_overall'])
        if ops_data['lmtd_overall'] > 0 and _ao > 0 else 0
    )
    ops_data['fouling_overall'] = 1 / ops_data['htc_overall'] if ops_data['htc_overall'] > 0 else 0
    ops_data['rf_overall'] = (
        (1 / ops_data['htc_overall']) - (1 / HTC_OVERALL_U_SOR) if ops_data['htc_overall'] > 0 else 0
    )

    # Simple (non-LMTD) cascade delta, shown for reference only.
    ops_data['dt_overall_simple'] = get_v('mra_t1') - get_v('brine_11')

    mra_data = {}
    coefs = st.session_state.mra_coef 
    model_type = coefs.get("model_type", "OLS")

    # Live predictor vector, built straight off MED_MRA_PARAMS so it can never fall out of step with
    # the order the model was trained in - a mismatch here would silently feed the wrong number into
    # the wrong coefficient and produce a confident, completely wrong prediction.
    live_input_arr = [get_v(v) for v in MRA_LIVE_VARS]

    # Per-plant reference values written by the last calibration, falling back to the provisional
    # figures until a calibration has actually been committed.
    live_baseline = {k: float(coefs.get(f"BASE_{k}", MRA_BASELINE[k])) for k in MRA_COEF_KEYS}

    # Flags the dashboard uses to be upfront about how trustworthy the number is.
    mra_data['calibrated'] = bool(coefs.get("calibrated", 0))
    mra_data['n_predictors'] = N_MRA_PREDICTORS

    def _ols_predict(coef_block):
        """Dot product of the coefficient block with the live inputs. Any predictor absent from the
        block contributes zero rather than raising, so a config saved by an older build still yields
        a usable number instead of crashing the whole dashboard."""
        total = float(coef_block.get("Intercept", 0.0))
        for _k, _v in zip(MRA_COEF_KEYS, live_input_arr):
            total += float(coef_block.get(_k, 0.0)) * float(_v)
        return total

    if model_type == "OLS":
        mra_data['Predicted'] = _ols_predict(coefs)
    else:
        # The trained RF/XGB model is a .pkl on the LOCAL disk, which Streamlit Cloud wipes on every
        # container restart. The calibration coefficients now persist in the Google Sheet, so
        # model_type correctly survives as "Random Forest"/"XGBoost" while the model file itself does
        # not. Previously a bare except swallowed that and silently returned 0. Now we fall back to the
        # OLS formula so a real number is still produced, and flag that the model needs retraining.
        mra_data['model_missing'] = False
        active_model = None
        try:
            active_model = joblib.load(AI_MODEL_FILE)
        except Exception:
            # Local .pkl is gone (ephemeral disk). Rebuild it from the copy kept in the sheet.
            if load_model_blob is not None:
                try:
                    active_model = load_model_blob(db_conn, LOCAL_CONFIG_FILE)
                    if active_model is not None:
                        try:
                            joblib.dump(active_model, AI_MODEL_FILE)  # re-cache locally for this session
                        except Exception:
                            pass
                except Exception:
                    active_model = None
        if active_model is not None:
            try:
                live_df = pd.DataFrame([live_input_arr], columns=MRA_COEF_KEYS)
                mra_data['Predicted'] = float(active_model.predict(live_df)[0])
            except Exception:
                # A stored model trained on the OLD 7-input set will reject this 10-column frame.
                # That is the correct outcome - better to drop to the OLS block than to coerce the
                # columns and predict from a model that was never fitted on these inputs.
                active_model = None
        if active_model is None:
            # Fall back to the CALIBRATED OLS block that is always stored alongside the AI selection,
            # and only drop to the provisional default if even that is absent.
            mra_data['model_missing'] = True
            _b = {k: coefs.get(k, MED_MRA_COEF_DEFAULT.get(k, 0.0)) for k in MED_MRA_COEF_DEFAULT}
            mra_data['Predicted'] = _ols_predict(_b)

    mra_data['Actual'] = ops_data['Gross Prod']
    mra_data['Residual'] = mra_data['Actual'] - mra_data['Predicted']

    var_data = []
    for _key, _label, _val in zip(MRA_COEF_KEYS, MRA_LABELS, live_input_arr):
        base = live_baseline[_key]
        dev = _val - base
        weight = coefs.get(_key, 0.0)
        # Feature importances are not coefficients, so a per-parameter m3/h impact is only
        # meaningful in OLS mode.
        impact = dev * weight if model_type == "OLS" else np.nan
        var_data.append([_label, base, _val, dev, weight, impact])

    mra_data['Variance_DF'] = pd.DataFrame(var_data, columns=["Parameter", "Baseline", "Live Input", "Deviation", "Regression Weight", "Impact (TPH)"])

    water_data = {'Feed': {}, 'Product': {}}
    for cat in ['Feed', 'Product']:
        for param, details in WATER_SPECS[cat].items():
            val = get_v(details['var'])
            status = "Pass" if details['lim'][0] <= val <= details['lim'][1] else "Fail"
            water_data[cat][param] = {'min': details['lim'][0], 'max': details['lim'][1], 'val': val, 'status': status, 'db_col': details['db_col']}
            
    chem_data = {
        'anti_ppm': get_v('chem_anti_ppm'), 
        'anti_cons': get_v('chem_anti_cons'), 
        'foam_ppm': get_v('chem_foam_ppm'), 
        'foam_cons': get_v('chem_foam_cons')
    }

    # --- TAB 0: INPUTS & PFD ---
    with tabs[0]:
        tab0_subtabs = st.tabs(["Data Entry", "Live PFD Monitor"])

        with tab0_subtabs[0]:
            st.subheader("Daily Data Entry")
            st.caption(
                "Five sections, one per source sheet in the plant workbook. Shared readings entered under "
                "**Operational** flow straight into the HTC sections - you only re-enter what's genuinely "
                "specific to each HTC calculation."
            )
            if mra_data['Predicted'] > 950:
                st.warning("MRA Prediction is unusually high (>950 m³/h). Check that 'Sea Water Feed' (~2100) wasn't entered into 'Sea Water Upper' (~550).")

            entry = st.tabs([
                "1 · Operational",
                "2 · HTC — 1st Effect",
                "3 · HTC — Overall",
                "4 · Feed & Brine",
            ])

            # ---------------------------------------------------------------- 1 · OPERATIONAL
            with entry[0]:
                st.caption("Source sheet: **Operational data**. Everything the plant logs daily from the DCS.")

                st.markdown("**Flows** — m³/h (steam in TPH)")
                f1, f2, f3, f4 = st.columns(4)
                with f1:
                    st.number_input("Sea Water Upper", key="in_sw_up", on_change=sync_var, args=('sw_upper', 'in_sw_up'))
                    st.number_input("Sea Water Lower", key="in_sw_low", on_change=sync_var, args=('sw_lower', 'in_sw_low'))
                with f2:
                    st.number_input("Sea Water Feed (total)", key="in_sw_tot", on_change=sync_var, args=('sw_total', 'in_sw_tot'))
                    st.number_input("Brine Water Return", key="in_brine", on_change=sync_var, args=('brine_ret', 'in_brine'))
                with f3:
                    st.number_input("Desal Production (net)", key="in_desal", on_change=sync_var, args=('desal', 'in_desal'))
                    st.number_input("Gross Production", key="in_gross", on_change=sync_var, args=('gross', 'in_gross'))
                with f4:
                    st.number_input("LP Steam Consumption (TPH)", key="in_steam", on_change=sync_var, args=('steam', 'in_steam'))
                    st.number_input("Condensate Return", key="in_cond_flow", on_change=sync_var, args=('cond_flow', 'in_cond_flow'))

                st.divider()
                st.markdown("**Temperatures** — °C")
                t1, t2, t3, t4 = st.columns(4)
                with t1:
                    st.number_input("1st Effect Vapour Temp", key="in_t1", on_change=sync_var, args=('mra_t1', 'in_t1'),
                                    help="Tag Z711TIT414. Feeds BOTH HTC calculations as the hot-side source temp.")
                    st.number_input("1st Effect Brine Temp", key="in_bt1", on_change=sync_var, args=('mra_bt1', 'in_bt1'),
                                    help="Tag Z711TIT401. Hot-side sink for the 1st Effect HTC.")
                with t2:
                    st.number_input("Condensate Temp", key="in_cond_temp", on_change=sync_var, args=('cond_temp', 'in_cond_temp'),
                                    help="Tag Z711TIT415. Cold-side source for BOTH HTC calculations.")
                    st.number_input("Brine Discharge Temp", key="in_brine_out", on_change=sync_var, args=('brine_out_t', 'in_brine_out'),
                                    help="Hot-side sink for the Overall HTC.")
                with t3:
                    st.number_input("11th Effect Brine Temp", key="in_brine_11", on_change=sync_var, args=('brine_11', 'in_brine_11'))
                    st.number_input("Steam Inlet Temp", key="in_steam_in_t", on_change=sync_var, args=('steam_in_t', 'in_steam_in_t'))
                with t4:
                    st.number_input("SW Condenser (FFC) I/L Temp", key="in_sw_in", on_change=sync_var, args=('sw_in_t', 'in_sw_in'))
                    st.number_input("SW Condenser (FFC) O/L Temp", key="in_sw_out", on_change=sync_var, args=('sw_out_t', 'in_sw_out'))

                st.divider()
                st.markdown("**Pressures, Cooling Water & Chemicals**")
                p1, p2, p3, p4 = st.columns(4)
                with p1:
                    st.number_input("1st Effect Vapour Pressure (mmHg)", key="in_press", on_change=sync_var, args=('mra_press', 'in_press'))
                    st.number_input("LP Steam Pressure (kg/cm²g)", key="in_stm_press", on_change=sync_var, args=('stm_press', 'in_stm_press'))
                with p2:
                    st.number_input("Sea Water Pressure (kg/cm²g)", key="in_sw_press", on_change=sync_var, args=('sw_press', 'in_sw_press'))
                    st.number_input("Brine Discharge Pressure (kg/cm²g)", key="in_brine_press", on_change=sync_var, args=('brine_press', 'in_brine_press'))
                with p3:
                    st.number_input("CW Supply Temp (°C)", key="in_cw_supply", on_change=sync_var, args=('cw_supply', 'in_cw_supply'))
                    st.number_input("CW Return Temp (°C)", key="in_cw_return", on_change=sync_var, args=('cw_return', 'in_cw_return'))
                    st.number_input("CW Flow (m³/h)", key="in_cw_flow", on_change=sync_var, args=('cw_flow', 'in_cw_flow'))
                with p4:
                    st.number_input("Antiscalant Residual (ppm)", key="in_anti_ppm", on_change=sync_var, args=('chem_anti_ppm', 'in_anti_ppm'))
                    st.number_input("Antiscalant Consumption (kg/hr)", key="in_anti_cons", on_change=sync_var, args=('chem_anti_cons', 'in_anti_cons'))
                    st.number_input("Antifoam Residual (ppm)", key="in_foam_ppm", on_change=sync_var, args=('chem_foam_ppm', 'in_foam_ppm'))
                    st.number_input("Antifoam Consumption (kg/hr)", key="in_foam_cons", on_change=sync_var, args=('chem_foam_cons', 'in_foam_cons'))

                st.divider()
                st.number_input("Condensate Conductivity (µS/cm)", key="in_cond_cond", on_change=sync_var, args=('cond_cond', 'in_cond_cond'))
                st.text_area("Remarks", key="t0_remarks", on_change=sync_var, args=('remarks', 't0_remarks'), height=68,
                             help="Mirrors the Remarks box on the Reporting tab - edit either one.")

                # Effect-wise temperature cascade is no longer recorded by the team, so the input UI is
                # hidden. skip_eff is forced True so any downstream logic treats it as intentionally absent.
                if not get_v('skip_eff'):
                    st.session_state.vars['skip_eff'] = True

            # ---------------------------------------------------------------- 2 · HTC 1st EFFECT
            with entry[1]:
                st.caption("Source sheet: **1st effect-HTC**. Heat transfer across the 1st effect tube bundle only.")
                st.success(
                    "**Already taken from Operational** — steam rate, 1st effect vapour temp, 1st effect brine temp, "
                    "condensate temp. Only the two genuinely 1st-effect-specific readings are below."
                )
                h1a, h1b = st.columns(2)
                with h1a:
                    st.number_input(
                        "Avg Brine Temp of Effects 4-5-6-7 (°C)", key="in_mid_effects_temp",
                        on_change=sync_var, args=('mid_effects_temp', 'in_mid_effects_temp'),
                        help="On the source sheet this column is labelled 'Feed Temp', but the tag row reads "
                             "'Avg of effects of 7,6,5,4'. It is the COLD-SIDE reference (ΔT2) for this calculation "
                             "— NOT a seawater temperature. Typically ~49 °C."
                    )
                    st.number_input(
                        "Feed Flow to 1st Effect (m³/h)", key="in_htc1_feed_flow",
                        on_change=sync_var, args=('htc1_feed_flow', 'in_htc1_feed_flow'),
                        help="Tag Z711FIT424 as recorded on the 1st-effect sheet (~514 m³/h). This is NOT the total "
                             "seawater feed (~2062) used on the Overall sheet."
                    )
                with h1b:
                    st.number_input("1st Effect Heat Transfer Area (m²)", key="in_area_1st",
                                    on_change=sync_var, args=('area_1st', 'in_area_1st'),
                                    help="π × 5.5 m × 31,244 tubes × 0.024 m OD = 12,950 m²")

                st.divider()
                d1, d2, d3, d4 = st.columns(4)
                d1.metric("ΔT1 (vapour − brine)", f"{ops_data['dt_1st']:.2f} °C")
                d2.metric("ΔT2 (condensate − eff 4-7)", f"{ops_data.get('dt2_1st', 0):.2f} °C")
                d3.metric("LMTD", f"{ops_data.get('lmtd_1st', 0):.2f} °C")
                d4.metric("1st Effect HTC", f"{ops_data['htc_1st']:.1f} W/m²K")

            # ---------------------------------------------------------------- 3 · HTC OVERALL
            with entry[2]:
                st.caption("Source sheet: **Overall-HTC**. Heat transfer across all 11 effects combined.")
                st.success(
                    "**Already taken from Operational** — steam rate, 1st effect vapour temp, brine discharge temp, "
                    "condensate temp, total seawater feed. Only the two Overall-specific readings are below."
                )
                hoa, hob = st.columns(2)
                with hoa:
                    st.number_input(
                        "Feed Temp to Cold Group (°C)", key="in_feed_cold",
                        on_change=sync_var, args=('feed_cold', 'in_feed_cold'),
                        help="On the source sheet this column is also labelled 'Feed Temp' — but here it means the "
                             "feed temperature into the cold group (~40 °C), a DIFFERENT measurement from the "
                             "'Feed Temp' on the 1st-effect sheet. It is the cold-side reference (ΔT2) here."
                    )
                with hob:
                    st.number_input("Overall Heat Transfer Area (m²)", key="in_area_overall",
                                    on_change=sync_var, args=('area_overall', 'in_area_overall'),
                                    help="11 effects × 12,950 m² × 1.15 correction = 163,818 m²")

                st.divider()
                o1, o2, o3, o4 = st.columns(4)
                o1.metric("ΔT1 (vapour − brine disch.)", f"{ops_data.get('dt1_overall', 0):.2f} °C")
                o2.metric("ΔT2 (condensate − cold grp)", f"{ops_data.get('dt2_overall', 0):.2f} °C")
                o3.metric("LMTD", f"{ops_data.get('lmtd_overall', 0):.2f} °C")
                o4.metric("Overall HTC", f"{ops_data['htc_overall']:.2f} W/m²K")

            # ---------------------------------------------------------------- 4 · FEED & BRINE
            with entry[3]:
                st.caption("Source sheet: **Feed & Brine Water Analysis**. Daily lab results.")
                st.checkbox("Skip water analysis today", key="in_skip_wq", on_change=sync_var, args=('skip_wq', 'in_skip_wq'))
                if not get_v('skip_wq'):
                    wf, wb_ = st.columns(2)
                    with wf:
                        st.markdown("**Feed Water (Sea Water)**")
                        for p, dd in WATER_SPECS["Feed"].items():
                            lo, hi = dd['lim']
                            st.number_input(f"{p}", key=f"in_{dd['var']}", on_change=sync_var,
                                            args=(dd['var'], f"in_{dd['var']}"), help=f"Specified limit: {lo} – {hi}")
                    with wb_:
                        st.markdown("**Brine Water**")
                        st.caption("No specified limits on the source sheet — tracked for trending.")
                        for p, dd in BRINE_SPECS.items():
                            st.number_input(f"{p}", key=f"in_{dd['var']}", on_change=sync_var,
                                            args=(dd['var'], f"in_{dd['var']}"))

            # NOTE: The "5 · Desal Product" manual-entry section was removed at Reliance's request.
            # Desal/product water quality is NO LONGER captured or shown on screen anywhere in the app,
            # but it is still ingested through the "Desal (Product) Analysis" bulk-upload tab and stored
            # against the Product_* columns, so the history stays queryable in the backend and the
            # figures continue to appear in the generated Word reports.

        with tab0_subtabs[1]:
            st.markdown("### Process Flow Diagram - Live Tags")
            if PIL_INSTALLED and (os.path.exists("Desal PFD (1).TIF") or os.path.exists("Desal PFD (1).tiff") or os.path.exists("Desal PFD.TIF")):
                try:
                    from PIL import Image
                    file_name = "Desal PFD (1).TIF" if os.path.exists("Desal PFD (1).TIF") else ("Desal PFD (1).tiff" if os.path.exists("Desal PFD (1).tiff") else "Desal PFD.TIF")
                    img = Image.open(file_name).convert("RGB")
                    buffered = BytesIO()
                    img.save(buffered, format="PNG")
                    img_str = base64.b64encode(buffered.getvalue()).decode()
                    
                    html_view = f"""
                    <div style="position: relative; width: 100%; max-width: 1200px; margin: auto; background: #fff; border: 2px solid #ddd; border-radius: 8px; overflow: hidden;">
                        <img src="data:image/png;base64,{img_str}" style="width: 100%; display: block;" alt="MED PFD"/>
                        
                        <div style="position: absolute; top: 5%; left: 2%; background: rgba(0,20,50,0.85); color: #00ff00; padding: 6px 12px; font-family: monospace; border: 1px solid #00ff00; border-radius: 4px; box-shadow: 0 0 8px #00ff00; font-size: 13px;">
                            <strong>SEA WATER SYSTEM</strong><br>
                            Sea Water Feed: {ops_data['SW Total']} m³/h<br>
                            Sea Water Upper: {ops_data['SW_Feed_1st']} m³/h<br>
                            Sea Water Lower: {get_v('sw_lower')} m³/h<br>
                            Sea Water cond I/L temp: {ops_data['SW In_overall']} °C<br>
                            Sea Water Condenser O/L Temp: {get_v('sw_out_t')} °C<br>
                            CW supply: {get_v('cw_supply')}
                        </div>
                        
                        <div style="position: absolute; top: 5%; right: 2%; background: rgba(50,0,0,0.85); color: #ff3333; padding: 6px 12px; font-family: monospace; border: 1px solid #ff3333; border-radius: 4px; box-shadow: 0 0 8px #ff3333; font-size: 13px;">
                            <strong>STEAM & 1ST EFFECT</strong><br>
                            LP Steam consumption: {ops_data['Steam']} TPH<br>
                            1st Effect Vapour Temp: {ops_data['Stm In_1st']} °C<br>
                            1st effect vapour pressure: {ops_data['Press_1st']} mmHg<br>
                            1st effect brine temp: {ops_data['Brine_1st']} °C<br>
                            Delta T: {ops_data['dt_1st']:.2f} °C
                        </div>

                        <div style="position: absolute; bottom: 5%; left: 2%; background: rgba(0,50,50,0.85); color: #00ffff; padding: 6px 12px; font-family: monospace; border: 1px solid #00ffff; border-radius: 4px; box-shadow: 0 0 8px #00ffff; font-size: 13px;">
                            <strong>PRODUCTION</strong><br>
                            Gross production: {ops_data['Gross Prod']} m³/h<br>
                            Desal production: {ops_data['Desal']} m³/h<br>
                            Condensate Return: {get_v('cond_flow')}<br>
                            condensate temp: {get_v('cond_temp')} °C
                        </div>
                        
                        <div style="position: absolute; bottom: 5%; right: 2%; background: rgba(50,25,0,0.85); color: #ff9900; padding: 6px 12px; font-family: monospace; border: 1px solid #ff9900; border-radius: 4px; box-shadow: 0 0 8px #ff9900; font-size: 13px;">
                            <strong>BRINE SYSTEM</strong><br>
                            Brine Water Return: {ops_data['Brine Return']} m³/h<br>
                            Brine Discharge Temp: {ops_data['Brine Out_overall']} °C<br>
                            CW Return: {get_v('cw_return')}
                        </div>
                    </div>
                    """
                    st.components.v1.html(html_view, height=800)
                except Exception as e:
                    st.error(f"Could not render TIF overlay. Error: {e}")
            else:
                st.info("Digital Twin HUD: Please upload 'Desal PFD (1).TIF' into the application directory to unlock the live interactive diagram overlay.")

    # --- TAB 1: FLOW KPIs & SOR MATRIX ---
    with tabs[1]:
        st.subheader("Performance Overview")

        anti_gm_m3 = (get_v('chem_anti_cons') / ops_data['SW Total']) * 1000 if ops_data['SW Total'] > 0 else 0
        foam_gm_m3 = (get_v('chem_foam_cons') / ops_data['SW Total']) * 1000 if ops_data['SW Total'] > 0 else 0
        has_anti_kg = get_v('chem_anti_cons') > 0
        has_foam_kg = get_v('chem_foam_cons') > 0

        # --- Headline KPI cards: the numbers Reliance/Chembond actually track day to day, up front ---
        st.markdown("##### Headline Performance Indicators")
        kpi1, kpi2, kpi3, kpi4, kpi5 = st.columns(5)
        kpi1.metric("GOR", f"{ops_data['GOR']:.2f}", f"{ops_data['GOR'] - 11.4:+.2f} vs SOR", help="Gain Output Ratio: Gross production / Steam consumption. SOR baseline: 11.4")
        kpi2.metric("STEC", f"{ops_data['STEC']:.1f} kWh/t", help="Specific Thermal Energy Consumption per tonne of distillate")
        kpi3.metric("Overall HTC", f"{ops_data['htc_overall']:.1f} W/m²K", help="Whole-plant heat transfer coefficient (steam condensation basis)")
        kpi4.metric("1st Effect HTC", f"{ops_data['htc_1st']:.1f} W/m²K", help="1st effect heat transfer coefficient (steam condensation basis)")
        kpi5.metric("Recovery", f"{ops_data['Recovery']:.1f}%", help="Gross production / Total sea water feed")
        st.divider()

        def color_diff(val):
            try:
                v = float(val)
                color = 'green' if v >= 0 else 'red'
                return f'color: {color}; font-weight: bold'
            except:
                return ''

        st.markdown("#### Sea Water")
        df_a = pd.DataFrame([
            {"Parameter": "Temp.", "UOM": "°C", "Design": "19-35", "SOR Base": 29.0, "Actual": get_v('sw_in_t'), "Difference": get_v('sw_in_t') - 29.0},
            {"Parameter": "Pressure", "UOM": "kg/cm2-g", "Design": "2.5", "SOR Base": 1.7, "Actual": get_v('sw_press'), "Difference": get_v('sw_press') - 1.7},
            {"Parameter": "Total sea water flow to desal unit", "UOM": "m3/hr", "Design": "2400", "SOR Base": 2112.0, "Actual": ops_data['SW Total'], "Difference": ops_data['SW Total'] - 2112.0}
        ])
        st.dataframe(df_a.style.map(color_diff, subset=['Difference']).format({"SOR Base": "{:.1f}", "Actual": "{:.1f}", "Difference": "{:+.1f}"}), use_container_width=True, hide_index=True)

        st.markdown("#### LP Steam")
        df_b = pd.DataFrame([
            {"Parameter": "Total Flow (Thermocompressor + NCG)", "UOM": "Tonne/hr", "Design": "97.5", "SOR Base": 76.94, "Actual": ops_data['Steam'], "Difference": ops_data['Steam'] - 76.94},
            {"Parameter": "Pressure", "UOM": "kg/cm2-g", "Design": "3.5", "SOR Base": 4.3, "Actual": get_v('stm_press'), "Difference": get_v('stm_press') - 4.3},
            {"Parameter": "Steam Inlet Temp.", "UOM": "°C", "Design": "176", "SOR Base": 176.0, "Actual": get_v('steam_in_t'), "Difference": get_v('steam_in_t') - 176.0},
            {"Parameter": "First Effect Vapour Temp.", "UOM": "°C", "Design": "69", "SOR Base": 68.47, "Actual": get_v('mra_t1'), "Difference": get_v('mra_t1') - 68.47}
        ])
        st.dataframe(df_b.style.map(color_diff, subset=['Difference']).format({"SOR Base": "{:.2f}", "Actual": "{:.2f}", "Difference": "{:+.2f}"}), use_container_width=True, hide_index=True)

        st.markdown("#### Cooling Water")
        df_c = pd.DataFrame([
            {"Parameter": "Flow", "UOM": "m3/hr", "Design": "4200", "SOR Base": 2726.0, "Actual": get_v('cw_flow'), "Difference": get_v('cw_flow') - 2726.0},
            {"Parameter": "Cooling Water Supply Temp", "UOM": "°C", "Design": "32", "SOR Base": 31.9, "Actual": get_v('cw_supply'), "Difference": get_v('cw_supply') - 31.9},
            {"Parameter": "Cooling Water Return Temp", "UOM": "°C", "Design": "41", "SOR Base": 37.5, "Actual": get_v('cw_return'), "Difference": get_v('cw_return') - 37.5}
        ])
        st.dataframe(df_c.style.map(color_diff, subset=['Difference']).format({"SOR Base": "{:.1f}", "Actual": "{:.1f}", "Difference": "{:+.1f}"}), use_container_width=True, hide_index=True)

        st.markdown("#### Desalinated Water")
        df_d = pd.DataFrame([
            {"Parameter": "Desal water production", "UOM": "m3/hr", "Design": "1000", "SOR Base": 824.0, "Actual": ops_data['Desal'], "Difference": ops_data['Desal'] - 824.0},
            {"Parameter": "Conductivity", "UOM": "microS/cm", "Design": "<15", "SOR Base": 2.5, "Actual": get_v('p_cond'), "Difference": get_v('p_cond') - 2.5}
        ])
        st.dataframe(df_d.style.map(color_diff, subset=['Difference']).format({"SOR Base": "{:.1f}", "Actual": "{:.1f}", "Difference": "{:+.1f}"}), use_container_width=True, hide_index=True)

        st.markdown("#### Brine Discharge")
        df_e = pd.DataFrame([
            {"Parameter": "Flow", "UOM": "m3/hr", "Design": "1400", "SOR Base": 1315.0, "Actual": ops_data['Brine Return'], "Difference": ops_data['Brine Return'] - 1315.0},
            {"Parameter": "Temp.", "UOM": "°C", "Design": "43.5", "SOR Base": 40.5, "Actual": ops_data['Brine Out_overall'], "Difference": ops_data['Brine Out_overall'] - 40.5},
            {"Parameter": "Pressure", "UOM": "kg/cm2-g", "Design": "6", "SOR Base": 1.3, "Actual": get_v('brine_press'), "Difference": get_v('brine_press') - 1.3}
        ])
        st.dataframe(df_e.style.map(color_diff, subset=['Difference']).format({"SOR Base": "{:.1f}", "Actual": "{:.1f}", "Difference": "{:+.1f}"}), use_container_width=True, hide_index=True)

        st.markdown("#### Condensate Return")
        df_f = pd.DataFrame([
            {"Parameter": "Quantity", "UOM": "m3/hr", "Design": "100", "SOR Base": 127.0, "Actual": get_v('cond_flow'), "Difference": get_v('cond_flow') - 127.0},
            {"Parameter": "Temp.", "UOM": "°C", "Design": "70", "SOR Base": 71.0, "Actual": get_v('cond_temp'), "Difference": get_v('cond_temp') - 71.0},
            {"Parameter": "Conductivity", "UOM": "microS/cm", "Design": "<15", "SOR Base": 3.0, "Actual": get_v('cond_cond'), "Difference": get_v('cond_cond') - 3.0}
        ])
        st.dataframe(df_f.style.map(color_diff, subset=['Difference']).format({"SOR Base": "{:.1f}", "Actual": "{:.1f}", "Difference": "{:+.1f}"}), use_container_width=True, hide_index=True)

        st.markdown("#### Plant Capacity Details")
        df_h = pd.DataFrame([
            {"Parameter": "Gross desal water production", "UOM": "tph", "Design": "1000", "SOR Base": 873.0, "Actual": ops_data['Gross Prod'], "Difference": ops_data['Gross Prod'] - 873.0},
            {"Parameter": "Conversion (Product to Feed)", "UOM": "%", "Design": "41.6", "SOR Base": 41.4, "Actual": ops_data['Conversion'] * 100, "Difference": (ops_data['Conversion'] * 100) - 41.4},
            {"Parameter": "GOR / Steam Economy", "UOM": "-", "Design": "10.5", "SOR Base": 11.4, "Actual": ops_data['GOR'], "Difference": ops_data['GOR'] - 11.4},
            {"Parameter": "Steam Economy (Steam/Desal)", "UOM": "Norms", "Design": "0.088", "SOR Base": 0.088, "Actual": ops_data['Economy'], "Difference": ops_data['Economy'] - 0.088},
            {"Parameter": "1st effect vapour temp.", "UOM": "°C", "Design": "74", "SOR Base": 72.0, "Actual": get_v('mra_t1'), "Difference": get_v('mra_t1') - 72.0},
            {"Parameter": "1st effect pressure", "UOM": "mm Hg", "Design": "248", "SOR Base": 256.0, "Actual": get_v('mra_press'), "Difference": get_v('mra_press') - 256.0},
            {"Parameter": "1st effect brine temp.", "UOM": "°C", "Design": "69", "SOR Base": 69.0, "Actual": get_v('mra_bt1'), "Difference": get_v('mra_bt1') - 69.0},
            {"Parameter": "11th effect brine temp", "UOM": "°C", "Design": "44", "SOR Base": 42.0, "Actual": get_v('brine_11'), "Difference": get_v('brine_11') - 42.0},
            {"Parameter": "Delta T (1st effect vapour temp -1st effect brine temp)", "UOM": "°C", "Design": "4", "SOR Base": 2.5, "Actual": ops_data['dt_1st'], "Difference": ops_data['dt_1st'] - 2.5},
            {"Parameter": "Overall delta T(1st eff brine temp - 11th eff brine temp)", "UOM": "°C", "Design": "25", "SOR Base": 27.1, "Actual": ops_data['dt_overall_simple'], "Difference": ops_data['dt_overall_simple'] - 27.1},
            {"Parameter": "Feed temp to cold group", "UOM": "°C", "Design": "40", "SOR Base": 37.0, "Actual": get_v('feed_cold'), "Difference": get_v('feed_cold') - 37.0}
        ])
        st.dataframe(df_h.style.map(color_diff, subset=['Difference']).format({"SOR Base": "{:.3f}", "Actual": "{:.3f}", "Difference": "{:+.3f}"}), use_container_width=True, hide_index=True)

        st.markdown("#### Chemical Dosing & Residual")
        st.caption("Dosing rate (gm/m³) needs a logged kg-consumption figure for the day; residual PPM comes from lab analysis and is tracked independently.")
        df_i = pd.DataFrame([
            {"Parameter": "Antiscalant (ID204)/IN-204AS", "UOM": "gm/m3 sea water", "Design": "7", "SOR Base": 10.5,
             "Actual": anti_gm_m3 if has_anti_kg else np.nan, "Difference": (anti_gm_m3 - 10.5) if has_anti_kg else np.nan,
             "Residual (PPM)": get_v('chem_anti_ppm')},
            {"Parameter": "Antifoam", "UOM": "gm/m3 sea water", "Design": "0.25", "SOR Base": 0.16,
             "Actual": foam_gm_m3 if has_foam_kg else np.nan, "Difference": (foam_gm_m3 - 0.16) if has_foam_kg else np.nan,
             "Residual (PPM)": get_v('chem_foam_ppm')}
        ])
        st.dataframe(
            df_i.style.map(color_diff, subset=['Difference']).format({"SOR Base": "{:.2f}", "Actual": "{:.2f}", "Difference": "{:+.2f}", "Residual (PPM)": "{:.2f}"}, na_rep="No kg data logged"),
            use_container_width=True, hide_index=True
        )
        if not has_anti_kg or not has_foam_kg:
            missing_chem = ([] if has_anti_kg else ["antiscalant"]) + ([] if has_foam_kg else ["antifoam"])
            st.info(f"No {' or '.join(missing_chem)} consumption (kg) is logged for this date, so the gm/m³ dosing rate can't be calculated - only the PPM residual is shown. Log daily kg consumption on the Chemical Dosing tab to enable this.")
        
        sor_export_dfs = {
            "Sea Water": df_a, "LP Steam": df_b, "Cooling Water": df_c, 
            "Desalinated Water": df_d, "Brine Discharge": df_e, 
            "Condensate Return": df_f, "Plant Capacity Details": df_h, 
            "Chemical Consumption": df_i
        }

    # --- TAB 2: OVERALL HTC ---
    with tabs[2]:
        st.subheader("Heat Transfer & Fouling")
        st.caption(
            "Both calculations use the steam-condensation basis: **U = Q / (A × LMTD)**, with "
            "**LMTD = (ΔT1 − ΔT2) / ln(ΔT1/ΔT2)**. They differ in which temperatures define ΔT1 and ΔT2, "
            "and in heat transfer area — exactly as in the two source sheets."
        )

        htc_headline = st.columns(4)
        _d1 = ops_data['htc_1st'] - HTC_1ST_U_SOR
        _d2 = ops_data['htc_overall'] - HTC_OVERALL_U_SOR
        htc_headline[0].metric("1st Effect HTC", f"{ops_data['htc_1st']:.1f} W/m²K",
                               f"{_d1:+.1f} vs SOR ({HTC_1ST_U_SOR:.0f})")
        htc_headline[1].metric("Overall HTC", f"{ops_data['htc_overall']:.2f} W/m²K",
                               f"{_d2:+.2f} vs SOR ({HTC_OVERALL_U_SOR:.1f})")
        htc_headline[2].metric("1st Effect Fouling Rf", f"{ops_data['rf_1st']:.6f}",
                               help="Rf = 1/U_actual − 1/U_SOR. Rising = fouling building up.")
        htc_headline[3].metric("Overall Fouling Rf", f"{ops_data['rf_overall']:.5f}",
                               help="Rf = 1/U_actual − 1/U_SOR. Rising = fouling building up.")

        if ops_data['htc_1st'] == 0 or ops_data['htc_overall'] == 0:
            st.warning(
                "An HTC reads 0, which means one of its required temperatures is missing or non-physical "
                "(ΔT ≤ 0). Check the Inputs tab — the calculator reports 0 rather than inventing a value."
            )

        st.divider()
        c1, c2 = st.columns(2)

        with c1:
            st.markdown("#### 1st Effect")
            st.caption("Source: `1st effect-HTC` · Area 12,950 m² (single tube bundle)")
            st.number_input("1st Effect Heat Transfer Area (m²)", key="t2_area_1st",
                            on_change=sync_var, args=('area_1st', 't2_area_1st'))
            st.number_input("Avg Brine Temp, Effects 4-5-6-7 (°C)", key="t2_mid_effects_temp",
                            on_change=sync_var, args=('mid_effects_temp', 't2_mid_effects_temp'),
                            help="Cold-side reference. The source sheet calls this 'Feed Temp' — it is not a seawater temp.")
            st.dataframe(pd.DataFrame([
                {"Step": "ΔT1  =  vapour − 1st eff. brine", "Value": ops_data['dt_1st'], "Unit": "°C"},
                {"Step": "ΔT2  =  condensate − eff 4-7 avg", "Value": ops_data['dt2_1st'], "Unit": "°C"},
                {"Step": "LMTD", "Value": ops_data['lmtd_1st'], "Unit": "°C"},
                {"Step": "Q  (heat duty)", "Value": ops_data['q_1st'] / 1000, "Unit": "kW"},
                {"Step": "A  (area)", "Value": get_v('area_1st'), "Unit": "m²"},
                {"Step": "U = Q / (A × LMTD)", "Value": ops_data['htc_1st'], "Unit": "W/m²K"},
            ]).style.format({"Value": "{:,.2f}"}), use_container_width=True, hide_index=True)

        with c2:
            st.markdown("#### Overall Plant")
            st.caption("Source: `Overall-HTC` · Area 163,818 m² (11 × 12,950 × 1.15)")
            st.number_input("Overall Heat Transfer Area (m²)", key="t2_area_overall",
                            on_change=sync_var, args=('area_overall', 't2_area_overall'))
            st.number_input("Feed Temp to Cold Group (°C)", key="t2_feed_cold",
                            on_change=sync_var, args=('feed_cold', 't2_feed_cold'),
                            help="Cold-side reference. The source sheet also calls this 'Feed Temp', but it is a "
                                 "different measurement from the 1st-effect sheet's column of the same name.")
            st.dataframe(pd.DataFrame([
                {"Step": "ΔT1  =  vapour − brine discharge", "Value": ops_data['dt1_overall'], "Unit": "°C"},
                {"Step": "ΔT2  =  condensate − cold grp feed", "Value": ops_data['dt2_overall'], "Unit": "°C"},
                {"Step": "LMTD", "Value": ops_data['lmtd_overall'], "Unit": "°C"},
                {"Step": "Q  (heat duty)", "Value": ops_data['q_overall'] / 1000, "Unit": "kW"},
                {"Step": "A  (area)", "Value": get_v('area_overall'), "Unit": "m²"},
                {"Step": "U = Q / (A × LMTD)", "Value": ops_data['htc_overall'], "Unit": "W/m²K"},
            ]).style.format({"Value": "{:,.2f}"}), use_container_width=True, hide_index=True)

        st.divider()
        st.markdown("#### Fouling Trend")
        _logs = st.session_state.daily_logs
        if _logs is not None and not _logs.empty and 'Date' in _logs.columns:
            tdf = _logs.copy()
            tdf['Date'] = standardize_dates(tdf['Date'])
            for c in ['1st Effect HTC', 'Overall HTC']:
                tdf[c] = pd.to_numeric(tdf.get(c), errors='coerce')
            tdf = tdf.dropna(subset=['Date']).sort_values('Date')
            tdf = tdf[(tdf['1st Effect HTC'].fillna(0) > 0) | (tdf['Overall HTC'].fillna(0) > 0)]
            if len(tdf) > 1:
                g1, g2 = st.columns(2)
                for col, metric, base, colr in (
                    (g1, '1st Effect HTC', HTC_1ST_U_SOR, '#1f77b4'),
                    (g2, 'Overall HTC', HTC_OVERALL_U_SOR, '#2ca02c'),
                ):
                    sub = tdf[tdf[metric] > 0]
                    if len(sub) > 1:
                        ch = alt.Chart(sub).mark_line(point=True, color=colr).encode(
                            x=alt.X('Date:T', title=None),
                            y=alt.Y(f'{metric}:Q', scale=alt.Scale(zero=False), title='W/m²K'),
                            tooltip=['Date:T', f'{metric}:Q'])
                        rule = alt.Chart(pd.DataFrame({'y': [base]})).mark_rule(
                            color='red', strokeDash=[4, 4]).encode(y='y:Q')
                        trend = ch.transform_regression('Date', metric).mark_line(
                            color='black', strokeDash=[5, 5])
                        col.markdown(f"**{metric}** (red = SOR baseline)")
                        col.altair_chart(ch + rule + trend, use_container_width=True)
                    else:
                        col.info(f"Not enough {metric} history yet.")
            else:
                st.info("No HTC history in the registry yet. Upload HTC data on the Bulk Uploads tab to build a trend.")
        else:
            st.info("No HTC history in the registry yet.")

    # --- TAB 3: WATER ANALYSIS TAB ---
    with tabs[3]:
        st.subheader("Water Quality Analysis")
        if not get_v('skip_wq'):
            w_col1, w_col2 = st.columns(2)
            with w_col1:
                st.markdown("**Intake Seawater Matrix**")
                for param, d in WATER_SPECS["Feed"].items():
                    c_in, c_chk = st.columns([2, 2])
                    with c_in: 
                        st.number_input(f"{param} ({d['lim'][0]}-{d['lim'][1]})", key=f"t3_{d['var']}", on_change=sync_var, args=(d['var'], f"t3_{d['var']}"))
                    c_chk.markdown(f"<div style='margin-top:30px'>{water_data['Feed'][param]['status']}</div>", unsafe_allow_html=True)
            with w_col2:
                # Brine replaces the old Product Distillate matrix here. Brine is already captured on the
                # Feed & Brine entry tab and via bulk upload, but until now had no display anywhere.
                # The source sheet lists no specified limits for brine, so these are trended against the
                # historical average rather than graded pass/fail.
                st.markdown("**Brine Matrix**")
                st.caption("No specified limits on the source sheet — tracked for trending against the reference average.")
                for param, d in BRINE_SPECS.items():
                    c_in, c_ref = st.columns([2, 2])
                    with c_in:
                        st.number_input(f"{param}", key=f"t3_{d['var']}", on_change=sync_var, args=(d['var'], f"t3_{d['var']}"))
                    _bval = get_v(d['var'])
                    _bavg = d.get('avg', 0.0)
                    if _bavg and _bval:
                        _bdev = (_bval - _bavg) / _bavg * 100
                        _bcol = "#D9822B" if abs(_bdev) > 15 else "#5A6B7B"
                        _btxt = f"<span style='color:{_bcol}'>{_bdev:+.1f}% vs ref {_bavg:g}</span>"
                    else:
                        _btxt = "<span style='color:#9AA5B1'>no reading</span>"
                    c_ref.markdown(f"<div style='margin-top:30px'>{_btxt}</div>", unsafe_allow_html=True)

    # --- TAB 4: CHEMICAL DOSING ---
    with tabs[4]:
        st.subheader("Chemical Treatment Monitoring")
        st.caption(
            "Kem Watreat r 3687 (antiscalant) and Kem Antifoam 1795. Dosing is derived from the daily tank-level "
            "drop uploaded on the Chemical Doses bulk tab; the figures below reflect the selected date's record."
        )

        # Pull this date's chemical record straight from the registry.
        chem_row = {}
        _logs = st.session_state.daily_logs
        if _logs is not None and not _logs.empty and 'Date' in _logs.columns:
            _lref = _logs.copy()
            _lref['Date'] = standardize_dates(_lref['Date']).dt.strftime('%Y-%m-%d')
            _match = _lref[_lref['Date'] == log_date_str]
            if not _match.empty:
                chem_row = _match.iloc[-1].to_dict()

        def _cv(col, fallback=0.0):
            try:
                v = float(chem_row.get(col))
                return v if pd.notna(v) else fallback
            except (TypeError, ValueError):
                return fallback

        as_kghr, as_ppm, as_lph = _cv('AS_KgHr'), _cv('AS_PPM'), _cv('AS_LPH')
        af_kghr, af_ppm, af_lph = _cv('AF_KgHr'), _cv('AF_PPM'), _cv('AF_LPH')
        # Fall back to the manually-tracked residual PPM fields if the dosing record isn't present.
        if as_ppm == 0: as_ppm = get_v('chem_anti_ppm')
        if af_ppm == 0: af_ppm = get_v('chem_foam_ppm')

        st.markdown("##### Dosing Snapshot")
        m1, m2, m3, m4 = st.columns(4)
        m1.metric("Antiscalant Dose", f"{as_kghr:.2f} kg/hr", f"{as_ppm:.2f} ppm")
        m2.metric("Antiscalant LPH", f"{as_lph:.2f} L/hr")
        m3.metric("Antifoam Dose", f"{af_kghr:.3f} kg/hr", f"{af_ppm:.3f} ppm")
        m4.metric("Antifoam LPH", f"{af_lph:.2f} L/hr")

        if not chem_row or (as_kghr == 0 and af_kghr == 0):
            st.info("No derived dosing record found for this date. Upload the Chemical Doses sheet on the Bulk Uploads tab to populate this.")

        st.divider()
        cc1, cc2 = st.columns(2)
        with cc1:
            st.markdown("### Kem Watreat r 3687 — Antiscalant")
            as_target = 10.5  # SOR baseline ppm
            df_as = pd.DataFrame([
                {"Metric": "Dose Rate (kg/hr)", "Value": as_kghr},
                {"Metric": "Dose Rate (LPH)", "Value": as_lph},
                {"Metric": "Residual (ppm)", "Value": as_ppm},
                {"Metric": "SOR Target (ppm)", "Value": as_target},
                {"Metric": "Deviation vs Target (ppm)", "Value": as_ppm - as_target},
            ])
            st.dataframe(df_as.style.format({"Value": "{:.3f}"}), use_container_width=True, hide_index=True)
            st.markdown("**MMC Stock (kg)**")
            st.dataframe(pd.DataFrame([{
                "Opening": _cv('AS_Stock_Open'), "Received": _cv('AS_Stock_Recd'),
                "Consumed": _cv('AS_Stock_Consumed'), "Closing": _cv('AS_Stock_Close')
            }]).style.format("{:.1f}"), use_container_width=True, hide_index=True)

        with cc2:
            st.markdown("### Kem Antifoam 1795 — Antifoam")
            af_target = 0.16
            df_af = pd.DataFrame([
                {"Metric": "Dose Rate (kg/hr)", "Value": af_kghr},
                {"Metric": "Dose Rate (LPH)", "Value": af_lph},
                {"Metric": "Residual (ppm)", "Value": af_ppm},
                {"Metric": "SOR Target (ppm)", "Value": af_target},
                {"Metric": "Deviation vs Target (ppm)", "Value": af_ppm - af_target},
            ])
            st.dataframe(df_af.style.format({"Value": "{:.3f}"}), use_container_width=True, hide_index=True)
            st.markdown("**MMC Stock (kg)**")
            st.dataframe(pd.DataFrame([{
                "Opening": _cv('AF_Stock_Open'), "Received": _cv('AF_Stock_Recd'),
                "Consumed": _cv('AF_Stock_Consumed'), "Closing": _cv('AF_Stock_Close')
            }]).style.format("{:.1f}"), use_container_width=True, hide_index=True)

        st.divider()
        st.markdown("#### Dosing Trend")
        if _logs is not None and not _logs.empty and 'Date' in _logs.columns:
            tdf = _logs.copy()
            tdf['Date'] = standardize_dates(tdf['Date'])
            for c in ['AS_KgHr', 'AF_KgHr', 'AS_PPM', 'AF_PPM']:
                tdf[c] = pd.to_numeric(tdf.get(c), errors='coerce')
            tdf = tdf.dropna(subset=['Date']).sort_values('Date')
            tdf = tdf[(tdf['AS_KgHr'].fillna(0) > 0) | (tdf['AF_KgHr'].fillna(0) > 0)]
            if len(tdf) > 1:
                g1, g2 = st.columns(2)
                for col, metric, title, colr in (
                    (g1, 'AS_KgHr', 'Antiscalant Dose (kg/hr)', '#1f77b4'),
                    (g2, 'AF_KgHr', 'Antifoam Dose (kg/hr)', '#ff7f0e'),
                ):
                    sub = tdf[tdf[metric] > 0]
                    if len(sub) > 1:
                        ch = alt.Chart(sub).mark_line(point=True, color=colr).encode(
                            x=alt.X('Date:T', title=None),
                            y=alt.Y(f'{metric}:Q', scale=alt.Scale(zero=False), title='kg/hr'),
                            tooltip=['Date:T', f'{metric}:Q'])
                        col.markdown(f"**{title}**")
                        col.altair_chart(ch, use_container_width=True)
                    else:
                        col.info(f"Not enough history for {title} yet.")
            else:
                st.info("No dosing history in the registry yet. Upload Chemical Doses data to build a trend.")
        else:
            st.info("No dosing history in the registry yet.")

    # --- TAB 5: MRA EVALUATION ENGINE ---
    with tabs[5]:
        st.subheader("Production Prediction")
        st.markdown("Modify process inputs to execute 'What-If' scenarios. Input limits dynamically unbind to prevent system crashes.")
        controls_col, calc_col = st.columns([1, 2])
        
        with controls_col:
            st.number_input("1st effect vapour pressure (mmHg)", key="t5_press", on_change=sync_var, args=('mra_press', 't5_press'))
            st.number_input("1st Effect Vapour Temp (°C)", key="t5_t1", on_change=sync_var, args=('mra_t1', 't5_t1'))
            st.number_input("Sea Water Upper (m³/h)", key="t5_sw_up", on_change=sync_var, args=('sw_upper', 't5_sw_up'))
            st.number_input("1st effect brine temp (°C)", key="t5_bt1", on_change=sync_var, args=('mra_bt1', 't5_bt1'))
            st.number_input("Brine Water Return (m³/h)", key="t5_bflow", on_change=sync_var, args=('brine_ret', 't5_bflow'))
            st.number_input("LP Steam consumption (TPH)", key="t5_steam", on_change=sync_var, args=('steam', 't5_steam'))
            st.number_input("Antiscalant PPM", key="t5_anti", on_change=sync_var, args=('chem_anti_ppm', 't5_anti'))

        with calc_col:
            if mra_data.get('model_missing'):
                st.warning(
                    f"The trained {model_type} model file is not available on this server (it is stored on "
                    "local disk, which resets when the app restarts). The prediction below is using the "
                    "baseline OLS formula instead. Retrain the model on the Model tab to restore it."
                )
            k1, k2, k3 = st.columns(3)
            k1.metric("Actual Gross SCADA", f"{mra_data['Actual']:.1f} m³/h")
            _pred_label = "OLS fallback" if mra_data.get('model_missing') else model_type
            k2.metric(f"Predicted ({_pred_label})", f"{mra_data['Predicted']:.1f} m³/h")
            
            diff_pct = (mra_data['Residual'] / mra_data['Predicted']) * 100 if mra_data['Predicted'] > 0 else 0
            if diff_pct <= -5.0: 
                k3.error(f"Residual Gap: {mra_data['Residual']:.1f} TPH ({diff_pct:.1f}%) - Shutdown/Acid Clean Required")
            elif diff_pct <= -4.0: 
                k3.warning(f"Residual Gap: {mra_data['Residual']:.1f} TPH ({diff_pct:.1f}%) - Optimize Scale Treatment Dosing")
            else: 
                k3.success(f"Residual Gap: {mra_data['Residual']:.1f} TPH ({diff_pct:.1f}%) - Operational Thermal Base Clean")
                
            if model_type != "OLS": 
                st.info("Machine Learning Evaluation Mode Active: Multi-variable parameter expansion is only available under pure linear OLS logic.")
            st.dataframe(mra_data['Variance_DF'].style.format({"Baseline": "{:.1f}", "Live Input": "{:.1f}", "Deviation": "{:+.1f}", "Regression Weight": "{:.3f}", "Impact (TPH)": "{:+.1f}"}, na_rep="-"), use_container_width=True, hide_index=True)

    # --- TAB 6: REPORTING & ANALYTICS ---
    with tabs[6]:
        st.subheader("Reports & Historical Data")
        rep_tabs = st.tabs(["Daily Execution Dashboard", "Master Historical Database", "Long-Term Performance Trends", "Interactive Explorer"])
        
        with rep_tabs[0]:
            m_col1, m_col2, m_col3, m_col4 = st.columns(4)
            m_col1.metric("Target Record Date", log_date.strftime('%d-%m-%Y')) 
            m_col2.metric("Gross Volumetric Production", f"{ops_data['Gross Prod']} m³/h", delta=f"{ops_data['Gross Prod'] - 1000:.0f} from Design" if ops_data['Gross Prod'] < 1000 else None)
            m_col3.metric("System GOR", f"{ops_data['GOR']:.2f}", delta=f"{ops_data['GOR'] - 10.5:.2f} from Target" if ops_data['GOR'] < 10.5 else None)
            
            diff_pct = (mra_data['Residual'] / mra_data['Predicted']) * 100 if mra_data['Predicted'] > 0 else 0
            if diff_pct <= -5.0: 
                delta_text, d_color = f"{diff_pct:.1f}% (Scaling Critical)", "inverse"
            elif diff_pct <= -4.0: 
                delta_text, d_color = f"{diff_pct:.1f}% (Deviation Warning)", "inverse"
            else: 
                delta_text, d_color = f"{diff_pct:.1f}% (Clean Baseline)", "normal"
                
            m_col4.metric("Twin MRA Performance Gap", f"{mra_data['Residual']:.1f} TPH", delta=delta_text, delta_color=d_color)
            
            st.divider()
            graph_col1, graph_col2 = st.columns(2)
            with graph_col1:
                if model_type == "OLS":
                    st.markdown("#### Parameter Deviation Impact (m³/h)")
                    impact_chart = alt.Chart(mra_data['Variance_DF']).mark_bar().encode(
                        x=alt.X('Impact (TPH):Q'), 
                        y=alt.Y('Parameter:N', sort='-x', title=''), 
                        color=alt.condition(alt.datum['Impact (TPH)'] > 0, alt.value('#2ca02c'), alt.value('#d62728')), 
                        tooltip=['Parameter', 'Impact (TPH)']
                    ).properties(height=300)
                    st.altair_chart(impact_chart, use_container_width=True)
                else:
                    st.markdown("#### Component Weight Importance (ML Mode)")
                    impact_chart = alt.Chart(mra_data['Variance_DF']).mark_bar(color='#1f77b4').encode(
                        x=alt.X('Regression Weight:Q', title="Importance Weight Matrix %"), 
                        y=alt.Y('Parameter:N', sort='-x', title=''), 
                        tooltip=['Parameter', 'Regression Weight']
                    ).properties(height=300)
                    st.altair_chart(impact_chart, use_container_width=True)

            with graph_col2:
                st.markdown("#### Mass Distribution Profile")
                unaccounted = ops_data['SW Total'] - (ops_data['Desal'] + ops_data['Brine Return'])
                mb_data = pd.DataFrame({'Stream': ['Product Net', 'Brine Blowdown', 'Loss Matrix'], 'Volume': [ops_data['Desal'], ops_data['Brine Return'], unaccounted if unaccounted > 0 else 0]})
                donut = alt.Chart(mb_data).mark_arc(innerRadius=50).encode(
                    theta=alt.Theta("Volume:Q"), 
                    color=alt.Color("Stream:N", scale=alt.Scale(scheme='set2')), 
                    tooltip=['Stream', 'Volume']
                ).properties(height=300)
                st.altair_chart(donut, use_container_width=True)

            st.divider()
            st.text_area("Remarks & Performance Observations", key="in_remarks", on_change=sync_var, args=('remarks', 'in_remarks'), placeholder="Record operational shift anomalies, sensor calibrations, or clean notes here...")
            
            st.markdown("### Record and Commit Log Payload")
            c_pwd, c_save, c_export, c_csv = st.columns([1.5, 1, 1, 1])
            with c_pwd: 
                pwd_append = st.text_input("Security Key Access", type="password", key="pwd_append", label_visibility="collapsed", placeholder="Enter Master Security Password to Commit")
            with c_save:
                if st.button("Save Operational Record", use_container_width=True):
                    if pwd_append == "12345678":
                        db_dict = {
                            "Date": [log_date_str], 
                            "Sea Water Upper": [get_v('sw_upper')], 
                            "Sea Water Lower": [get_v('sw_lower')],
                            "Sea Water Feed": [ops_data['SW Total']], 
                            "Sea Water Pressure": [get_v('sw_press')],
                            "Brine Water Return": [ops_data['Brine Return']], 
                            "Desal production": [ops_data['Desal']], 
                            "LP Steam consumption": [ops_data['Steam']],
                            "LP Steam Pressure": [get_v('stm_press')],
                            "Condensate Return": [get_v('cond_flow')], 
                            "condensate temp": [get_v('cond_temp')],
                            "Condensate Conductivity": [get_v('cond_cond')],
                            "1st Effect Vapour Temp": [get_v('mra_t1')], 
                            "1st effect brine temp": [get_v('mra_bt1')], 
                            "11th Effect Brine Temp": [get_v('brine_11')],
                            "Feed Temp to Cold Group": [get_v('feed_cold')],
                            "Intermediate Effects Avg Brine Temp": [get_v('mid_effects_temp')],
                            "Delta T": [ops_data['dt_1st']], 
                            "1st effect vapour pressure": [get_v('mra_press')], 
                            "Brine Discharge Temp": [get_v('brine_out_t')], 
                            "Brine Discharge Pressure": [get_v('brine_press')],
                            "Sea Water cond I/L temp": [get_v('sw_in_t')], 
                            "Sea Water Condenser O/L Temp": [get_v('sw_out_t')], 
                            "CW supply": [get_v('cw_supply')], 
                            "CW Return": [get_v('cw_return')], 
                            "CW Flow": [get_v('cw_flow')],
                            "Gross production": [ops_data['Gross Prod']],
                            "GOR": [round(ops_data['GOR'], 2)], 
                            "STEC": [round(ops_data['STEC'], 2)],
                            "Overall HTC": [round(ops_data['htc_overall'], 2)], 
                            "1st Effect HTC": [round(ops_data['htc_1st'], 2)], 
                            "Residual": [round(mra_data['Residual'], 1)], 
                            "Antiscalant (kg)": [chem_data['anti_cons']], 
                            "Antifoam (kg)": [chem_data['foam_cons']], 
                            "Anti_PPM": [get_v('chem_anti_ppm')], 
                            "Foam_PPM": [get_v('chem_foam_ppm')], 
                            "Area_1st": [get_v('area_1st')], 
                            "Area_Overall": [get_v('area_overall')], 
                            "Remarks": [get_v('remarks')]
                        }
                        for cat in ['Feed', 'Product']:
                            for param, details in WATER_SPECS[cat].items(): 
                                db_dict[details['db_col']] = [get_v(details['var'])]
                        for param, details in BRINE_SPECS.items():
                            db_dict[details['db_col']] = [get_v(details['var'])]

                        # Persist the HTC sheets' own inputs + derived values so a manually-entered day
                        # appears on the HTC trends exactly like a bulk-uploaded one.
                        db_dict.update({
                            "Steam Inlet Temp": [get_v('steam_in_t')],
                            "HTC1_Feed_Flow": [get_v('htc1_feed_flow')],
                            "HTC1_Steam_TPH": [get_v('steam')],
                            "HTC1_Feed_Temp_Eff4to7": [get_v('mid_effects_temp')],
                            "HTC1_Brine_Temp": [get_v('mra_bt1')],
                            "HTC1_Vapor_Temp": [get_v('mra_t1')],
                            "HTC1_Cond_Temp": [get_v('cond_temp')],
                            "HTC1_dT1": [round(ops_data['dt_1st'], 3)],
                            "HTC1_dT2": [round(ops_data['dt2_1st'], 3)],
                            "HTC1_LMTD": [round(ops_data['lmtd_1st'], 3)],
                            "HTC1_Rf": [round(ops_data['rf_1st'], 8)],
                            "HTCO_Feed_Flow": [get_v('sw_total')],
                            "HTCO_Steam_TPH": [get_v('steam')],
                            "HTCO_Feed_Temp_ColdGrp": [get_v('feed_cold')],
                            "HTCO_Brine_Disch_Temp": [get_v('brine_out_t')],
                            "HTCO_Vapor_Temp": [get_v('mra_t1')],
                            "HTCO_Cond_Temp": [get_v('cond_temp')],
                            "HTCO_dT1": [round(ops_data['dt1_overall'], 3)],
                            "HTCO_dT2": [round(ops_data['dt2_overall'], 3)],
                            "HTCO_LMTD": [round(ops_data['lmtd_overall'], 3)],
                            "HTCO_Rf": [round(ops_data['rf_overall'], 8)],
                        })
                        
                        new_log = pd.DataFrame(db_dict)
                        st.session_state.daily_logs = pd.concat([st.session_state.daily_logs, new_log], ignore_index=True)
                        
                        # MASTER DATE FIX: Standardize before dropping duplicates to eradicate "ghost" format duplication
                        st.session_state.daily_logs['Date'] = standardize_dates(st.session_state.daily_logs['Date']).dt.strftime('%Y-%m-%d')
                        st.session_state.daily_logs = st.session_state.daily_logs.dropna(subset=['Date'])
                        st.session_state.daily_logs = st.session_state.daily_logs.drop_duplicates(subset=['Date'], keep='last').reset_index(drop=True)
                        
                        save_database(db_conn, st.session_state.daily_logs, LOCAL_DB_FILE)
                        st.success("Operational record successfully integrated into file engine!")
                        time.sleep(1.0)
                        st.rerun()  
                    elif pwd_append != "": 
                        st.error("Master verification credential failed.")
            with c_export:
                # Generate only on demand. st.download_button needs its data upfront, so calling the
                # report generator inline rebuilt the entire Word document on EVERY script rerun -
                # i.e. on every keystroke while this tab was open. That was the single heaviest
                # per-interaction cost in the app and a direct contributor to CPU throttling.
                if st.button("Generate Word Report (.docx)", use_container_width=True, key="gen_daily_docx"):
                    with st.spinner("Building report..."):
                        st.session_state.daily_docx = generate_comprehensive_report(
                            log_date, ops_data, sor_export_dfs, water_data, chem_data,
                            mra_data, get_v('skip_wq'), get_v('remarks'))
                        st.session_state.daily_docx_date = log_date_str
                if st.session_state.get('daily_docx') and st.session_state.get('daily_docx_date') == log_date_str:
                    st.download_button("Download Word Report", data=st.session_state.daily_docx,
                                       file_name=f"MED4_ExecutiveReport_{log_date_str}.docx",
                                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                       use_container_width=True)
                st.caption(f"Report engine {REPORT_VERSION}")
            
            with c_csv:
                csv_file = generate_daily_csv(log_date, ops_data, water_data, chem_data, mra_data, st.session_state.vars)
                st.download_button("Export Tabular Values (.csv)", data=csv_file, file_name=f"MED4_DataRecord_{log_date_str}.csv", mime="text/csv", use_container_width=True)

        with rep_tabs[1]:
            st.markdown("#### Master System Registry Database")
            display_cols = [c for c in EXACT_DB_COLUMNS if c in st.session_state.daily_logs.columns]
            edited_db = st.data_editor(st.session_state.daily_logs[display_cols] if not st.session_state.daily_logs.empty else st.session_state.daily_logs, num_rows="dynamic", use_container_width=True)
            c_sync_pwd, c_sync, c_dl = st.columns([2, 1, 1])
            with c_sync_pwd: 
                pwd_sync = st.text_input("Database Write-Access Password", type="password", key="pwd_sync", label_visibility="collapsed", placeholder="Enter Database Master Password to Save Modifications")
            with c_sync:
                if st.button("Synchronize Registry", use_container_width=True):
                    if pwd_sync == "12345678":
                        # MASTER DATE FIX: Standardize manually edited database
                        edited_db['Date'] = standardize_dates(edited_db['Date']).dt.strftime('%Y-%m-%d')
                        st.session_state.daily_logs = edited_db.dropna(subset=['Date']).drop_duplicates(subset=['Date'], keep='last').reset_index(drop=True)
                        
                        save_database(db_conn, st.session_state.daily_logs, LOCAL_DB_FILE)
                        st.success("Master registry records updated successfully!")
                    else: 
                        st.error("System modification credentials failed.")
            with c_dl:
                st.download_button("Download Database Offline Backup", data=st.session_state.daily_logs.to_csv(index=False).encode('utf-8'), file_name=f"MED4_MasterRegistry_Backup.csv", mime='text/csv', use_container_width=True)

            st.divider()
            st.markdown("#### Aggregated Monthly Performance Generator")
            if not st.session_state.daily_logs.empty:
                df_logs = st.session_state.daily_logs.copy()
                
                df_logs['Date'] = standardize_dates(df_logs['Date'])
                df_logs = df_logs.dropna(subset=['Date'])
                
                month_data = df_logs[(df_logs['Date'].dt.month == log_date.month) & (df_logs['Date'].dt.year == log_date.year)].copy()
                if not month_data.empty:
                    if st.button("Compile and Generate Monthly Summary (.docx)", use_container_width=True):
                        monthly_doc = generate_monthly_report(month_data, log_date.strftime('%B'), str(log_date.year))
                        st.download_button("Download Monthly Briefing Document", data=monthly_doc, file_name=f"MED4_MonthlySummary_{log_date.strftime('%b_%Y')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        with rep_tabs[2]:
            if not st.session_state.daily_logs.empty:
                df_logs = st.session_state.daily_logs.copy()
                
                df_logs['Date'] = standardize_dates(df_logs['Date'])
                df_logs = df_logs.dropna(subset=['Date'])
                
                if not df_logs.empty:
                    df_logs['Total SW Feed (m3/h)'] = pd.to_numeric(df_logs.get('Sea Water Feed', 0), errors='coerce')
                    df_logs['Recovery (%)'] = np.where(df_logs['Total SW Feed (m3/h)'] > 0, (pd.to_numeric(df_logs.get('Gross production', 0), errors='coerce') / df_logs['Total SW Feed (m3/h)']) * 100, 0)
                    
                    df_logs['Actual Production'] = pd.to_numeric(df_logs.get('Gross production', 0), errors='coerce')
                    df_logs['Residual_Val'] = pd.to_numeric(df_logs.get('Residual', 0), errors='coerce')
                    df_logs['Predicted Production'] = df_logs['Actual Production'] - df_logs['Residual_Val']
                    df_logs['Overall_HTC_Val'] = pd.to_numeric(df_logs.get('Overall HTC', 0), errors='coerce')
                    df_logs['GOR_Val'] = pd.to_numeric(df_logs.get('GOR', 0), errors='coerce')
                    df_logs['STEC_Val'] = pd.to_numeric(df_logs.get('STEC', np.nan), errors='coerce')
                    
                    min_date = df_logs['Date'].min().date() 
                    max_date = df_logs['Date'].max().date()
                    
                    st.markdown("##### Performance Evaluation Horizon Filter")
                    d_col1, d_col2 = st.columns(2)
                    with d_col1: 
                        start_date = st.date_input("Start Threshold Date", min_date, key="start_d1")
                    with d_col2: 
                        end_date = st.date_input("End Threshold Date", max_date, key="end_d1")
                    
                    mask = (df_logs['Date'].dt.date >= start_date) & (df_logs['Date'].dt.date <= end_date)
                    df_filtered = df_logs.loc[mask]
                    
                    q_col1, q_col2 = st.columns(2)
                    with q_col1:
                        st.markdown("#### Performance Recovery Rate Deviation Trend")
                        if len(df_filtered) > 1:
                            rec_chart = alt.Chart(df_filtered).mark_circle().encode(x=alt.X('Date:T', title="Evaluation Timeline"), y=alt.Y('Recovery (%):Q', scale=alt.Scale(zero=False)))
                            st.altair_chart(rec_chart + rec_chart.transform_regression('Date', 'Recovery (%)').mark_line(color='red'), use_container_width=True)
                    with q_col2:
                        st.markdown("#### Seawater Coefficient Degradation Rate (HTC)")
                        if len(df_filtered) > 1:
                            htc_chart = alt.Chart(df_filtered).mark_line(point=True, color='orange').encode(x=alt.X('Date:T', title="Evaluation Timeline"), y=alt.Y('Overall_HTC_Val:Q', scale=alt.Scale(zero=False), title="Overall HTC (W/m²K)"))
                            st.altair_chart(htc_chart + htc_chart.transform_regression('Date', 'Overall_HTC_Val').mark_line(color='black'), use_container_width=True)

                    st.divider()
                    
                    q_col3, q_col4 = st.columns(2)
                    with q_col3:
                        st.markdown("#### Actual Mass Output vs Normalized Twin Output")
                        if len(df_filtered) > 1:
                            fold_df = df_filtered[['Date', 'Actual Production', 'Predicted Production']].melt('Date', var_name='Metric', value_name='Mass Flow Volume (m³/h)')
                            prod_chart = alt.Chart(fold_df).mark_line(point=True).encode(
                                x=alt.X('Date:T', title="Evaluation Timeline"), y=alt.Y('Mass Flow Volume (m³/h):Q', scale=alt.Scale(zero=False)),
                                color=alt.Color('Metric:N', scale=alt.Scale(domain=['Actual Production', 'Predicted Production'], range=['#1f77b4', '#ff7f0e'])),
                                strokeDash=alt.condition(alt.datum.Metric == 'Predicted Production', alt.value([5, 5]), alt.value([0])),
                                tooltip=['Date:T', 'Metric', 'Mass Flow Volume (m³/h)']
                            )
                            st.altair_chart(prod_chart, use_container_width=True)
                    with q_col4:
                        st.markdown("#### Specific Unit Thermal Efficiency GOR Performance")
                        if len(df_filtered) > 1:
                            gor_chart = alt.Chart(df_filtered).mark_line(point=True, color='green').encode(
                                x=alt.X('Date:T', title="Evaluation Timeline"), y=alt.Y('GOR_Val:Q', scale=alt.Scale(zero=False), title="Gain Output Ratio"),
                                tooltip=['Date:T', 'GOR_Val']
                            )
                            st.altair_chart(gor_chart + gor_chart.transform_regression('Date', 'GOR_Val').mark_line(color='red', strokeDash=[5, 5]), use_container_width=True)

                    st.divider()

                    st.markdown("#### Specific Thermal Energy Consumption (STEC) Trend")
                    df_stec = df_filtered.dropna(subset=['STEC_Val'])
                    if len(df_stec) > 1:
                        stec_chart = alt.Chart(df_stec).mark_line(point=True, color='purple').encode(
                            x=alt.X('Date:T', title="Evaluation Timeline"), y=alt.Y('STEC_Val:Q', scale=alt.Scale(zero=False), title="STEC (kWh/ton)"),
                            tooltip=['Date:T', 'STEC_Val']
                        )
                        st.altair_chart(stec_chart + stec_chart.transform_regression('Date', 'STEC_Val').mark_line(color='black', strokeDash=[5, 5]), use_container_width=True)
                    else:
                        st.info("No STEC data available yet for the selected range. Rows saved before this update won't have a stored STEC value.")

                    # ---- Report for the selected horizon -------------------------------------
                    # Uses the SAME start/end filter driving the charts above, so the document
                    # always covers exactly what is on screen. Generated straight into a download
                    # button (no intermediate "compile" click) so it is one action, not two.
                    st.divider()
                    st.markdown("#### Performance Report for the Selected Period")

                    if df_filtered.empty:
                        st.info("No records fall inside the selected dates, so there is nothing to report on yet.")
                    else:
                        _n_run = int((pd.to_numeric(df_filtered.get('Gross production', 0), errors='coerce').fillna(0) > 0).sum())
                        _span_days = (end_date - start_date).days + 1

                        # If the selection happens to be exactly one whole calendar month, label it as a
                        # monthly report so it reads naturally; otherwise use the explicit date span.
                        _is_full_month = (
                            start_date.day == 1
                            and end_date.month == start_date.month
                            and end_date.year == start_date.year
                            and (end_date + datetime.timedelta(days=1)).month != end_date.month
                        )
                        if _is_full_month:
                            _label = start_date.strftime('%B %Y')
                            _kind = "Monthly"
                            _fname = f"MED4_MonthlyReport_{start_date.strftime('%b_%Y')}.docx"
                        else:
                            _label = f"{start_date.strftime('%d %B %Y')} to {end_date.strftime('%d %B %Y')}"
                            _kind = "Period"
                            _fname = f"MED4_PeriodReport_{start_date.strftime('%Y%m%d')}_{end_date.strftime('%Y%m%d')}.docx"

                        rc1, rc2, rc3 = st.columns(3)
                        rc1.metric("Days in selection", _span_days)
                        rc2.metric("Records in range", len(df_filtered))
                        rc3.metric("Operating days", _n_run)

                        st.caption(f"Report period: **{_label}**")

                        if _n_run == 0:
                            st.warning(
                                "None of the records in this range show production above zero, so the report "
                                "would have no operating days to average. Widen the dates or check that "
                                "production data was logged for this period."
                            )
                        else:
                            if _n_run < 6:
                                st.info(
                                    f"Only {_n_run} operating days in range. The report will still generate, but "
                                    "the opening-versus-closing fouling trend needs at least 6 days and will be "
                                    "left out."
                                )
                            try:
                                _period_doc = generate_period_report(df_filtered, _label, period_kind=_kind)
                                st.download_button(
                                    "Download Performance Report (.docx)",
                                    data=_period_doc,
                                    file_name=_fname,
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    use_container_width=True,
                                    type="primary",
                                    key="lt_period_report_dl"
                                )
                            except Exception as _e:
                                st.error(f"Report could not be generated: {_e}")
                else:
                    st.info("No valid dates found in registry to draw charts.")

        with rep_tabs[3]:
            st.markdown("#### Multivariable Cross-Correlation Explorer")
            if not st.session_state.daily_logs.empty:
                exp_df = st.session_state.daily_logs.copy()
                
                exp_df['Date'] = standardize_dates(exp_df['Date'])
                exp_df = exp_df.dropna(subset=['Date'])
                
                if not exp_df.empty:
                    min_date2 = exp_df['Date'].min().date() 
                    max_date2 = exp_df['Date'].max().date()
                    
                    d_col1, d_col2 = st.columns(2)
                    with d_col1: 
                        start_date2 = st.date_input("Start Horizon Date", min_date2, key="start_d2")
                    with d_col2: 
                        end_date2 = st.date_input("End Horizon Date", max_date2, key="end_d2")
                    
                    mask2 = (exp_df['Date'].dt.date >= start_date2) & (exp_df['Date'].dt.date <= end_date2)
                    exp_df = exp_df.loc[mask2]
                    
                    num_cols = [col for col in exp_df.columns if col not in ['Date']]
                    x_c, y_c, t_c = st.columns(3)
                    with x_c: 
                        exp_x = st.selectbox("Select Independent Domain X-Axis", ['Date'] + num_cols, index=0)
                    with y_c: 
                        exp_y = st.selectbox("Select Dependent Variable Y-Axis", num_cols, index=0)
                    with t_c: 
                        exp_type = st.selectbox("Select Functional Chart Variant", ["Line Chart", "Scatter Plot", "Bar Chart"])
                    
                    if exp_type == "Line Chart": 
                        chart = alt.Chart(exp_df).mark_line(point=True).encode(x=alt.X(f"{exp_x}{':T' if exp_x == 'Date' else ':Q'}"), y=alt.Y(f"{exp_y}:Q", scale=alt.Scale(zero=False)), tooltip=[exp_x, exp_y])
                    elif exp_type == "Scatter Plot": 
                        chart = alt.Chart(exp_df).mark_circle(size=80).encode(x=alt.X(f"{exp_x}{':T' if exp_x == 'Date' else ':Q'}"), y=alt.Y(f"{exp_y}:Q", scale=alt.Scale(zero=False)), tooltip=[exp_x, exp_y])
                    else: 
                        chart = alt.Chart(exp_df).mark_bar().encode(x=alt.X(f"{exp_x}{':T' if exp_x == 'Date' else ':N'}"), y=alt.Y(f"{exp_y}:Q"), tooltip=[exp_x, exp_y])
                    st.altair_chart(chart.interactive(), use_container_width=True)
                else:
                    st.info("No active historical registry values detected to perform correlation modeling.")

    # --- TAB 7: AI MODEL SELECTOR ---
    with tabs[7]:
        st.subheader("Prediction Model Setup")
        if not SKLEARN_INSTALLED:
            st.error("Mathematical package 'scikit-learn' is missing from file dependencies.")
        else:
            from sklearn.linear_model import LinearRegression
            from sklearn.ensemble import RandomForestRegressor
            from sklearn.metrics import r2_score
            
            st.warning("Ephemeral Server Parameter Caution: Since this tracking node runs on temporary testing cloud containers, manual machine-learning logic selection targets revert back to historical OLS baseline models after inactive shutdown flags are generated. Selection options remain permanently hardlocked upon local internal node integration.")
            st.markdown("### Manage Baseline Evaluation Multipliers")
            st.markdown(f"**Current Evaluator Logic Subroutine:** `{model_type}`")
            c_reset, _ = st.columns([1, 1])
            with c_reset:
                if st.button("Reset to Default Coefficients", use_container_width=True):
                    st.session_state.mra_coef = MED_MRA_COEF_DEFAULT.copy()
                    _ok = save_config(db_conn, st.session_state.mra_coef, LOCAL_CONFIG_FILE)
                    if _ok:
                        st.success("Coefficients reset and saved to the cloud sheet.")
                    else:
                        st.warning("Coefficients reset, but the cloud sheet was unreachable - saved locally only.")
                    time.sleep(1.5)
                    st.rerun()

            st.divider()
            st.markdown("### Prediction Model Builder")
            st.markdown(
                "Calibrate the prediction model against a period of known-good operation. Anchoring to a "
                "cleaning date means the baseline represents this plant when clean, so later deviations "
                "measure real fouling rather than drift from a generic design assumption."
            )

            # Both the required columns and the predictor count come from MED_MRA_PARAMS, so the
            # template, the validation and the fit can never disagree about what the model expects.
            req_cols = ["Date", "Gross production"] + MRA_DB_COLUMNS
            N_PREDICTORS = N_MRA_PREDICTORS

            if not bool(st.session_state.mra_coef.get("calibrated", 0)):
                st.warning(
                    f"**This model has not been calibrated yet.** It is running on a provisional "
                    f"coefficient set: the six inputs carried over from the previous model keep their "
                    f"old values, and the four new inputs ({', '.join(MRA_LABELS[4:5] + MRA_LABELS[7:])}) "
                    f"currently contribute nothing at all. Predictions are indicative only until you "
                    f"calibrate below against real plant data."
                )

            calib_src = st.radio(
                "Calibration data source",
                ["Historical data after a cleaning", "Upload a CSV"],
                horizontal=True, key="calib_source"
            )

            uploaded_file = None

            if calib_src == "Historical data after a cleaning":
                cc1, cc2 = st.columns([2, 1])
                with cc1:
                    clean_date = st.date_input(
                        "Last cleaning date", value=datetime.date.today() - datetime.timedelta(days=90),
                        format="DD/MM/YYYY", key="calib_clean_date",
                        help="Calibration uses operating data starting the day after this cleaning."
                    )
                with cc2:
                    window_days = st.selectbox(
                        "Window", [30, 60, 90, 120, 180, 365], index=3, key="calib_window",
                        help="Days of post-cleaning data to calibrate on. The model now fits "
                             f"{N_MRA_PREDICTORS} inputs, so roughly {N_MRA_PREDICTORS * 10} usable "
                             "rows are needed for a stable fit - noticeably more than the previous "
                             "7-input model required."
                    )

                hist = st.session_state.daily_logs
                if hist is None or hist.empty or 'Date' not in hist.columns:
                    st.info("No historical data in the registry yet.")
                else:
                    h = hist.copy()
                    h['_d'] = standardize_dates(h['Date'])
                    start = pd.Timestamp(clean_date)
                    end = start + pd.Timedelta(days=int(window_days))
                    h = h.dropna(subset=['_d'])
                    h = h[(h['_d'] > start) & (h['_d'] <= end)]

                    missing_cols = [c for c in req_cols if c != "Date" and c not in h.columns]
                    if missing_cols:
                        st.error(f"The registry is missing required columns: {', '.join(missing_cols)}")
                    else:
                        sel = h[['_d'] + [c for c in req_cols if c != "Date"]].copy()
                        sel.rename(columns={'_d': 'Date'}, inplace=True)
                        for c in req_cols:
                            if c != "Date":
                                sel[c] = pd.to_numeric(sel[c], errors='coerce')
                        # A logged-but-empty row is stored as 0; those are not real operating points
                        # and would drag the regression toward the origin, so drop them.
                        sel = sel.dropna(subset=[c for c in req_cols if c != "Date"])
                        sel = sel[(sel['Gross production'] > 0) & (sel['LP Steam consumption'] > 0)]

                        n_rows = len(sel)
                        ratio = n_rows / N_PREDICTORS if N_PREDICTORS else 0
                        m1, m2, m3 = st.columns(3)
                        m1.metric("Usable rows", n_rows)
                        m2.metric("Parameters fitted", N_PREDICTORS)
                        m3.metric("Rows per parameter", f"{ratio:.1f}")

                        if n_rows == 0:
                            st.warning("No usable rows in that window. Check the cleaning date, or upload data for this period first.")
                        elif ratio < 5:
                            st.error(
                                f"Only {ratio:.1f} rows per parameter. This is far too few - the model would "
                                "memorise noise instead of learning plant behaviour, and later predictions would "
                                "drift unpredictably. Use a longer window or a period with more complete logging."
                            )
                        elif ratio < 10:
                            st.warning(
                                f"{ratio:.1f} rows per parameter is on the low side (10+ is comfortable). The fit "
                                "may look better than it really is. Prefer a longer window if the plant stayed "
                                "clean, and treat a very high R² here with suspicion."
                            )
                        else:
                            st.success(f"{n_rows} usable rows covering {window_days} days after cleaning.")

                        if n_rows > 0:
                            sel['Date'] = sel['Date'].dt.strftime('%Y-%m-%d')
                            with st.expander("Preview calibration data"):
                                st.dataframe(sel, use_container_width=True, hide_index=True)
                            # Hand the selected slice to the existing training pipeline as an in-memory CSV,
                            # so the upload path and this path share identical validation and fitting code.
                            uploaded_file = io.StringIO(sel.to_csv(index=False))
            else:
                template_df = pd.DataFrame(columns=req_cols)
                st.download_button(label="Download Training Template", data=template_df.to_csv(index=False).encode('utf-8'), file_name='MED4_ML_CalibrationTemplate.csv', mime='text/csv')
                st.divider()
                uploaded_file = st.file_uploader("Upload Training Data", type=["csv"], key="mra_trainer")
            
            if uploaded_file is not None:
                try:
                    df_train = pd.read_csv(uploaded_file)
                    if not all(col in df_train.columns for col in req_cols): 
                        st.error("The data is missing one or more required columns.")
                    else:
                        for col in req_cols:
                            if col != "Date":
                                if df_train[col].dtype == object: 
                                    df_train[col] = pd.to_numeric(df_train[col].astype(str).str.replace(',', '', regex=False), errors='coerce')
                        
                        df_train = df_train.dropna(subset=[c for c in req_cols if c != "Date"])
                        st.caption(f"Fitting models on {len(df_train)} rows.")
                        
                        if len(df_train) > 0:
                            X = df_train[MRA_DB_COLUMNS].copy()
                            X.columns = MRA_COEF_KEYS   # fit under the coefficient keys the live path uses
                            Y = df_train["Gross production"]
                            
                            model_ols = LinearRegression(fit_intercept=True).fit(X, Y)
                            r2_ols = r2_score(Y, model_ols.predict(X))

                            # Ridge is fitted on standardised inputs, then the coefficients are
                            # converted back to raw units so they drop straight into the same live
                            # prediction path as OLS. Ridge exists here specifically because several
                            # of these predictors are near-duplicates of each other physically; the
                            # penalty shares weight between correlated inputs instead of letting OLS
                            # assign them huge offsetting coefficients.
                            model_ridge = None
                            r2_ridge = np.nan
                            ridge_block = None
                            try:
                                from sklearn.linear_model import RidgeCV
                                _mu, _sd = X.mean(), X.std(ddof=0).replace(0, 1.0)
                                Xs = (X - _mu) / _sd
                                model_ridge = RidgeCV(alphas=np.logspace(-3, 3, 25)).fit(Xs, Y)
                                r2_ridge = r2_score(Y, model_ridge.predict(Xs))
                                _raw = model_ridge.coef_ / _sd.values
                                ridge_block = {"Intercept": float(model_ridge.intercept_ - np.sum(_raw * _mu.values))}
                                for _i, _k in enumerate(MRA_COEF_KEYS):
                                    ridge_block[_k] = float(_raw[_i])
                            except Exception:
                                model_ridge = None

                            model_rf = RandomForestRegressor(n_estimators=100, random_state=42).fit(X, Y)
                            r2_rf = r2_score(Y, model_rf.predict(X))
                            
                            if XGB_INSTALLED:
                                import xgboost as xgb
                                model_xgb = xgb.XGBRegressor(n_estimators=100, random_state=42).fit(X, Y)
                                r2_xgb = r2_score(Y, model_xgb.predict(X))
                            
                            st.markdown("### Algorithm Accuracy Evaluation Matrix")
                            m1, m2, m3, m4 = st.columns(4)
                            m1.metric("1. Linear OLS Fit (R² Coefficient)", f"{r2_ols * 100:.2f}%")
                            m2.metric("2. Ridge (Penalised Linear) R²", f"{r2_ridge * 100:.2f}%" if pd.notna(r2_ridge) else "n/a")
                            m3.metric("3. Random Forest Tree Logic (R²)", f"{r2_rf * 100:.2f}%")
                            if XGB_INSTALLED:
                                m4.metric("4. Extreme Gradient Boost XGB (R²)", f"{r2_xgb * 100:.2f}%")
                            else:
                                m4.warning("Advanced Gradient boosting library dependency not activated.")

                            st.caption(
                                "These R² figures are measured on the same rows the models were trained on, so "
                                "they show fit, not predictive accuracy. Random Forest and XGBoost will almost "
                                "always look near-perfect here; that is memorisation, not skill."
                            )

                            # ---- Collinearity diagnostic ------------------------------------------
                            # Several of these predictors measure nearly the same physical thing:
                            # 1st effect pressure and vapour temperature are a saturation pair, steam
                            # flow and condensate flow are the same stream in and out, and the brine
                            # temperatures track each other closely. OLS can still report a high R²
                            # while assigning individual coefficients that are unstable and can flip
                            # sign - which would make the daily "Parameter Deviation Impact" chart
                            # actively misleading. VIF exposes that before the model is committed.
                            st.markdown("#### Predictor Independence Check (VIF)")
                            try:
                                _c = np.corrcoef(X.values, rowvar=False)
                                _vifs = np.diag(np.linalg.pinv(_c))
                                vif_df = pd.DataFrame({"Parameter": MRA_LABELS, "VIF": np.round(_vifs, 1)})
                                vif_df["Assessment"] = np.where(
                                    vif_df["VIF"] >= 10, "Severe - coefficient unreliable",
                                    np.where(vif_df["VIF"] >= 5, "Moderate - interpret with care", "Acceptable")
                                )
                                vif_df = vif_df.sort_values("VIF", ascending=False)
                                st.dataframe(vif_df, use_container_width=True, hide_index=True)

                                _n_severe = int((vif_df["VIF"] >= 10).sum())
                                if _n_severe:
                                    st.warning(
                                        f"{_n_severe} of {N_MRA_PREDICTORS} predictors show severe collinearity "
                                        "(VIF ≥ 10). The overall prediction stays usable, but the individual "
                                        "coefficients for those inputs are not trustworthy on their own and the "
                                        "per-parameter impact breakdown should not be read as cause and effect. "
                                        "Ridge is the safer choice here than plain OLS."
                                    )
                                else:
                                    st.success("No severe collinearity detected. OLS coefficients can be read individually.")
                            except Exception as _ve:
                                st.info(f"Collinearity check unavailable for this dataset ({_ve}).")
                            
                            st.markdown("#### Dynamic Feature Sensitivity Weights / Scaling Coefficients")
                            comp_dict = {
                                "Parameter": MRA_LABELS,
                                "OLS (Coefficients)": np.round(model_ols.coef_, 4),
                                "Random Forest (Importance %)": np.round(model_rf.feature_importances_ * 100, 2)
                            }
                            if ridge_block is not None:
                                comp_dict["Ridge (Coefficients)"] = np.round([ridge_block[k] for k in MRA_COEF_KEYS], 4)
                            if XGB_INSTALLED: 
                                comp_dict["XGBoost (Importance %)"] = np.round(model_xgb.feature_importances_ * 100, 2)
                            
                            st.dataframe(pd.DataFrame(comp_dict).style.format(precision=4), use_container_width=True, hide_index=True)
                            
                            st.markdown("### Commit & Lock Mathematical Subroutine Target")
                            opts = ["OLS (Linear)"]
                            if ridge_block is not None:
                                opts.append("Ridge (Penalised Linear)")
                            opts.append("Random Forest")
                            if XGB_INSTALLED:
                                opts.append("XGBoost")

                            selected_model = st.radio("Configure Active Live Prediction Logic Block:", opts)

                            if st.button("Confirm & Activate Model", type="primary", use_container_width=True):
                                # The OLS fit is always available here, so persist it EVERY time regardless
                                # of which model is selected. That guarantees a real calibrated fallback
                                # exists if an AI model is ever unavailable - previously selecting an AI
                                # model wiped the OLS coefficients and left only the factory baseline.
                                ols_block = {"Intercept": float(model_ols.intercept_)}
                                for _i, _k in enumerate(MRA_COEF_KEYS):
                                    ols_block[_k] = float(model_ols.coef_[_i])

                                # Store the calibration-period mean of every predictor. The deviation
                                # column on the daily dashboard is meant to measure against THIS plant
                                # in clean condition, so the reference has to come from the calibration
                                # window rather than a generic design figure.
                                base_block = {f"BASE_{_k}": float(X[_k].mean()) for _k in MRA_COEF_KEYS}

                                meta = {
                                    "calibrated": 1,
                                    "n_predictors": N_MRA_PREDICTORS,
                                    "calib_rows": int(len(df_train)),
                                    "calib_r2": float(r2_ols),
                                }

                                _blob_ok = True
                                if selected_model in ("OLS (Linear)", "Ridge (Penalised Linear)"):
                                    # Ridge coefficients have already been converted back to raw units,
                                    # so they slot into the identical live prediction path as OLS.
                                    block = ols_block if selected_model == "OLS (Linear)" else ridge_block
                                    new_coefs = dict(block)
                                    new_coefs.update(base_block)
                                    new_coefs.update(meta)
                                    new_coefs["model_type"] = "OLS"
                                    new_coefs["fit_method"] = selected_model
                                    if selected_model == "Ridge (Penalised Linear)":
                                        new_coefs["calib_r2"] = float(r2_ridge)
                                    st.session_state.mra_coef = new_coefs
                                    _ok = save_config(db_conn, new_coefs, LOCAL_CONFIG_FILE)
                                else:
                                    target_m = model_rf if selected_model == "Random Forest" else model_xgb
                                    # Feature importances are a DIFFERENT quantity from regression
                                    # coefficients, so they are stored under AI_-prefixed keys. Reusing the
                                    # plain names would silently corrupt the OLS block.
                                    ai_coefs = dict(ols_block)
                                    ai_coefs.update(base_block)
                                    ai_coefs.update(meta)
                                    ai_coefs["model_type"] = selected_model
                                    ai_coefs["fit_method"] = selected_model
                                    for _i, _k in enumerate(MRA_COEF_KEYS):
                                        ai_coefs[f"AI_{_k}"] = float(target_m.feature_importances_[_i])
                                    st.session_state.mra_coef = ai_coefs
                                    _ok = save_config(db_conn, ai_coefs, LOCAL_CONFIG_FILE)
                                    # Persist the fitted model itself so it survives container restarts.
                                    _blob_ok = save_model_blob(db_conn, target_m, LOCAL_CONFIG_FILE, AI_MODEL_FILE)

                                if _ok and _blob_ok:
                                    st.success(f"{selected_model} activated on {N_MRA_PREDICTORS} predictors. Calibration and model saved to the cloud sheet.")
                                elif _ok and not _blob_ok:
                                    st.warning(f"{selected_model} activated and calibration saved, but the model file could not be stored in the sheet - it may need retraining after a restart.")
                                else:
                                    st.warning(f"{selected_model} activated, but the cloud sheet was unreachable - saved locally only and may reset on restart.")
                                time.sleep(1.5)
                                st.rerun()
                        else: 
                            st.error("Structural data parsing produced empty float ranges inside parameters.")
                except Exception as e: 
                    st.error(f"Structural data matrix crash: {e}")

    # --- TAB 8: BULK EXCEL UPLOADER PANEL ---
    with tabs[8]:
        st.subheader("Bulk Data Upload")
        st.caption(
            "Each uploader mirrors ONE tab of the plant workbook exactly. Upload only the INPUT columns - "
            "every derived value (LMTD, HTC, GOR, STEC, Recovery, Fouling) is recomputed by the calculator "
            "and any such column in your file is ignored. Uploads are merged by date: loading HTC data never "
            "overwrites Operational or Water Quality data for the same day, and vice versa."
        )

        def _clean_num(df, cols):
            """Coerce to numeric. The plant sheets use '-' for 'not measured'; that must become NaN
            (genuinely missing), never 0, because a 0 temperature would silently corrupt an LMTD."""
            for c in cols:
                if c not in df.columns:
                    df[c] = np.nan
                df[c] = pd.to_numeric(
                    df[c].astype(str).str.replace(',', '', regex=False).str.strip().replace({'-': np.nan, '': np.nan}),
                    errors='coerce'
                )
            return df

        def _lmtd(dt1, dt2):
            """LMTD = (dT1 - dT2) / ln(dT1/dT2), exactly as in both HTC sheets (col N).
            Valid only when both driving forces are present and positive. Note dT2 > dT1 in this
            plant's data, which the formula handles fine (both numerator and log go negative)."""
            valid = dt1.notna() & dt2.notna() & (dt1 > 0) & (dt2 > 0) & (dt1 != dt2)
            ratio = np.where(valid, dt1 / dt2, 1.0)
            logr = np.log(np.where(ratio > 0, ratio, 1.0))
            return np.where(valid & (logr != 0), (dt1 - dt2) / logr, np.nan)

        def _backfill_from_db(d, mapping):
            """For any HTC input the uploaded file doesn't supply, fall back to the value already
            stored in the master registry for that same date (typically loaded by the Operational
            upload, which shares most of these readings). HTC is only left blank when a value is
            available NOWHERE - not merely because it was absent from this one file.
            mapping: {htc_column_in_this_file: master_registry_column}"""
            logs = st.session_state.daily_logs
            if logs is None or logs.empty or 'Date' not in logs.columns:
                return d, []
            ref = logs.copy()
            ref['Date'] = standardize_dates(ref['Date']).dt.strftime('%Y-%m-%d')
            ref = ref.dropna(subset=['Date']).drop_duplicates(subset=['Date'], keep='last').set_index('Date')
            filled = []
            for htc_col, db_col in mapping.items():
                if db_col not in ref.columns:
                    continue
                src = pd.to_numeric(d['Date'].map(ref[db_col]), errors='coerce')
                if htc_col not in d.columns:
                    d[htc_col] = np.nan
                n_before = int(d[htc_col].isna().sum())
                d[htc_col] = d[htc_col].fillna(src)
                n_after = int(d[htc_col].isna().sum())
                if n_before > n_after:
                    filled.append(f"{htc_col.split('_', 1)[-1].replace('_', ' ')} ({n_before - n_after})")
            return d, filled

        bulk_subtabs = st.tabs([
            "A) Operational Data", "B) 1st Effect HTC", "C) Overall HTC", "D) Water Quality", "E) Chemical Doses"
        ])

        # ===================================================================================
        # A) OPERATIONAL DATA  <-  'Operational data' sheet
        # ===================================================================================
        with bulk_subtabs[0]:
            st.markdown(
                "Source: **`Operational data`** tab. Upload the sheet as-is. The calculator recomputes "
                "**Recovery**, **Conversion**, **GOR**, **Steam Economy**, **Overall Delta T** and **STEC** "
                "from the raw readings - the versions of those columns already in your sheet are ignored."
            )
            st.download_button(
                "Download Operational Template", key='dl_op',
                data=pd.DataFrame(columns=OPERATIONAL_BULK_HEADERS).to_csv(index=False).encode('utf-8'),
                file_name='MED4_Operational_Template.csv', mime='text/csv'
            )
            st.divider()
            op_file = st.file_uploader("Upload Operational Data CSV", type=["csv"], key="op_up")

            if op_file is not None:
                try:
                    d = pd.read_csv(op_file)
                    if 'Parameter' in d.columns:
                        d = d[~d['Parameter'].astype(str).isin(['Design', 'Unit', 'TAG', 'SOR/  Base line'])]
                    d.rename(columns={
                        'Parameter': 'Date',
                        'Sea water Upper': 'Sea Water Upper', 'Sea water Lower': 'Sea Water Lower',
                        'Sea water feed': 'Sea Water Feed', 'Brine return': 'Brine Water Return',
                        ' Desal Production': 'Desal production', 'Desal Production': 'Desal production',
                        'LP Steam Consumption': 'LP Steam consumption',
                        'Condensate return': 'Condensate Return', 'Condensate Temp': 'condensate temp',
                        "1'st effect vapour Temp": '1st Effect Vapour Temp',
                        '1st Effect Brine Temp': '1st effect brine temp',
                        '1st Effect Vapour pres': '1st effect vapour pressure',
                        'Steam Inlet Temp': 'Steam Inlet Temp',
                        'Brine DischargeTemp': 'Brine Discharge Temp',
                        'Sea water cond (FFC) I/L temp': 'Sea Water cond I/L temp',
                        'Sea water cond (FFC) o/L temp': 'Sea Water Condenser O/L Temp',
                        'CW (FCC) supply': 'CW supply', 'CW (FCC) return': 'CW Return',
                        'Gross desal water production': 'Gross production',
                        '11 effect brine Temp': '11th Effect Brine Temp',
                        'Antiscalant residual (Cold group)': 'Anti_PPM',
                        'Antiscalant residual': 'Anti_PPM',
                        'Antiscalant residual (Hot group)': 'Anti_PPM_Hot',
                        'Antiscalant residual (Brine)': 'Anti_PPM_Brine',
                        'Unnamed: 27': 'Anti_PPM_Hot', 'Unnamed: 28': 'Anti_PPM_Brine',
                        'Remarks': 'Remarks', 'REMARKS': 'Remarks',
                    }, inplace=True)

                    op_inputs = [
                        'Sea Water Upper', 'Sea Water Lower', 'Sea Water Feed', 'Brine Water Return',
                        'Desal production', 'LP Steam consumption', 'Condensate Return', 'condensate temp',
                        '1st Effect Vapour Temp', '1st effect brine temp', '1st effect vapour pressure',
                        'Steam Inlet Temp', 'Brine Discharge Temp', 'Sea Water cond I/L temp',
                        'Sea Water Condenser O/L Temp', 'CW supply', 'CW Return', 'Gross production',
                        '11th Effect Brine Temp', 'Anti_PPM', 'Anti_PPM_Hot', 'Anti_PPM_Brine',
                    ]
                    d = _clean_num(d, op_inputs)
                    d['Date'] = standardize_dates(d['Date']).dt.strftime('%Y-%m-%d')
                    d = d.dropna(subset=['Date'])

                    if len(d) == 0:
                        st.error("No valid dated rows found.")
                    else:
                        steam = d['LP Steam consumption']
                        gross = d['Gross production']
                        desal = d['Desal production']
                        swfeed = d['Sea Water Feed']

                        # Derived - recomputed, never trusted from the file.
                        d['Delta T'] = d['1st Effect Vapour Temp'] - d['1st effect brine temp']
                        d['Overall Delta T'] = d['1st Effect Vapour Temp'] - d['11th Effect Brine Temp']
                        d['GOR'] = np.where(steam > 0, gross / steam, np.nan)
                        d['Recovery'] = np.where(swfeed > 0, (gross / swfeed) * 100, np.nan)
                        d['Conversion'] = d['Recovery'] / 100
                        d['Steam Economy'] = np.where(desal > 0, steam / desal, np.nan)
                        d['STEC'] = np.where(
                            desal > 0, ((steam * 1000) / 3600 * LATENT_HEAT_STEAM_KJ_KG) / desal, np.nan
                        )

                        out_cols = ['Date'] + op_inputs + [
                            'Delta T', 'Overall Delta T', 'GOR', 'Recovery', 'Conversion',
                            'Steam Economy', 'STEC'
                        ]
                        ready = d[out_cols].copy()
                        ready['Remarks'] = d.get('Remarks', pd.Series("", index=d.index)).fillna("")

                        st.success(f"Recomputed operational KPIs for {len(ready)} rows.")
                        st.dataframe(ready.style.format(precision=2), use_container_width=True, hide_index=True)

                        cp, cs = st.columns([2, 2])
                        pw = cp.text_input("Password", type="password", key="pw_op",
                                           label_visibility="collapsed", placeholder="Master password to sync")
                        if cs.button("Update Database (Operational)", use_container_width=True, key="b_op"):
                            if pw == "12345678":
                                st.session_state.daily_logs = upsert_daily_logs(st.session_state.daily_logs, ready)
                                save_database(db_conn, st.session_state.daily_logs, LOCAL_DB_FILE)
                                st.success("Operational data synced. HTC and Water Quality untouched.")
                                time.sleep(1.2); st.rerun()
                            elif pw != "":
                                st.error("Incorrect password.")
                except Exception as e:
                    st.error(f"Error processing file: {e}")

        # ===================================================================================
        # B) 1st EFFECT HTC  <-  '1st effect-HTC' sheet
        # ===================================================================================
        with bulk_subtabs[1]:
            st.markdown(
                "Source: **`1st effect-HTC`** tab. Upload only columns **A-K** (the process inputs). "
                "The calculator recomputes ΔT1, ΔT2, LMTD, Heat Duty, HTC and Fouling."
            )
            st.info(
                "**Column meanings on this sheet** (they differ from the Overall-HTC sheet):\n\n"
                "- **Feed flow** = feed to the 1st effect (~514 m³/hr), tag Z711FIT424\n"
                "- **Feed Temp** = *average brine temp of effects 4, 5, 6 and 7* (~49 °C) - this is the cold-side reference, not a seawater temp\n"
                "- **Brine Temp** = 1st effect brine temp (~66 °C), tag Z711TIT401\n"
                "- **1st effect vapor temp** = ~69 °C, tag Z711TIT414\n"
                "- **Condensate temperature** = ~75 °C, tag Z711TIT415\n"
                "- **Heat Transfer Area** = 12,950 m² (leave blank to use this default)\n\n"
                "ΔT1 = vapor − brine · ΔT2 = condensate − Feed Temp(eff 4-7)"
            )
            st.download_button(
                "Download 1st Effect HTC Template", key='dl_h1',
                data=pd.DataFrame(columns=HTC_1ST_BULK_HEADERS).to_csv(index=False).encode('utf-8'),
                file_name='MED4_1stEffect_HTC_Template.csv', mime='text/csv'
            )
            st.divider()
            h1_file = st.file_uploader("Upload 1st Effect HTC CSV", type=["csv"], key="h1_up")

            if h1_file is not None:
                try:
                    d = pd.read_csv(h1_file)
                    first = d.columns[0]
                    d = d[~d[first].astype(str).isin(['Unit ', 'Unit', 'Tag', 'Desigen', 'Design', 'SOR/  Base line'])]
                    d.rename(columns={
                        first: 'Date',
                        'Feed flow': 'HTC1_Feed_Flow', 'Product flow ': 'HTC1_Product_Flow',
                        'Product flow': 'HTC1_Product_Flow',
                        'Condensate Flow ': 'HTC1_Cond_Flow', 'Condensate Flow': 'HTC1_Cond_Flow',
                        'Steam consumption rate': 'HTC1_Steam_TPH',
                        'Feed Temp': 'HTC1_Feed_Temp_Eff4to7',
                        'Brine Temp': 'HTC1_Brine_Temp',
                        '1st effect vapor temp': 'HTC1_Vapor_Temp',
                        ' Condensate temperature': 'HTC1_Cond_Temp',
                        'Condensate temperature': 'HTC1_Cond_Temp',
                        'Heat Transfer Area ': 'HTC1_Area', 'Heat Transfer Area': 'HTC1_Area',
                    }, inplace=True)

                    h1_inputs = ['HTC1_Feed_Flow', 'HTC1_Product_Flow', 'HTC1_Cond_Flow', 'HTC1_Steam_TPH',
                                 'HTC1_Feed_Temp_Eff4to7', 'HTC1_Brine_Temp', 'HTC1_Vapor_Temp',
                                 'HTC1_Cond_Temp', 'HTC1_Area']
                    d = _clean_num(d, h1_inputs)
                    d['Date'] = standardize_dates(d['Date']).dt.strftime('%Y-%m-%d')
                    d = d.dropna(subset=['Date'])

                    # Pull anything this file didn't supply from the registry (Operational upload shares
                    # steam rate, vapour/brine/condensate temps, product and condensate flows).
                    d, filled = _backfill_from_db(d, {
                        'HTC1_Steam_TPH': 'LP Steam consumption',
                        'HTC1_Vapor_Temp': '1st Effect Vapour Temp',
                        'HTC1_Brine_Temp': '1st effect brine temp',
                        'HTC1_Cond_Temp': 'condensate temp',
                        'HTC1_Product_Flow': 'Desal production',
                        'HTC1_Cond_Flow': 'Condensate Return',
                        'HTC1_Feed_Temp_Eff4to7': 'HTC1_Feed_Temp_Eff4to7',
                        'HTC1_Feed_Flow': 'HTC1_Feed_Flow',
                    })
                    if filled:
                        st.info("Filled from existing registry data: " + ", ".join(filled))

                    if len(d) == 0:
                        st.error("No valid dated rows found.")
                    else:
                        d['HTC1_Area'] = d['HTC1_Area'].fillna(HTC_1ST_AREA)

                        # ΔT1 = 1st effect vapor temp - 1st effect brine temp   (sheet col L)
                        # ΔT2 = condensate temp - avg brine temp of effects 4-7 (sheet col M)
                        d['HTC1_dT1'] = d['HTC1_Vapor_Temp'] - d['HTC1_Brine_Temp']
                        d['HTC1_dT2'] = d['HTC1_Cond_Temp'] - d['HTC1_Feed_Temp_Eff4to7']
                        d['HTC1_LMTD'] = _lmtd(d['HTC1_dT1'], d['HTC1_dT2'])

                        # Steam-condensation heat duty (sheet cols V,W,X):
                        # ms(kg/hr) = TPH*1000 ; W(kJ/hr) = ms*lambda ; Q(W) = W*1000/3600
                        d['HTC1_Q_Steam'] = (d['HTC1_Steam_TPH'] * 1000 * LATENT_HEAT_STEAM_KJ_KG * 1000) / 3600

                        # U (steam condensation basis) = Q / (A * LMTD)   (sheet col AA)
                        denom = d['HTC1_Area'] * d['HTC1_LMTD']
                        d['1st Effect HTC'] = np.where(
                            d['HTC1_Q_Steam'].notna() & pd.notna(denom) & (denom > 0),
                            d['HTC1_Q_Steam'] / denom, np.nan
                        )
                        d['HTC1_Fouling'] = np.where(d['1st Effect HTC'] > 0, 1 / d['1st Effect HTC'], np.nan)
                        # Rf = 1/U_actual - 1/U_SOR_baseline   (sheet col AC)
                        d['HTC1_Rf'] = np.where(
                            d['1st Effect HTC'] > 0,
                            (1 / d['1st Effect HTC']) - (1 / HTC_1ST_U_SOR), np.nan
                        )
                        d['Area_1st'] = d['HTC1_Area']

                        ready = d[['Date'] + h1_inputs + [
                            'HTC1_dT1', 'HTC1_dT2', 'HTC1_LMTD', 'HTC1_Q_Steam',
                            '1st Effect HTC', 'HTC1_Fouling', 'HTC1_Rf', 'Area_1st'
                        ]].copy()

                        n_bad = int(ready['1st Effect HTC'].isna().sum())
                        st.success(f"Computed 1st Effect HTC for {len(ready) - n_bad} of {len(ready)} rows.")
                        if n_bad:
                            st.warning(f"{n_bad} row(s) left blank - missing one of: steam rate, vapor temp, "
                                       f"brine temp, condensate temp, or Feed Temp (eff 4-7).")
                        st.dataframe(ready.style.format(precision=2), use_container_width=True, hide_index=True)

                        cp, cs = st.columns([2, 2])
                        pw = cp.text_input("Password", type="password", key="pw_h1",
                                           label_visibility="collapsed", placeholder="Master password to sync")
                        if cs.button("Update Database (1st Effect HTC)", use_container_width=True, key="b_h1"):
                            if pw == "12345678":
                                st.session_state.daily_logs = upsert_daily_logs(st.session_state.daily_logs, ready)
                                save_database(db_conn, st.session_state.daily_logs, LOCAL_DB_FILE)
                                st.success("1st Effect HTC synced. Operational, Overall HTC and Water Quality untouched.")
                                time.sleep(1.2); st.rerun()
                            elif pw != "":
                                st.error("Incorrect password.")
                except Exception as e:
                    st.error(f"Error processing file: {e}")

        # ===================================================================================
        # C) OVERALL HTC  <-  'Overall-HTC ' sheet
        # ===================================================================================
        with bulk_subtabs[2]:
            st.markdown(
                "Source: **`Overall-HTC`** tab. Upload only columns **A-K** (the process inputs). "
                "The calculator recomputes ΔT1, ΔT2, LMTD, Heat Duty, HTC and Fouling."
            )
            st.info(
                "**Column meanings on this sheet** (they differ from the 1st-effect sheet):\n\n"
                "- **Feed flow** = *total* seawater feed (~2062 m³/hr), tag Z711FIT424\n"
                "- **Feed Temp** = *feed temp to the cold group* (~40 °C) - the cold-side reference\n"
                "- **Brine discharge Temp** = ~42 °C, tag Z711TIT401\n"
                "- **1st effect vapor temp** = ~69 °C, tag Z711TIT414\n"
                "- **Condensate temperature** = ~75 °C\n"
                "- **Heat Transfer Area** = 163,818 m² (11 × 12,950 × 1.15; leave blank to use this default)\n\n"
                "ΔT1 = vapor − brine discharge · ΔT2 = condensate − Feed Temp(cold group)"
            )
            st.download_button(
                "Download Overall HTC Template", key='dl_ho',
                data=pd.DataFrame(columns=HTC_OVERALL_BULK_HEADERS).to_csv(index=False).encode('utf-8'),
                file_name='MED4_Overall_HTC_Template.csv', mime='text/csv'
            )
            st.divider()
            ho_file = st.file_uploader("Upload Overall HTC CSV", type=["csv"], key="ho_up")

            if ho_file is not None:
                try:
                    d = pd.read_csv(ho_file)
                    first = d.columns[0]
                    d = d[~d[first].astype(str).isin(['Unit ', 'Unit', 'Tag', 'Desigen', 'Design', 'SOR/  Base line'])]
                    d.rename(columns={
                        first: 'Date',
                        'Feed flow': 'HTCO_Feed_Flow', 'Product flow ': 'HTCO_Product_Flow',
                        'Product flow': 'HTCO_Product_Flow',
                        'Condensate Flow ': 'HTCO_Cond_Flow', 'Condensate Flow': 'HTCO_Cond_Flow',
                        'Steam consumption rate': 'HTCO_Steam_TPH',
                        'Feed Temp': 'HTCO_Feed_Temp_ColdGrp',
                        'Brine discharge Temp': 'HTCO_Brine_Disch_Temp',
                        'Brine Discharge Temp': 'HTCO_Brine_Disch_Temp',
                        '1st effect vapor temp': 'HTCO_Vapor_Temp',
                        ' Condensate temperature': 'HTCO_Cond_Temp',
                        'Condensate temperature': 'HTCO_Cond_Temp',
                    }, inplace=True)
                    # Area column header on this sheet carries the formula in its name, so match by prefix.
                    for c in list(d.columns):
                        if str(c).strip().startswith('Heat Transfer Area'):
                            d.rename(columns={c: 'HTCO_Area'}, inplace=True)

                    ho_inputs = ['HTCO_Feed_Flow', 'HTCO_Product_Flow', 'HTCO_Cond_Flow', 'HTCO_Steam_TPH',
                                 'HTCO_Feed_Temp_ColdGrp', 'HTCO_Brine_Disch_Temp', 'HTCO_Vapor_Temp',
                                 'HTCO_Cond_Temp', 'HTCO_Area']
                    d = _clean_num(d, ho_inputs)
                    d['Date'] = standardize_dates(d['Date']).dt.strftime('%Y-%m-%d')
                    d = d.dropna(subset=['Date'])

                    # Pull anything this file didn't supply from the registry (Operational upload shares
                    # steam rate, vapour/condensate temps, brine discharge temp, seawater feed, flows).
                    d, filled = _backfill_from_db(d, {
                        'HTCO_Steam_TPH': 'LP Steam consumption',
                        'HTCO_Vapor_Temp': '1st Effect Vapour Temp',
                        'HTCO_Brine_Disch_Temp': 'Brine Discharge Temp',
                        'HTCO_Cond_Temp': 'condensate temp',
                        'HTCO_Feed_Flow': 'Sea Water Feed',
                        'HTCO_Product_Flow': 'Desal production',
                        'HTCO_Cond_Flow': 'Condensate Return',
                        'HTCO_Feed_Temp_ColdGrp': 'HTCO_Feed_Temp_ColdGrp',
                    })
                    if filled:
                        st.info("Filled from existing registry data: " + ", ".join(filled))

                    if len(d) == 0:
                        st.error("No valid dated rows found.")
                    else:
                        d['HTCO_Area'] = d['HTCO_Area'].fillna(HTC_OVERALL_AREA)

                        # ΔT1 = 1st effect vapor temp - brine discharge temp     (sheet col L)
                        # ΔT2 = condensate temp - feed temp to cold group        (sheet col M)
                        d['HTCO_dT1'] = d['HTCO_Vapor_Temp'] - d['HTCO_Brine_Disch_Temp']
                        d['HTCO_dT2'] = d['HTCO_Cond_Temp'] - d['HTCO_Feed_Temp_ColdGrp']
                        d['HTCO_LMTD'] = _lmtd(d['HTCO_dT1'], d['HTCO_dT2'])

                        d['HTCO_Q_Steam'] = (d['HTCO_Steam_TPH'] * 1000 * LATENT_HEAT_STEAM_KJ_KG * 1000) / 3600

                        denom = d['HTCO_Area'] * d['HTCO_LMTD']
                        d['Overall HTC'] = np.where(
                            d['HTCO_Q_Steam'].notna() & pd.notna(denom) & (denom > 0),
                            d['HTCO_Q_Steam'] / denom, np.nan
                        )
                        d['HTCO_Fouling'] = np.where(d['Overall HTC'] > 0, 1 / d['Overall HTC'], np.nan)
                        d['HTCO_Rf'] = np.where(
                            d['Overall HTC'] > 0,
                            (1 / d['Overall HTC']) - (1 / HTC_OVERALL_U_SOR), np.nan
                        )
                        d['Area_Overall'] = d['HTCO_Area']

                        ready = d[['Date'] + ho_inputs + [
                            'HTCO_dT1', 'HTCO_dT2', 'HTCO_LMTD', 'HTCO_Q_Steam',
                            'Overall HTC', 'HTCO_Fouling', 'HTCO_Rf', 'Area_Overall'
                        ]].copy()

                        n_bad = int(ready['Overall HTC'].isna().sum())
                        st.success(f"Computed Overall HTC for {len(ready) - n_bad} of {len(ready)} rows.")
                        if n_bad:
                            st.warning(f"{n_bad} row(s) left blank - missing one of: steam rate, vapor temp, "
                                       f"brine discharge temp, condensate temp, or Feed Temp (cold group).")
                        st.dataframe(ready.style.format(precision=2), use_container_width=True, hide_index=True)

                        cp, cs = st.columns([2, 2])
                        pw = cp.text_input("Password", type="password", key="pw_ho",
                                           label_visibility="collapsed", placeholder="Master password to sync")
                        if cs.button("Update Database (Overall HTC)", use_container_width=True, key="b_ho"):
                            if pw == "12345678":
                                st.session_state.daily_logs = upsert_daily_logs(st.session_state.daily_logs, ready)
                                save_database(db_conn, st.session_state.daily_logs, LOCAL_DB_FILE)
                                st.success("Overall HTC synced. Operational, 1st Effect HTC and Water Quality untouched.")
                                time.sleep(1.2); st.rerun()
                            elif pw != "":
                                st.error("Incorrect password.")
                except Exception as e:
                    st.error(f"Error processing file: {e}")

        # ===================================================================================
        # D) WATER QUALITY  <-  'Feed & Brine Water Analysis' + 'Desal Analysis' sheets
        # ===================================================================================
        with bulk_subtabs[3]:
            st.markdown(
                "Two separate lab sheets, uploaded independently. Both are pure lab readings - nothing is "
                "derived from them, so they are stored exactly as supplied (`-` becomes blank, not 0)."
            )

            st.markdown("##### Feed & Brine Water Analysis")
            st.download_button(
                "Download Feed & Brine Template", key='dl_fb',
                data=pd.DataFrame(columns=FEEDBRINE_BULK_HEADERS).to_csv(index=False).encode('utf-8'),
                file_name='MED4_FeedBrine_Template.csv', mime='text/csv'
            )
            fb_file = st.file_uploader("Upload Feed & Brine Analysis CSV", type=["csv"], key="fb_up")

            if fb_file is not None:
                try:
                    d = pd.read_csv(fb_file)
                    first = d.columns[0]
                    d = d[~d[first].astype(str).isin(['UOM', 'Specified Limit'])]
                    fb_map = {
                        first: 'Date',
                        'pH': 'Feed_pH', 'Turbidity': 'Feed_Turbidity', 'TSS': 'Feed_TSS',
                        'Conductivity': 'Feed_Cond', 'TDS': 'Feed_TDS',
                        'Total Alkalinity': 'Feed_Alkalinity', 'Calcium Hardness': 'Feed_Calcium',
                        'Mg Hardness': 'Feed_MgHardness', 'Total Hardness': 'Feed_TotalHardness',
                        'Silica': 'Feed_Silica', 'Chloride ': 'Feed_Chlorides', 'Chloride': 'Feed_Chlorides',
                        'Sulphate': 'Feed_Sulphate', 'Sulphide': 'Feed_Sulphide',
                        'Brine pH': 'Brine_pH', 'Brine Turbidity': 'Brine_Turbidity',
                        'Brine Conductivity': 'Brine_Cond', 'Brine TDS': 'Brine_TDS',
                        'Brine Total Alkalinity': 'Brine_Alkalinity',
                        'Brine Calcium Hardness': 'Brine_Calcium', 'Brine Mg Hardness': 'Brine_MgHardness',
                        'Brine Total Hardness': 'Brine_TotalHardness', 'Brine Silica': 'Brine_Silica',
                        'Brine Chloride': 'Brine_Chlorides',
                        # Raw sheet exports the brine block with duplicate names, which pandas suffixes '.1'
                        'pH.1': 'Brine_pH', 'Turbidity.1': 'Brine_Turbidity',
                        'Conductivity.1': 'Brine_Cond', 'TDS.1': 'Brine_TDS',
                        'Total Alkalinity.1': 'Brine_Alkalinity', 'Calcium Hardness.1': 'Brine_Calcium',
                        'Mg Hardness.1': 'Brine_MgHardness', 'Total Hardness.1': 'Brine_TotalHardness',
                        'Silica.1': 'Brine_Silica', 'Chloride .1': 'Brine_Chlorides',
                        'Chloride.1': 'Brine_Chlorides',
                        'REMARKS': 'Remarks',
                    }
                    d.rename(columns=fb_map, inplace=True)
                    fb_cols = [c for c in dict.fromkeys(fb_map.values()) if c not in ('Date', 'Remarks')]
                    d = _clean_num(d, fb_cols)
                    d['Date'] = standardize_dates(d['Date']).dt.strftime('%Y-%m-%d')
                    d = d.dropna(subset=['Date'])

                    if len(d) == 0:
                        st.error("No valid dated rows found.")
                    else:
                        ready = d[['Date'] + fb_cols].copy()
                        ready['Remarks'] = d.get('Remarks', pd.Series("", index=d.index)).fillna("")
                        st.success(f"Prepared Feed & Brine analysis for {len(ready)} rows.")
                        st.dataframe(ready.style.format(precision=2), use_container_width=True, hide_index=True)

                        cp, cs = st.columns([2, 2])
                        pw = cp.text_input("Password", type="password", key="pw_fb",
                                           label_visibility="collapsed", placeholder="Master password to sync")
                        if cs.button("Update Database (Feed & Brine)", use_container_width=True, key="b_fb"):
                            if pw == "12345678":
                                st.session_state.daily_logs = upsert_daily_logs(st.session_state.daily_logs, ready)
                                save_database(db_conn, st.session_state.daily_logs, LOCAL_DB_FILE)
                                st.success("Feed & Brine analysis synced.")
                                time.sleep(1.2); st.rerun()
                            elif pw != "":
                                st.error("Incorrect password.")
                except Exception as e:
                    st.error(f"Error processing Feed & Brine file: {e}")

            st.divider()
            st.markdown("##### Desal (Product) Analysis")
            st.download_button(
                "Download Desal Analysis Template", key='dl_ds',
                data=pd.DataFrame(columns=DESAL_BULK_HEADERS).to_csv(index=False).encode('utf-8'),
                file_name='MED4_Desal_Template.csv', mime='text/csv'
            )
            ds_file = st.file_uploader("Upload Desal Analysis CSV", type=["csv"], key="ds_up")

            if ds_file is not None:
                try:
                    d = pd.read_csv(ds_file)
                    first = d.columns[0]
                    d = d[~d[first].astype(str).isin(['UOM', 'Specified Limit'])]
                    ds_map = {
                        first: 'Date',
                        'pH': 'Product_pH', 'Turbidity': 'Product_Turbidity',
                        'Conductivity': 'Product_Cond', 'TDS': 'Product_TDS',
                        'Total Alkalinity': 'Product_Alkalinity', 'Calcium Hardness': 'Product_Calcium',
                        'Mg Hardness': 'Product_MgHardness', 'Total Hardness': 'Product_TotalHardness',
                        'Chloride ': 'Product_Chlorides', 'Chloride': 'Product_Chlorides',
                        'Total Iron ': 'Product_Iron', 'Total Iron': 'Product_Iron',
                        'Silica': 'Product_Silica', 'Sulphate': 'Product_Sulphate',
                        'REMARKS': 'Remarks',
                    }
                    d.rename(columns=ds_map, inplace=True)
                    ds_cols = [c for c in dict.fromkeys(ds_map.values()) if c not in ('Date', 'Remarks')]
                    d = _clean_num(d, ds_cols)
                    d['Date'] = standardize_dates(d['Date']).dt.strftime('%Y-%m-%d')
                    d = d.dropna(subset=['Date'])

                    if len(d) == 0:
                        st.error("No valid dated rows found.")
                    else:
                        ready = d[['Date'] + ds_cols].copy()
                        ready['Remarks'] = d.get('Remarks', pd.Series("", index=d.index)).fillna("")
                        st.success(f"Prepared Desal product analysis for {len(ready)} rows.")
                        st.dataframe(ready.style.format(precision=2), use_container_width=True, hide_index=True)

                        cp, cs = st.columns([2, 2])
                        pw = cp.text_input("Password", type="password", key="pw_ds",
                                           label_visibility="collapsed", placeholder="Master password to sync")
                        if cs.button("Update Database (Desal Analysis)", use_container_width=True, key="b_ds"):
                            if pw == "12345678":
                                st.session_state.daily_logs = upsert_daily_logs(st.session_state.daily_logs, ready)
                                save_database(db_conn, st.session_state.daily_logs, LOCAL_DB_FILE)
                                st.success("Desal product analysis synced.")
                                time.sleep(1.2); st.rerun()
                            elif pw != "":
                                st.error("Incorrect password.")
                except Exception as e:
                    st.error(f"Error processing Desal file: {e}")

        # ===================================================================================
        # E) CHEMICAL DOSES  <-  'Chemicals doses' sheet
        # ===================================================================================
        with bulk_subtabs[4]:
            st.markdown(
                "Source sheet: **Chemicals doses**. Upload the raw tank readings (initial / top-up / final "
                "levels and hours) plus MMC stock movements. The calculator recomputes **LPH**, **Kg/hr** and "
                "**PPM** for both antiscalant (Kem Watreat r 3687) and antifoam (Kem Antifoam 1795) exactly as "
                "the sheet does."
            )
            st.caption(
                "Level drop = Initial + Top-up − Final · LPH = (drop ÷ hrs) × 23 · "
                "AS Kg/hr = LPH × 1.20 · AF Kg/hr = LPH × 0.02 · PPM = Kg/hr × 1000 ÷ seawater feed. "
                "PPM uses the seawater feed already stored for that date from the Operational upload."
            )
            st.download_button(
                "Download Chemical Doses Template", key='dl_chem',
                data=pd.DataFrame(columns=CHEM_BULK_HEADERS).to_csv(index=False).encode('utf-8'),
                file_name='MED4_ChemicalDoses_Template.csv', mime='text/csv'
            )
            st.divider()
            chem_file = st.file_uploader("Upload Chemical Doses CSV", type=["csv"], key="chem_up")

            if chem_file is not None:
                try:
                    d = pd.read_csv(chem_file)
                    first = d.columns[0]
                    d = d[~d[first].astype(str).isin(['Unit', 'UOM', 'Initail', 'Top-up', 'Final'])]
                    d.rename(columns={
                        first: 'Date',
                        'AS Initial': 'AS_Initial', 'AS Top-up': 'AS_Topup', 'AS Final': 'AS_Final', 'AS Nos of Hrs': 'AS_Hours',
                        'AF Initial': 'AF_Initial', 'AF Top-up': 'AF_Topup', 'AF Final': 'AF_Final', 'AF Nos of Hrs': 'AF_Hours',
                        'AS Stock Opening': 'AS_Stock_Open', 'AS Stock Received': 'AS_Stock_Recd',
                        'AS Stock Consumed': 'AS_Stock_Consumed', 'AS Stock Closing': 'AS_Stock_Close',
                        'AF Stock Opening': 'AF_Stock_Open', 'AF Stock Received': 'AF_Stock_Recd',
                        'AF Stock Consumed': 'AF_Stock_Consumed', 'AF Stock Closing': 'AF_Stock_Close',
                    }, inplace=True)

                    chem_num = ['AS_Initial', 'AS_Topup', 'AS_Final', 'AS_Hours',
                                'AF_Initial', 'AF_Topup', 'AF_Final', 'AF_Hours',
                                'AS_Stock_Open', 'AS_Stock_Recd', 'AS_Stock_Consumed', 'AS_Stock_Close',
                                'AF_Stock_Open', 'AF_Stock_Recd', 'AF_Stock_Consumed', 'AF_Stock_Close']
                    d = _clean_num(d, chem_num)
                    d['Date'] = standardize_dates(d['Date']).dt.strftime('%Y-%m-%d')
                    d = d.dropna(subset=['Date'])

                    if len(d) == 0:
                        st.error("No valid dated rows found.")
                    else:
                        # Pull the seawater feed already stored for each date (from Operational upload)
                        # so PPM is computed against the real feed rather than a guess.
                        feed_map = {}
                        logs = st.session_state.daily_logs
                        if logs is not None and not logs.empty and 'Date' in logs.columns:
                            ref = logs.copy()
                            ref['Date'] = standardize_dates(ref['Date']).dt.strftime('%Y-%m-%d')
                            ref = ref.dropna(subset=['Date']).drop_duplicates(subset=['Date'], keep='last').set_index('Date')
                            if 'Sea Water Feed' in ref.columns:
                                feed_map = pd.to_numeric(ref['Sea Water Feed'], errors='coerce').to_dict()
                        feed = d['Date'].map(feed_map)

                        # --- Antiscalant derivations ---
                        d['AS_LevelDrop'] = d['AS_Initial'].fillna(0) + d['AS_Topup'].fillna(0) - d['AS_Final'].fillna(0)
                        d['AS_LPH'] = np.where(d['AS_Hours'] > 0, (d['AS_LevelDrop'] / d['AS_Hours']) * LITRES_PER_LEVEL_UNIT, np.nan)
                        d['AS_KgHr'] = d['AS_LPH'] * AS_DENSITY
                        d['AS_PPM'] = np.where((feed > 0) & d['AS_KgHr'].notna(), d['AS_KgHr'] * 1000 / feed, np.nan)

                        # --- Antifoam derivations ---
                        d['AF_LevelDrop'] = d['AF_Initial'].fillna(0) + d['AF_Topup'].fillna(0) - d['AF_Final'].fillna(0)
                        d['AF_LPH'] = np.where(d['AF_Hours'] > 0, (d['AF_LevelDrop'] / d['AF_Hours']) * LITRES_PER_LEVEL_UNIT, np.nan)
                        d['AF_KgHr'] = d['AF_LPH'] * AF_DENSITY
                        d['AF_PPM'] = np.where((feed > 0) & d['AF_KgHr'].notna(), d['AF_KgHr'] * 1000 / feed, np.nan)

                        # Also feed the legacy KPI-tab fields so the Chemicals tab and SOR chem section populate.
                        d['Anti_PPM'] = d['AS_PPM']
                        d['Foam_PPM'] = d['AF_PPM']
                        d['Antiscalant (kg)'] = d['AS_KgHr']
                        d['Antifoam (kg)'] = d['AF_KgHr']

                        keep = ['Date',
                                'AS_Initial', 'AS_Topup', 'AS_Final', 'AS_LevelDrop', 'AS_Hours', 'AS_LPH', 'AS_KgHr', 'AS_PPM',
                                'AF_Initial', 'AF_Topup', 'AF_Final', 'AF_LevelDrop', 'AF_Hours', 'AF_LPH', 'AF_KgHr', 'AF_PPM',
                                'AS_Stock_Open', 'AS_Stock_Recd', 'AS_Stock_Consumed', 'AS_Stock_Close',
                                'AF_Stock_Open', 'AF_Stock_Recd', 'AF_Stock_Consumed', 'AF_Stock_Close',
                                'Anti_PPM', 'Foam_PPM', 'Antiscalant (kg)', 'Antifoam (kg)']
                        ready = d[keep].copy()
                        ready['Remarks'] = d.get('Remarks', pd.Series("", index=d.index)).fillna("")

                        n_nofeed = int(ready['AS_PPM'].isna().sum())
                        st.success(f"Computed chemical dosing for {len(ready)} rows.")
                        if n_nofeed:
                            st.warning(f"PPM left blank for {n_nofeed} row(s) with no seawater feed stored yet - "
                                       f"upload the matching Operational data first, then re-upload this file to fill PPM.")
                        show_cols = ['Date', 'AS_LevelDrop', 'AS_LPH', 'AS_KgHr', 'AS_PPM',
                                     'AF_LevelDrop', 'AF_LPH', 'AF_KgHr', 'AF_PPM', 'AS_Stock_Close', 'AF_Stock_Close']
                        st.dataframe(ready[show_cols].style.format(precision=3), use_container_width=True, hide_index=True)

                        cp, cs = st.columns([2, 2])
                        pw = cp.text_input("Password", type="password", key="pw_chem",
                                           label_visibility="collapsed", placeholder="Master password to sync")
                        if cs.button("Update Database (Chemical Doses)", use_container_width=True, key="b_chem"):
                            if pw == "12345678":
                                st.session_state.daily_logs = upsert_daily_logs(st.session_state.daily_logs, ready)
                                save_database(db_conn, st.session_state.daily_logs, LOCAL_DB_FILE)
                                st.success("Chemical dosing synced. Operational, HTC and Water Quality untouched.")
                                time.sleep(1.2); st.rerun()
                            elif pw != "":
                                st.error("Incorrect password.")
                except Exception as e:
                    st.error(f"Error processing Chemical Doses file: {e}")

    render_chatbot()
