# requirements: pandas, numpy, python-docx, altair, gspread, oauth2client, scikit-learn, xgboost, joblib, Pillow
import streamlit as st
import datetime
import pandas as pd
import numpy as np
import os
import json
import time
import altair as alt
import joblib
import re
from calculator_tab import show_matrix_calculator
from projection_engine import UtilityProjectionEngine
from med_suite import render_med_suite

try:
    import gspread
    from oauth2client.service_account import ServiceAccountCredentials
    GSPREAD_INSTALLED = True
except ImportError:
    GSPREAD_INSTALLED = False

try:
    from sklearn.linear_model import LinearRegression
    from sklearn.ensemble import RandomForestRegressor
    from sklearn.metrics import r2_score
    SKLEARN_INSTALLED = True
except ImportError:
    SKLEARN_INSTALLED = False

try:
    import xgboost as xgb
    XGB_INSTALLED = True
except ImportError:
    XGB_INSTALLED = False

try:
    from PIL import Image
    PIL_INSTALLED = True
except ImportError:
    PIL_INSTALLED = False

st.set_page_config(page_title="Chembond Water Technologies Limited | Enterprise Hub", layout="wide")

# ==========================================
# 1. CLOUD "GHOST SHEET" & CONFIG ENGINE
# ==========================================
GOOGLE_SHEET_NAME = "MED4_Cloud_Database"
LOCAL_DB_FILE = "MED4_Master_Database.csv"
LOCAL_CONFIG_FILE = "mra_config.json"
AI_MODEL_FILE = "mra_ai_model.pkl"

RO_LOCAL_DB_FILE = "RO_Master_Database.csv"
RO_LOCAL_CONFIG_FILE = "ro_mra_config.json"
RO_AI_MODEL_FILE = "ro_mra_ai_model.pkl"

MRA_COEF_2014 = {
    "model_type": "OLS",
    "Intercept": -161.5638, "Press_1st": 0.6136, "Temp_1st": 3.6392, 
    "SW_Upper": 0.8111, "Brine_Temp_1st": -7.6638, "Brine_Flow": -0.2329, 
    "LP_Steam": 8.2539, "Steam_Temp": 2.1924, "Anti_PPM": -7.0301
}

RO_MRA_COEF_BASE = {
    "model_type": "OLS", "Intercept": 50.0, "Feed_Flow": 0.85, 
    "Feed_TDS": -0.01, "Coag_PPM": 2.5, "SMBS_PPM": 1.1
}

RO_MRA_BASELINE = {
    "Feed_Flow": 450.0, "Feed_TDS": 2000.0, "Coag_PPM": 2.0, "SMBS_PPM": 3.0
}

@st.cache_resource(ttl=600)
def init_db_connection():
    if not GSPREAD_INSTALLED: 
        return {"type": "local", "client": None}
    if "gcp_service_account" in st.secrets:
        try:
            scope = ['https://spreadsheets.google.com/feeds', 'https://www.googleapis.com/auth/drive']
            creds_dict = dict(st.secrets["gcp_service_account"])
            if "\\n" in creds_dict["private_key"]: 
                creds_dict["private_key"] = creds_dict["private_key"].replace("\\n", "\n")
            creds = ServiceAccountCredentials.from_json_keyfile_dict(creds_dict, scope)
            client = gspread.authorize(creds)
            return {"type": "cloud", "client": client.open(GOOGLE_SHEET_NAME).sheet1}
        except: pass
    return {"type": "local", "client": None}

def load_database(db, target_file=LOCAL_DB_FILE):
    if db["type"] == "cloud" and target_file == LOCAL_DB_FILE:
        try:
            records = db["client"].get_all_records()
            if records: return pd.DataFrame(records)
        except: pass
    if os.path.exists(target_file): return pd.read_csv(target_file)
    return pd.DataFrame()

def save_database(db, df, target_file=LOCAL_DB_FILE):
    if 'Date' in df.columns:
        # By this point 'Date' has already been normalized to unambiguous ISO
        # YYYY-MM-DD by the calling code. Parse it as ISO first so we never
        # re-mangle an already-clean date. Re-parsing a clean ISO string with
        # dayfirst=True was swapping the month/day components (e.g. 2026-07-09
        # -> 2026-09-07, and 2026-01-15 -> NaT), which silently dropped every
        # date with day-of-month > 12. Only fall back to a dayfirst guess for
        # any leftover string that isn't already ISO.
        parsed = pd.to_datetime(df['Date'], format='%Y-%m-%d', errors='coerce')
        parsed = parsed.fillna(pd.to_datetime(df['Date'], errors='coerce', dayfirst=True))
        df['Date'] = parsed.dt.strftime('%Y-%m-%d')
    df = df.fillna(0)
    if db["type"] == "cloud" and target_file == LOCAL_DB_FILE:
        try:
            db["client"].clear()
            db["client"].update([df.columns.values.tolist()] + df.values.tolist())
            df.to_csv(target_file, index=False)
            return True
        except: pass
    df.to_csv(target_file, index=False)
    return True

def load_config(db, target_file=LOCAL_CONFIG_FILE, baseline_dict=MRA_COEF_2014):
    if os.path.exists(target_file):
        try:
            with open(target_file, "r") as f: 
                return json.load(f)
        except: pass
    return baseline_dict.copy()

def save_config(db, coef_dict, target_file=LOCAL_CONFIG_FILE):
    with open(target_file, "w") as f: json.dump(coef_dict, f)

db_conn = init_db_connection()

RO_EXACT_DB_COLUMNS = [
    "Date", "Feed Flow", "Permeate Flow", "Feed TDS", "Permeate TDS", 
    "Clarifier TSS", "PDMF TSS", "SDMF TSS", "Softener Hardness", "HRU Hardness",
    "Cartridge SDI", "Permeate pH", "Permeate COD", "Coagulant PPM", "Flocculant PPM", "SMBS PPM",
    "Recovery", "Rejection", "Residual", "Remarks"
]

# RO REPORT & CSV EXPORT GENERATORS
from io import BytesIO
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_ro_daily_csv(date, state, mra):
    data_dict = {
        "Date": date.strftime('%d/%m/%Y'),
        "Feed Flow": state.ro_feed_flow, "Permeate Flow": state.ro_perm_flow,
        "Feed TDS": state.ro_feed_tds, "Permeate TDS": state.ro_perm_tds,
        "Clarifier TSS": state.ro_clarifier_tss, "PDMF TSS": state.ro_pdmf_tss,
        "SDMF TSS": state.ro_sdmf_tss, "Softener Hardness": state.ro_soft_hard,
        "HRU Hardness": state.ro_hru_hard, "Cartridge SDI": state.ro_sdi,
        "Permeate pH": state.ro_perm_ph, "Permeate COD": state.ro_perm_cod,
        "Coagulant PPM": state.ro_coag_ppm, "Flocculant PPM": state.ro_floc_ppm,
        "SMBS PPM": state.ro_smbs_ppm, 
        "Recovery": round((state.ro_perm_flow / state.ro_feed_flow * 100) if state.ro_feed_flow > 0 else 0, 2),
        "Rejection": round(((state.ro_feed_tds - state.ro_perm_tds) / state.ro_feed_tds * 100) if state.ro_feed_tds > 0 else 0, 2),
        "Residual": round(mra['Residual'], 2), "Remarks": state.ro_remarks
    }
    df = pd.DataFrame([data_dict])
    return df.to_csv(index=False).encode('utf-8')

def generate_ro_comprehensive_report(date, state, mra):
    doc = Document()
    doc.add_heading('HERO RO Plant Daily Operational Report', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.add_run('Prepared by: ').bold = True
    p.add_run('Chembond Water Technologies Limited\n')
    p.add_run('Date: ').bold = True
    p.add_run(date.strftime('%d/%m/%Y'))
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_heading('1. Executive Summary', level=1)
    recovery = (state.ro_perm_flow / state.ro_feed_flow * 100) if state.ro_feed_flow > 0 else 0
    rejection = ((state.ro_feed_tds - state.ro_perm_tds) / state.ro_feed_tds * 100) if state.ro_feed_tds > 0 else 0
    doc.add_paragraph(f"On {date.strftime('%d/%m/%Y')}, the HERO RO Plant processed {state.ro_feed_flow} m³/h of feed, achieving a Permeate Flow of {state.ro_perm_flow} m³/h. The overall plant recovery was {recovery:.1f}% with a salt rejection rate of {rejection:.1f}%.")

    doc.add_heading('2. Quality Guarantees', level=1)
    t_wq = doc.add_table(rows=1, cols=4); t_wq.style = 'Table Grid'
    for i, h in enumerate(['Parameter', 'Limit/Spec', 'Actual', 'Status']): t_wq.rows[0].cells[i].text = h
    
    guarantees = [
        ("Clarifier Outlet TSS", "< 10 ppm", state.ro_clarifier_tss, "Pass" if state.ro_clarifier_tss < 10.0 else "Fail"),
        ("PDMF Outlet TSS", "< 3 ppm", state.ro_pdmf_tss, "Pass" if state.ro_pdmf_tss < 3.0 else "Fail"),
        ("SDMF Outlet TSS", "< 1 ppm", state.ro_sdmf_tss, "Pass" if state.ro_sdmf_tss < 1.0 else "Fail"),
        ("Softener O/L Hardness", "< 5 ppm", state.ro_soft_hard, "Pass" if state.ro_soft_hard < 5.0 else "Fail"),
        ("HRU O/L Hardness", "< 1 ppm", state.ro_hru_hard, "Pass" if state.ro_hru_hard < 1.0 else "Fail"),
        ("Cartridge Filter SDI", "< 3", state.ro_sdi, "Pass" if state.ro_sdi < 3.0 else "Fail"),
        ("RO Permeate COD", "< 10 ppm", state.ro_perm_cod, "Pass" if state.ro_perm_cod < 10.0 else "Fail"),
        ("RO Permeate pH", "7.0 - 7.5", state.ro_perm_ph, "Pass" if 7.0 <= state.ro_perm_ph <= 7.5 else "Fail")
    ]
    for param, limit, val, status in guarantees:
        rc = t_wq.add_row().cells
        rc[0].text, rc[1].text, rc[2].text, rc[3].text = str(param), str(limit), str(val), str(status)

    doc.add_heading('3. MRA Fouling Indicator (Normalized Flow)', level=1)
    diff_pct = (mra['Residual'] / mra['Predicted']) * 100 if mra['Predicted'] > 0 else 0
    doc.add_paragraph(f"Actual Permeate: {mra['Actual']:.1f} m³/h | Normalized Predicted Permeate: {mra['Predicted']:.1f} m³/h | Difference: {diff_pct:.1f}%")
    if diff_pct <= -5.0: doc.add_paragraph(f"STATUS: MEMBRANE FOULING DETECTED ({diff_pct:.1f}% flow loss). Initiate CIP protocol.").runs[0].font.color.rgb = RGBColor(255, 0, 0)
    elif diff_pct <= -4.0: doc.add_paragraph(f"STATUS: WARNING ({diff_pct:.1f}% flow loss). Verify antiscalant and SMBS dosing.").runs[0].font.color.rgb = RGBColor(255, 140, 0)
    else: doc.add_paragraph(f"STATUS: CLEAN ({diff_pct:.1f}% flow loss). Membranes operating normally.").runs[0].font.color.rgb = RGBColor(0, 128, 0)
    
    if state.ro_remarks and str(state.ro_remarks).strip() != "":
        doc.add_heading('4. Remarks & Observations', level=1)
        doc.add_paragraph(str(state.ro_remarks))

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def generate_ro_monthly_report(df_month, month_str, year_str):
    doc = Document()
    doc.add_heading(f'HERO RO Monthly Performance Summary: {month_str} {year_str}', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading('1. Monthly Aggregation', level=1)
    t_agg = doc.add_table(rows=1, cols=4); t_agg.style = 'Table Grid'
    for i, h in enumerate(['Metric', 'Minimum', 'Maximum', 'Average']): t_agg.rows[0].cells[i].text = h
    metrics = [("Permeate Flow (m³/h)", df_month['Permeate Flow']), ("Plant Recovery (%)", df_month['Recovery']), ("Salt Rejection (%)", df_month['Rejection']), ("Permeate TDS (ppm)", df_month['Permeate TDS'])]
    for name, series in metrics:
        rc = t_agg.add_row().cells
        rc[0].text, rc[1].text, rc[2].text, rc[3].text = name, f"{pd.to_numeric(series, errors='coerce').min():.2f}", f"{pd.to_numeric(series, errors='coerce').max():.2f}", f"{pd.to_numeric(series, errors='coerce').mean():.2f}"
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

# GLOBAL DATA LOADERS (Shared across apps)
if 'daily_logs' not in st.session_state: st.session_state.daily_logs = load_database(db_conn, LOCAL_DB_FILE)
if 'ro_daily_logs' not in st.session_state: st.session_state.ro_daily_logs = load_database(db_conn, RO_LOCAL_DB_FILE)

if 'mra_coef' not in st.session_state: st.session_state.mra_coef = load_config(db_conn, LOCAL_CONFIG_FILE, MRA_COEF_2014)
if 'ro_mra_coef' not in st.session_state: st.session_state.ro_mra_coef = load_config(db_conn, RO_LOCAL_CONFIG_FILE, RO_MRA_COEF_BASE)

if "messages" not in st.session_state:
    st.session_state.messages = [{"role": "assistant", "content": "Hello! I am the Chembond Water Assistant. Ask me anything about how the calculations work."}]

def render_chatbot():
    st.sidebar.divider()
    st.sidebar.markdown("### Chembond Water Assistant")
    
    chat_container = st.sidebar.container(height=350)
    for message in st.session_state.messages:
        chat_container.chat_message(message["role"]).markdown(message["content"])

    if prompt := st.sidebar.chat_input("Ask a question about formulas..."):
        st.session_state.messages.append({"role": "user", "content": prompt})
        chat_container.chat_message("user").markdown(prompt)

        p_lower = prompt.lower()
        if "password" in p_lower: response = "For security reasons, I cannot provide the Master Password. Please contact the plant administrator."
        elif "auto" in p_lower or "dose" in p_lower or "optimal" in p_lower or "calculate" in p_lower and "auto" in p_lower:
            response = "The **Auto-Calculate Optimal Dose** feature is currently in development! In a future update, it will use real-time feed water chemistry, Concentration Factors, and scaling indices to scientifically recommend the exact PPM needed."
        elif re.search(r'\bgor\b', p_lower) or "gain output ratio" in p_lower:
            response = "**GOR (Gain Output Ratio)** is calculated as:\n`Gross Production (m³/h) / LP Steam (TPH)`\n\nIt represents the 'fuel economy' of the plant—how many tons of water are produced per ton of steam."
        elif "recovery" in p_lower: response = "**System Recovery** is calculated as:\n`(Gross Production / Total SW Feed) * 100` for MED, or `(Permeate Flow / Feed Flow) * 100` for RO."
        elif re.search(r'\blmtd\b', p_lower) or "log mean" in p_lower:
            response = ("**LMTD (Log Mean Temperature Difference)** now uses the full log-mean formula, matching "
                        "the plant's HTC sheets:\n\n`LMTD = (ΔT1 − ΔT2) / ln(ΔT1 / ΔT2)`\n\n"
                        "The two driving forces ΔT1 and ΔT2 differ between the two HTC calculations — ask about "
                        "**1st Effect HTC** or **Overall HTC** for the specifics.")
        elif "1st effect htc" in p_lower:
            response = ("**1st Effect HTC (U)** uses the steam-condensation basis, mirroring the `1st effect-HTC` sheet:\n\n"
                        "`U = Q / (A × LMTD)`\n\nWhere:\n"
                        "* `Q` = `Steam(TPH) × 1000 × 2330 × 1000 / 3600` (W)\n"
                        "* `A` = **12,950 m²** (1st effect tube bundle)\n"
                        "* `ΔT1` = `1st Effect Vapour Temp − 1st Effect Brine Temp`\n"
                        "* `ΔT2` = `Condensate Temp − Avg Brine Temp of Effects 4-5-6-7`\n"
                        "* `LMTD` = `(ΔT1 − ΔT2) / ln(ΔT1/ΔT2)`\n\n"
                        "Typical value ~330 W/m²K. Note ΔT2's reference is the effects 4-7 average, not a seawater temp.")
        elif "overall htc" in p_lower or "htc" in p_lower:
            response = ("**Overall HTC (U)** uses the steam-condensation basis, mirroring the `Overall-HTC` sheet:\n\n"
                        "`U = Q / (A × LMTD)`\n\nWhere:\n"
                        "* `Q` = `Steam(TPH) × 1000 × 2330 × 1000 / 3600` (W)\n"
                        "* `A` = **163,818 m²** (11 effects × 12,950 × 1.15)\n"
                        "* `ΔT1` = `1st Effect Vapour Temp − Brine Discharge Temp`\n"
                        "* `ΔT2` = `Condensate Temp − Feed Temp to Cold Group`\n"
                        "* `LMTD` = `(ΔT1 − ΔT2) / ln(ΔT1/ΔT2)`\n\n"
                        "Typical value ~9 W/m²K. The 'Feed Temp' here means the cold-group feed (~40 °C), which is a "
                        "different measurement from the same-named column on the 1st-effect sheet.")
        elif "fouling factor" in p_lower or re.search(r'\brf\b', p_lower):
            response = ("**Fouling** is tracked two ways:\n\n"
                        "* **Fouling Factor** = `1 / HTC`\n"
                        "* **Fouling Resistance (Rf)** = `1/U_actual − 1/U_SOR_baseline`\n\n"
                        "Rf rising over time means scale is building up on the tubes. The HTC tab charts both HTCs "
                        "against their SOR baselines so you can watch the trend.")
        elif "stec" in p_lower:
            response = ("**STEC (Specific Thermal Energy Consumption)** is:\n"
                        "`STEC = (Steam(TPH) × 1000 / 3600 × 2330) / Desal Production`\n\n"
                        "It's the thermal energy in kWh consumed per tonne of distillate produced.")
        elif "chemical" in p_lower or "antiscalant" in p_lower or "antifoam" in p_lower or "dosing" in p_lower or re.search(r'\bppm\b', p_lower):
            response = ("**Chemical dosing** is derived from the daily tank-level drop, matching the Chemicals doses sheet:\n\n"
                        "* `Level Drop` = `Initial + Top-up − Final`\n"
                        "* `LPH` = `(Level Drop ÷ Hours) × 23`\n"
                        "* `Antiscalant Kg/hr` = `LPH × 1.20` · `Antifoam Kg/hr` = `LPH × 0.02`\n"
                        "* `PPM` = `Kg/hr × 1000 ÷ Seawater Feed`\n\n"
                        "Upload the **Chemical Doses** sheet on the Bulk Uploads tab; the Chemicals tab then shows dose "
                        "rates, residual PPM, MMC stock, and trends for both chemicals.")
        elif "bulk" in p_lower or "upload" in p_lower:
            response = ("The **Bulk Uploads** tab has five separate uploaders, one per source sheet: **Operational Data**, "
                        "**1st Effect HTC**, **Overall HTC**, **Water Quality**, and **Chemical Doses**. Each recomputes "
                        "its own KPIs and merges by date, so uploading one never overwrites another. Missing readings are "
                        "left genuinely blank (shown as 0) rather than filled with a guessed value.")
        elif re.search(r'\bols\b', p_lower) or "linear regression" in p_lower:
            response = "**OLS (Ordinary Least Squares)** is the standard mathematical method used to draw a straight line of best fit through data points. It creates the 'Digital Twin' of the plant's clean physics."
        elif "xgboost" in p_lower or "random forest" in p_lower or "ai" in p_lower:
            response = "**Random Forest and XGBoost** are advanced AI models that use Decision Trees instead of linear math. They are highly accurate at tracking complex plant behavior, but they don't give you simple linear 'coefficients' like OLS does. Your selected model is saved and will persist across reboots."
        elif "residual" in p_lower:
            response = "**Residual** is calculated as:\n`Actual Production - Predicted Production`\n\nA negative residual means the plant is underperforming compared to its clean digital twin, indicating scale/fouling is blocking performance."
        elif "fouling" in p_lower or "alert" in p_lower or "status" in p_lower:
            response = "The software calculates a **% Difference**:\n`(Residual / Predicted) * 100`\n\n* **Better than -4%:** CLEAN\n* **-4% to -5%:** WARNING\n* **Worse than -5%:** FOULING DETECTED"
        elif "remarks" in p_lower or "observation" in p_lower:
            response = "You can add custom notes, TT errors, or shift observations in the **Remarks & Observations** box in the Reporting Tab. These automatically save to the database and print on the Daily Word Report!"
        else:
            response = ("I am the Chembond Water Assistant. I can explain the current formulas for **1st Effect HTC, "
                        "Overall HTC, LMTD, Fouling (Rf), GOR, STEC, Recovery, Chemical Dosing (antiscalant & antifoam), "
                        "Residuals, Fouling alerts, OLS**, and **AI Models**. What would you like to know?")

        st.session_state.messages.append({"role": "assistant", "content": response})
        chat_container.chat_message("assistant").markdown(response)

# ==========================================
# MAIN APPLICATION HUBS AND ROUTER
# ==========================================
def main():
    try: st.sidebar.image("chembond_logo.png", use_container_width=True)
    except: st.sidebar.markdown("### CHEMBOND WATER TECHNOLOGIES LIMITED") 
        
    st.sidebar.divider()
    st.sidebar.markdown("### Utility Network")

    nav_options = ["Central Hub", "RO Plant", "Multi-Effect Distillation (MED)", "Projection Engine"]
    if 'utility_choice' not in st.session_state:
        st.session_state.utility_choice = "Central Hub"

    # Bind the selectbox to utility_choice via its key so the two share ONE piece of state.
    # Buttons elsewhere set st.session_state.utility_choice and rerun; because the selectbox reads
    # its value from that same key, there's no separate widget value to override the button's choice.
    st.sidebar.selectbox(
        "Select System", nav_options,
        key="utility_choice"
    )

    utility_choice = st.session_state.utility_choice

    if utility_choice == "Central Hub":
        st.title("Site Utility Management Suite")
        st.markdown("Central monitoring platform for Reliance utility systems. Select a system to open its full interface.")
        st.markdown("---")

        # --- MED summary from the registry ---
        med_status, med_kpis = "No Data", {}
        try:
            if not st.session_state.daily_logs.empty:
                logs = st.session_state.daily_logs.copy()
                # Registry dates are stored ISO (YYYY-MM-DD), so parse ISO first. Using dayfirst=True on
                # an ISO string drops any day-of-month > 12 as NaT (which made 'latest' wrongly settle on
                # the 12th) and swaps day/month for day <= 12. Only fall back to a dayfirst guess for any
                # value that isn't already ISO.
                _d = pd.to_datetime(logs['Date'], format='%Y-%m-%d', errors='coerce')
                _d = _d.fillna(pd.to_datetime(logs['Date'], errors='coerce', dayfirst=True))
                logs['_d'] = _d
                logs = logs.dropna(subset=['_d']).sort_values('_d')
                last = logs.iloc[-1]

                def _mnum(col, default=0.0):
                    try:
                        return float(str(last.get(col, default)).replace(',', '').strip() or default)
                    except (ValueError, TypeError):
                        return default

                med_kpis = {
                    "date": last['_d'].strftime('%d-%m-%Y'),
                    "gross": _mnum('Gross production'), "gor": _mnum('GOR'),
                    "htc_overall": _mnum('Overall HTC'), "recovery": _mnum('Recovery'),
                }
                gross_val = _mnum('Gross production', 1) or 1.0
                med_diff = (_mnum('Residual') / gross_val) * 100
                med_status = ("Attention Required" if med_diff <= -5.0
                              else "Minor Deviation" if med_diff <= -4.0
                              else "Good Working Condition")
        except Exception:
            med_status = "Data Format Error"

        # --- RO summary from the registry ---
        ro_status, ro_kpis = "No Data", {}
        try:
            if not st.session_state.ro_daily_logs.empty:
                rlogs = st.session_state.ro_daily_logs.copy()
                rlast = rlogs.iloc[-1]

                def _rnum(col, default=0.0):
                    try:
                        return float(str(rlast.get(col, default)).replace(',', '').strip() or default)
                    except (ValueError, TypeError):
                        return default

                ro_kpis = {
                    "perm": _rnum('Permeate Flow'), "recovery": _rnum('Recovery'),
                    "rejection": _rnum('Rejection'),
                }
                flow_val = _rnum('Permeate Flow', 1) or 1.0
                ro_diff = (_rnum('Residual') / flow_val) * 100
                ro_status = ("Attention Required" if ro_diff <= -5.0
                             else "Minor Deviation" if ro_diff <= -4.0
                             else "Good Working Condition")
        except Exception:
            ro_status = "Data Format Error"

        # Map a status to a color for the tile's status band.
        def _status_color(s):
            if s == "Good Working Condition": return "#1F9D55"   # green
            if s == "Minor Deviation": return "#D9822B"          # amber
            if s == "Attention Required": return "#D64545"       # red
            return "#7A8A99"                                      # grey (no data / error)

        tile_css = """
        <div style="border:1px solid #D6E6F5; border-radius:12px; padding:0; overflow:hidden;
                    background:#FFFFFF; box-shadow:0 1px 4px rgba(0,0,0,0.06); height:100%;">
          <div style="background:{band}; color:#FFFFFF; padding:10px 18px; font-weight:600; font-size:15px;">
            {title}
          </div>
          <div style="padding:16px 18px;">
            <div style="font-size:13px; color:#5A6B7B; margin-bottom:4px;">{subtitle}</div>
            <div style="font-size:15px; font-weight:600; color:{band}; margin-bottom:14px;">{status}</div>
            {body}
          </div>
        </div>
        """

        def _kpi_row(pairs):
            cells = "".join(
                f'<div style="flex:1; min-width:90px;">'
                f'<div style="font-size:20px; font-weight:700; color:#1B2A38;">{val}</div>'
                f'<div style="font-size:11px; color:#7A8A99; text-transform:uppercase; letter-spacing:0.4px;">{lbl}</div>'
                f'</div>'
                for lbl, val in pairs
            )
            return f'<div style="display:flex; gap:14px; flex-wrap:wrap;">{cells}</div>'

        col_med, col_ro = st.columns(2, gap="large")

        with col_med:
            if med_kpis:
                body = _kpi_row([
                    ("Gross Prod.", f"{med_kpis['gross']:.0f} m³/h"),
                    ("GOR", f"{med_kpis['gor']:.2f}"),
                    ("Overall HTC", f"{med_kpis['htc_overall']:.1f}"),
                    ("Recovery", f"{med_kpis['recovery']:.1f}%"),
                ])
                subtitle = f"Unit MED-4  ·  as of {med_kpis['date']}"
            else:
                body = '<div style="font-size:13px; color:#7A8A99;">No data loaded yet. Open the suite to upload plant data.</div>'
                subtitle = "Unit MED-4"
            st.markdown(tile_css.format(
                band=_status_color(med_status), title="Multi-Effect Distillation (MED)",
                subtitle=subtitle, status=med_status, body=body), unsafe_allow_html=True)
            st.write("")
            if st.button("Open MED-4 Suite", use_container_width=True, key="open_med", type="primary"):
                st.session_state.utility_choice = "Multi-Effect Distillation (MED)"
                st.rerun()

        with col_ro:
            if ro_kpis:
                body = _kpi_row([
                    ("Permeate", f"{ro_kpis['perm']:.0f} m³/h"),
                    ("Recovery", f"{ro_kpis['recovery']:.1f}%"),
                    ("Rejection", f"{ro_kpis['rejection']:.1f}%"),
                ])
                subtitle = "HERO Plant"
            else:
                body = '<div style="font-size:13px; color:#7A8A99;">No data loaded yet. Open the suite to upload plant data.</div>'
                subtitle = "HERO Plant"
            st.markdown(tile_css.format(
                band=_status_color(ro_status), title="Reverse Osmosis (RO)",
                subtitle=subtitle, status=ro_status, body=body), unsafe_allow_html=True)
            st.write("")
            if st.button("Open RO Suite", use_container_width=True, key="open_ro", type="primary"):
                st.session_state.utility_choice = "RO Plant"
                st.rerun()

        render_chatbot()
        return

    elif utility_choice == "Projection Engine":
        engine = UtilityProjectionEngine()
        engine.render_engine()
        return 

    elif utility_choice == "Multi-Effect Distillation (MED)":
        # Pass necessary backend hooks into the isolated MED suite file
        render_med_suite(
            db_conn=db_conn, 
            LOCAL_DB_FILE=LOCAL_DB_FILE, 
            LOCAL_CONFIG_FILE=LOCAL_CONFIG_FILE, 
            AI_MODEL_FILE=AI_MODEL_FILE, 
            save_database=save_database, 
            save_config=save_config, 
            render_chatbot=render_chatbot,
            SKLEARN_INSTALLED=SKLEARN_INSTALLED,
            XGB_INSTALLED=XGB_INSTALLED,
            PIL_INSTALLED=PIL_INSTALLED
        )
        return

    # ------------------------------------------
    # RO PLANT ENGINE (HERO)
    # ------------------------------------------
    elif utility_choice == "RO Plant":
        st.title("Reverse Osmosis (HERO) Suite")
        st.markdown("Monitor high-recovery RO metrics, pretreatment guarantees, and antiscalant/coagulant dosing for the SEZ RO facility.")
        
        log_date = st.sidebar.date_input("Date", datetime.date.today(), format="DD/MM/YYYY")
        log_date_str = log_date.strftime('%Y-%m-%d')
        
        # Setup session states for RO parameters safely
        ro_vars = {
            'ro_feed_flow': 450.0, 'ro_perm_flow': 385.0, 'ro_feed_tds': 2000.0, 'ro_perm_tds': 90.0,
            'ro_clarifier_tss': 8.0, 'ro_pdmf_tss': 2.0, 'ro_sdmf_tss': 0.5, 'ro_soft_hard': 4.0,
            'ro_hru_hard': 0.5, 'ro_sdi': 2.5, 'ro_perm_ph': 7.2, 'ro_perm_cod': 8.0,
            'ro_coag_ppm': 2.0, 'ro_floc_ppm': 1.0, 'ro_smbs_ppm': 3.0, 'ro_remarks': ""
        }
        for k, v in ro_vars.items():
            if k not in st.session_state: st.session_state[k] = v

        # RO Database Autoloader
        if 'ro_last_selected_date' not in st.session_state: st.session_state.ro_last_selected_date = None

        if log_date_str != st.session_state.ro_last_selected_date:
            st.session_state.ro_last_selected_date = log_date_str
            if not st.session_state.ro_daily_logs.empty and 'Date' in st.session_state.ro_daily_logs.columns:
                db_dates = pd.to_datetime(st.session_state.ro_daily_logs['Date'], errors='coerce').dt.strftime('%Y-%m-%d').values
                if log_date_str in db_dates:
                    row_idx = np.where(db_dates == log_date_str)[0][0]
                    row = st.session_state.ro_daily_logs.iloc[row_idx]
                    
                    db_to_var_mapping = {
                        'ro_feed_flow': 'Feed Flow', 'ro_perm_flow': 'Permeate Flow',
                        'ro_feed_tds': 'Feed TDS', 'ro_perm_tds': 'Permeate TDS',
                        'ro_clarifier_tss': 'Clarifier TSS', 'ro_pdmf_tss': 'PDMF TSS',
                        'ro_sdmf_tss': 'SDMF TSS', 'ro_soft_hard': 'Softener Hardness',
                        'ro_hru_hard': 'HRU Hardness', 'ro_sdi': 'Cartridge SDI',
                        'ro_perm_ph': 'Permeate pH', 'ro_perm_cod': 'Permeate COD',
                        'ro_coag_ppm': 'Coagulant PPM', 'ro_floc_ppm': 'Flocculant PPM',
                        'ro_smbs_ppm': 'SMBS PPM', 'ro_remarks': 'Remarks'
                    }
                    
                    loaded_vars = False
                    for var_key, col_name in db_to_var_mapping.items():
                        if col_name in row.index and pd.notna(row[col_name]):
                            try:
                                val_str = str(row[col_name]).strip()
                                if val_str and val_str.lower() not in ['nan', 'none', 'null', 'na']:
                                    if var_key == 'ro_remarks': st.session_state[var_key] = val_str
                                    else: st.session_state[var_key] = float(val_str.replace(',', ''))
                                    loaded_vars = True
                            except: pass 
                    if loaded_vars: st.sidebar.success(f"Loaded RO data for {log_date.strftime('%d-%m-%Y')}")

        # --- RO MRA CALCULATION ---
        ro_mra_data = {}
        ro_coefs = st.session_state.ro_mra_coef 
        ro_model_type = ro_coefs.get("model_type", "OLS")
        
        ro_live_input = [st.session_state.ro_feed_flow, st.session_state.ro_feed_tds, st.session_state.ro_coag_ppm, st.session_state.ro_smbs_ppm]
        
        if ro_model_type == "OLS":
            ro_mra_data['Predicted'] = (
                ro_coefs["Intercept"] + 
                (ro_coefs["Feed_Flow"] * ro_live_input[0]) + 
                (ro_coefs["Feed_TDS"] * ro_live_input[1]) + 
                (ro_coefs["Coag_PPM"] * ro_live_input[2]) + 
                (ro_coefs["SMBS_PPM"] * ro_live_input[3])
            )
        else:
            try:
                active_model = joblib.load(RO_AI_MODEL_FILE)
                live_df = pd.DataFrame([ro_live_input], columns=["Feed_Flow", "Feed_TDS", "Coag_PPM", "SMBS_PPM"])
                ro_mra_data['Predicted'] = float(active_model.predict(live_df)[0])
            except: ro_mra_data['Predicted'] = 0.0
                
        ro_mra_data['Actual'] = st.session_state.ro_perm_flow
        ro_mra_data['Residual'] = ro_mra_data['Actual'] - ro_mra_data['Predicted']

        ro_var_data = []
        ro_param_keys = ["Feed_Flow", "Feed_TDS", "Coag_PPM", "SMBS_PPM"]
        ro_param_names = ["Feed Flow", "Feed TDS", "Coagulant PPM", "SMBS PPM"]
        
        for i in range(4):
            dev = ro_live_input[i] - RO_MRA_BASELINE[ro_param_keys[i]]
            weight = ro_coefs.get(ro_param_keys[i], 0.0) 
            if ro_model_type == "OLS": impact = dev * weight
            else: impact = np.nan 
            ro_var_data.append([ro_param_names[i], RO_MRA_BASELINE[ro_param_keys[i]], ro_live_input[i], dev, weight, impact])
            
        ro_mra_data['Variance_DF'] = pd.DataFrame(ro_var_data, columns=["Parameter", "Baseline", "Live Input", "Deviation", "Regression Weight", "Impact (m³/h)"])

        # --- THE FIX: MERGED RO TABS ---
        ro_tabs = st.tabs(["Inputs", "Performance", "Chemical Dosing", "Prediction", "Reports", "Model", "Bulk Upload"])
        
        with ro_tabs[0]:
            st.subheader("HERO Plant Inputs")
            col1, col2 = st.columns(2)
            with col1:
                st.session_state.ro_feed_flow = st.number_input("Feed Flow (m³/hr)", value=st.session_state.ro_feed_flow, key="in_ro_feed_flow")
                st.session_state.ro_feed_tds = st.number_input("Feed TDS (ppm)", value=st.session_state.ro_feed_tds, key="in_ro_feed_tds")
            with col2:
                st.session_state.ro_perm_flow = st.number_input("Permeate Flow (m³/hr)", value=st.session_state.ro_perm_flow, key="in_ro_perm_flow")
                st.session_state.ro_perm_tds = st.number_input("Permeate TDS (ppm)", value=st.session_state.ro_perm_tds, key="in_ro_perm_tds")
                
            st.markdown("### Pre-treatment Parameters")
            p1, p2, p3 = st.columns(3)
            with p1:
                st.session_state.ro_clarifier_tss = st.number_input("Clarifier Outlet TSS (ppm)", value=st.session_state.ro_clarifier_tss, key="in_ro_clarifier_tss")
                st.session_state.ro_pdmf_tss = st.number_input("PDMF Outlet TSS (ppm)", value=st.session_state.ro_pdmf_tss, key="in_ro_pdmf_tss")
                st.session_state.ro_sdmf_tss = st.number_input("SDMF Outlet TSS (ppm)", value=st.session_state.ro_sdmf_tss, key="in_ro_sdmf_tss")
            with p2:
                st.session_state.ro_soft_hard = st.number_input("Softener Outlet Hardness (ppm)", value=st.session_state.ro_soft_hard, key="in_ro_soft_hard")
                st.session_state.ro_hru_hard = st.number_input("HRU Outlet Hardness (ppm)", value=st.session_state.ro_hru_hard, key="in_ro_hru_hard")
                st.session_state.ro_sdi = st.number_input("Cartridge Filter SDI", value=st.session_state.ro_sdi, key="in_ro_sdi")
            with p3:
                st.session_state.ro_perm_ph = st.number_input("Permeate pH", value=st.session_state.ro_perm_ph, key="in_ro_perm_ph")
                st.session_state.ro_perm_cod = st.number_input("Permeate COD (ppm)", value=st.session_state.ro_perm_cod, key="in_ro_perm_cod")
                
        with ro_tabs[1]:
            st.subheader("HERO Key Performance Indicators")
            recovery = (st.session_state.ro_perm_flow / st.session_state.ro_feed_flow * 100) if st.session_state.ro_feed_flow > 0 else 0
            rejection = ((st.session_state.ro_feed_tds - st.session_state.ro_perm_tds) / st.session_state.ro_feed_tds * 100) if st.session_state.ro_feed_tds > 0 else 0
            
            k1, k2, k3 = st.columns(3)
            rec_delta = f"{recovery - 85.0:.1f}% from Target" if recovery < 85.0 else "Target Met"
            k1.metric("Overall Plant Recovery", f"{recovery:.1f} %", delta=rec_delta, delta_color="normal" if recovery >= 85.0 else "inverse")
            k2.metric("Salt Rejection", f"{rejection:.1f} %")
            k3.metric("Permeate TDS", f"{st.session_state.ro_perm_tds:.1f} ppm", delta="Target: <150 ppm", delta_color="off")
            
            st.divider()
            st.subheader("Treatment Parameter Check")
            guarantees = [
                ("Clarifier Outlet TSS", st.session_state.ro_clarifier_tss, "< 10 ppm", 10.0),
                ("PDMF Outlet TSS", st.session_state.ro_pdmf_tss, "< 3 ppm", 3.0),
                ("SDMF Outlet TSS", st.session_state.ro_sdmf_tss, "< 1 ppm", 1.0),
                ("Softener O/L Hardness", st.session_state.ro_soft_hard, "< 5 ppm as CaCO3", 5.0),
                ("HRU O/L Hardness", st.session_state.ro_hru_hard, "< 1 ppm as CaCO3", 1.0),
                ("Cartridge Filter SDI", st.session_state.ro_sdi, "< 3", 3.0),
                ("RO Permeate COD", st.session_state.ro_perm_cod, "< 10 ppm", 10.0)
            ]
            for name, val, target_text, limit in guarantees:
                col_name, col_val, col_target, col_status = st.columns([2, 1, 1, 1])
                col_name.write(f"**{name}**")
                col_val.write(f"{val}")
                col_target.write(target_text)
                if val < limit: col_status.success("Pass")
                else: col_status.error("Fail")
                    
            c_name, c_val, c_target, c_status = st.columns([2, 1, 1, 1])
            c_name.write("**RO Permeate pH**")
            c_val.write(f"{st.session_state.ro_perm_ph}")
            c_target.write("7.0 - 7.5")
            if 7.0 <= st.session_state.ro_perm_ph <= 7.5: c_status.success("Pass")
            else: c_status.error("Fail")
                
        with ro_tabs[2]:
            st.subheader("Chemical Dosing Control")
            st.info("AI-driven Optimal Dose Recommendations currently in development.")
            c1, c2, c3 = st.columns(3)
            with c1:
                st.markdown("### Coagulant (Clarifier)")
                st.session_state.ro_coag_ppm = st.number_input("Target Dosing (PPM)", value=st.session_state.ro_coag_ppm, key="in_ro_coag_ppm")
                st.info(f"**Requirement:** {(st.session_state.ro_feed_flow * st.session_state.ro_coag_ppm) / 1000:.2f} kg/hr")
            with c2:
                st.markdown("### Flocculant (Clarifier)")
                st.session_state.ro_floc_ppm = st.number_input("Target Dosing (PPM)", value=st.session_state.ro_floc_ppm, key="in_ro_floc_ppm")
                st.info(f"**Requirement:** {(st.session_state.ro_feed_flow * st.session_state.ro_floc_ppm) / 1000:.2f} kg/hr")
            with c3:
                st.markdown("### SMBS (RO Feed)")
                st.session_state.ro_smbs_ppm = st.number_input("Target Dosing (PPM)", value=st.session_state.ro_smbs_ppm, key="in_ro_smbs_ppm")
                st.info(f"**Requirement:** {(st.session_state.ro_feed_flow * st.session_state.ro_smbs_ppm) / 1000:.2f} kg/hr")

        with ro_tabs[3]:
            st.subheader("Production Prediction")
            st.markdown("Modify process inputs to execute 'What-If' scenarios and check for membrane fouling.")
            controls_col, calc_col = st.columns([1, 2])
            
            with controls_col:
                st.session_state.ro_feed_flow = st.number_input("Feed Flow (m³/hr)", value=st.session_state.ro_feed_flow, key="t5_ro_feed_flow")
                st.session_state.ro_feed_tds = st.number_input("Feed TDS (ppm)", value=st.session_state.ro_feed_tds, key="t5_ro_feed_tds")
                st.session_state.ro_coag_ppm = st.number_input("Coagulant PPM", value=st.session_state.ro_coag_ppm, key="t5_ro_coag_ppm")
                st.session_state.ro_smbs_ppm = st.number_input("SMBS PPM", value=st.session_state.ro_smbs_ppm, key="t5_ro_smbs_ppm")

            with calc_col:
                k1, k2, k3 = st.columns(3)
                k1.metric("Actual Permeate SCADA", f"{ro_mra_data['Actual']:.1f} m³/h")
                k2.metric(f"Normalized Twin ({ro_model_type})", f"{ro_mra_data['Predicted']:.1f} m³/h")
                
                ro_diff_pct = (ro_mra_data['Residual'] / ro_mra_data['Predicted']) * 100 if ro_mra_data['Predicted'] > 0 else 0
                if ro_diff_pct <= -5.0: k3.error(f"Residual: {ro_mra_data['Residual']:.1f} m³/h ({ro_diff_pct:.1f}%) - Membrane CIP Required")
                elif ro_diff_pct <= -4.0: k3.warning(f"Residual: {ro_mra_data['Residual']:.1f} m³/h ({ro_diff_pct:.1f}%) - Warning: Scaling detected")
                else: k3.success(f"Residual: {ro_mra_data['Residual']:.1f} m³/h ({ro_diff_pct:.1f}%) - CLEAN")
                    
                if ro_model_type != "OLS": st.info("AI model active. Variance breakdown is only available for OLS models.")
                st.dataframe(ro_mra_data['Variance_DF'].style.format({"Baseline": "{:.1f}", "Live Input": "{:.1f}", "Deviation": "{:+.1f}", "Regression Weight": "{:.3f}", "Impact (m³/h)": "{:+.1f}"}, na_rep="-"), use_container_width=True, hide_index=True)

        with ro_tabs[4]:
            st.subheader("Reporting & Data Logging")
            rep_tabs = st.tabs(["Daily Dashboard", "Historical Database", "Performance Trends", "Data Explorer"])
            
            with rep_tabs[0]:
                m_col1, m_col2, m_col3, m_col4 = st.columns(4)
                m_col1.metric("Target Record Date", log_date.strftime('%d/%m/%Y')) 
                m_col2.metric("Plant Recovery", f"{(st.session_state.ro_perm_flow / st.session_state.ro_feed_flow * 100) if st.session_state.ro_feed_flow > 0 else 0:.1f} %")
                m_col3.metric("Salt Rejection", f"{((st.session_state.ro_feed_tds - st.session_state.ro_perm_tds) / st.session_state.ro_feed_tds * 100) if st.session_state.ro_feed_tds > 0 else 0:.1f} %")
                
                ro_diff_pct = (ro_mra_data['Residual'] / ro_mra_data['Predicted']) * 100 if ro_mra_data['Predicted'] > 0 else 0
                if ro_diff_pct <= -5.0: delta_text, d_color = f"{ro_diff_pct:.1f}% (Fouling Critical)", "inverse"
                elif ro_diff_pct <= -4.0: delta_text, d_color = f"{ro_diff_pct:.1f}% (Deviation Warning)", "inverse"
                else: delta_text, d_color = f"{ro_diff_pct:.1f}% (Clean Baseline)", "normal"
                    
                m_col4.metric("Normalized Flow Gap", f"{ro_mra_data['Residual']:.1f} m³/h", delta=delta_text, delta_color=d_color)
                
                st.divider()
                st.text_area("Remarks & Performance Observations", key="ro_in_remarks", placeholder="Record operational shift anomalies, CEB/CIP clean notes here...")
                
                st.markdown("### Save Daily Record")
                c_pwd, c_save, c_export, c_csv = st.columns([1.5, 1, 1, 1])
                with c_pwd: pwd_append = st.text_input("Master Password", type="password", key="ro_pwd_append", label_visibility="collapsed", placeholder="Enter master password to save")
                with c_save:
                    if st.button("Save Operational Record", use_container_width=True):
                        if pwd_append == "12345678":
                            db_dict = {
                                "Date": [log_date_str], 
                                "Feed Flow": [st.session_state.ro_feed_flow], "Permeate Flow": [st.session_state.ro_perm_flow],
                                "Feed TDS": [st.session_state.ro_feed_tds], "Permeate TDS": [st.session_state.ro_perm_tds],
                                "Clarifier TSS": [st.session_state.ro_clarifier_tss], "PDMF TSS": [st.session_state.ro_pdmf_tss],
                                "SDMF TSS": [st.session_state.ro_sdmf_tss], "Softener Hardness": [st.session_state.ro_soft_hard],
                                "HRU Hardness": [st.session_state.ro_hru_hard], "Cartridge SDI": [st.session_state.ro_sdi],
                                "Permeate pH": [st.session_state.ro_perm_ph], "Permeate COD": [st.session_state.ro_perm_cod],
                                "Coagulant PPM": [st.session_state.ro_coag_ppm], "Flocculant PPM": [st.session_state.ro_floc_ppm],
                                "SMBS PPM": [st.session_state.ro_smbs_ppm], 
                                "Recovery": [round((st.session_state.ro_perm_flow / st.session_state.ro_feed_flow * 100) if st.session_state.ro_feed_flow > 0 else 0, 2)],
                                "Rejection": [round(((st.session_state.ro_feed_tds - st.session_state.ro_perm_tds) / st.session_state.ro_feed_tds * 100) if st.session_state.ro_feed_tds > 0 else 0, 2)],
                                "Residual": [round(ro_mra_data['Residual'], 2)], "Remarks": [st.session_state.get('ro_in_remarks', '')]
                            }
                            new_log = pd.DataFrame(db_dict)
                            st.session_state.ro_daily_logs = pd.concat([st.session_state.ro_daily_logs, new_log], ignore_index=True)
                            save_database(db_conn, st.session_state.ro_daily_logs, RO_LOCAL_DB_FILE)
                            st.success("Record saved.")
                        elif pwd_append != "": st.error("Incorrect password.")
                with c_export:
                    word_file = generate_ro_comprehensive_report(log_date, st.session_state, ro_mra_data)
                    st.download_button("Export Word Document (.docx)", data=word_file, file_name=f"RO_ExecutiveReport_{log_date_str}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                with c_csv:
                    csv_file = generate_ro_daily_csv(log_date, st.session_state, ro_mra_data)
                    st.download_button("Export Tabular Values (.csv)", data=csv_file, file_name=f"RO_DataRecord_{log_date_str}.csv", mime="text/csv", use_container_width=True)

            with rep_tabs[1]:
                st.markdown("#### Master System Registry Database")
                display_cols = [c for c in RO_EXACT_DB_COLUMNS if c in st.session_state.ro_daily_logs.columns]
                edited_db = st.data_editor(st.session_state.ro_daily_logs[display_cols] if not st.session_state.ro_daily_logs.empty else st.session_state.ro_daily_logs, num_rows="dynamic", use_container_width=True)
                c_sync_pwd, c_sync, c_dl = st.columns([2, 1, 1])
                with c_sync_pwd: pwd_sync = st.text_input("Master Password", type="password", key="ro_pwd_sync", label_visibility="collapsed", placeholder="Enter master password to save")
                with c_sync:
                    if st.button("Synchronize Registry", use_container_width=True, key="ro_sync_btn"):
                        if pwd_sync == "12345678":
                            st.session_state.ro_daily_logs = edited_db
                            save_database(db_conn, st.session_state.ro_daily_logs, RO_LOCAL_DB_FILE)
                            st.success("Database updated.")
                        else: st.error("Incorrect password.")
                with c_dl:
                    st.download_button("Download Database Offline Backup", data=st.session_state.ro_daily_logs.to_csv(index=False).encode('utf-8'), file_name=f"RO_MasterRegistry_Backup.csv", mime='text/csv', use_container_width=True, key="ro_dl_btn")

                st.divider()
                st.markdown("#### Aggregated Monthly Performance Generator")
                if not st.session_state.ro_daily_logs.empty:
                    df_logs = st.session_state.ro_daily_logs.copy()
                    df_logs['Date'] = pd.to_datetime(df_logs['Date'], format='%Y-%m-%d', errors='coerce').fillna(pd.to_datetime(df_logs['Date'], errors='coerce', dayfirst=True))
                    month_data = df_logs[(df_logs['Date'].dt.month == log_date.month) & (df_logs['Date'].dt.year == log_date.year)].copy()
                    if not month_data.empty:
                        if st.button("Compile and Generate Monthly Summary (.docx)", use_container_width=True, key="ro_month_btn"):
                            monthly_doc = generate_ro_monthly_report(month_data, log_date.strftime('%B'), str(log_date.year))
                            st.download_button("Download Monthly Briefing Document", data=monthly_doc, file_name=f"RO_MonthlySummary_{log_date.strftime('%b_%Y')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="ro_dl_month")

            with rep_tabs[2]:
                if not st.session_state.ro_daily_logs.empty:
                    df_logs = st.session_state.ro_daily_logs.copy()
                    df_logs['Date'] = pd.to_datetime(df_logs['Date'], format='%Y-%m-%d', errors='coerce').fillna(pd.to_datetime(df_logs['Date'], errors='coerce', dayfirst=True))
                    
                    min_date = df_logs['Date'].min().date() if not df_logs['Date'].isnull().all() else datetime.date(2023, 1, 1)
                    max_date = df_logs['Date'].max().date() if not df_logs['Date'].isnull().all() else datetime.date.today()
                    
                    st.markdown("##### Date Range")
                    d_col1, d_col2 = st.columns(2)
                    with d_col1: start_date = st.date_input("Start Threshold Date", min_date, key="ro_start_d1")
                    with d_col2: end_date = st.date_input("End Threshold Date", max_date, key="ro_end_d1")
                    
                    mask = (df_logs['Date'].dt.date >= start_date) & (df_logs['Date'].dt.date <= end_date)
                    df_filtered = df_logs.loc[mask]
                    
                    q_col1, q_col2 = st.columns(2)
                    with q_col1:
                        st.markdown("#### Performance Recovery Trend")
                        if len(df_filtered) > 1:
                            rec_chart = alt.Chart(df_filtered).mark_circle().encode(x=alt.X('Date:T', title="Evaluation Timeline"), y=alt.Y('Recovery:Q', scale=alt.Scale(zero=False)))
                            st.altair_chart(rec_chart + rec_chart.transform_regression('Date', 'Recovery').mark_line(color='red'), use_container_width=True)
                    with q_col2:
                        st.markdown("#### Normalized Performance Gap")
                        if len(df_filtered) > 1:
                            htc_chart = alt.Chart(df_filtered).mark_line(point=True, color='orange').encode(x=alt.X('Date:T', title="Evaluation Timeline"), y=alt.Y('Residual:Q', scale=alt.Scale(zero=False), title="Fouling Residual (m³/h)"))
                            st.altair_chart(htc_chart + htc_chart.transform_regression('Date', 'Residual').mark_line(color='black'), use_container_width=True)

            with rep_tabs[3]:
                st.markdown("#### Data Explorer")
                if not st.session_state.ro_daily_logs.empty:
                    exp_df = st.session_state.ro_daily_logs.copy()
                    exp_df['Date'] = pd.to_datetime(exp_df['Date'], format='%Y-%m-%d', errors='coerce').fillna(pd.to_datetime(exp_df['Date'], errors='coerce', dayfirst=True))
                    
                    min_date2 = exp_df['Date'].min().date() if not exp_df['Date'].isnull().all() else datetime.date(2023, 1, 1)
                    max_date2 = exp_df['Date'].max().date() if not exp_df['Date'].isnull().all() else datetime.date.today()
                    
                    d_col1, d_col2 = st.columns(2)
                    with d_col1: start_date2 = st.date_input("Start Horizon Date", min_date2, key="ro_start_d2")
                    with d_col2: end_date2 = st.date_input("End Horizon Date", max_date2, key="ro_end_d2")
                    
                    mask2 = (exp_df['Date'].dt.date >= start_date2) & (exp_df['Date'].dt.date <= end_date2)
                    exp_df = exp_df.loc[mask2]
                    
                    num_cols = [col for col in exp_df.columns if col not in ['Date', 'Remarks']]
                    x_c, y_c, t_c = st.columns(3)
                    with x_c: exp_x = st.selectbox("Select Independent Domain X-Axis", ['Date'] + num_cols, index=0, key="ro_x")
                    with y_c: exp_y = st.selectbox("Select Dependent Variable Y-Axis", num_cols, index=0, key="ro_y")
                    with t_c: exp_type = st.selectbox("Select Functional Chart Variant", ["Line Chart", "Scatter Plot", "Bar Chart"], key="ro_chart_type")
                    
                    if exp_type == "Line Chart": chart = alt.Chart(exp_df).mark_line(point=True).encode(x=alt.X(f"{exp_x}{':T' if exp_x == 'Date' else ':Q'}"), y=alt.Y(f"{exp_y}:Q", scale=alt.Scale(zero=False)), tooltip=[exp_x, exp_y])
                    elif exp_type == "Scatter Plot": chart = alt.Chart(exp_df).mark_circle(size=80).encode(x=alt.X(f"{exp_x}{':T' if exp_x == 'Date' else ':Q'}"), y=alt.Y(f"{exp_y}:Q", scale=alt.Scale(zero=False)), tooltip=[exp_x, exp_y])
                    else: chart = alt.Chart(exp_df).mark_bar().encode(x=alt.X(f"{exp_x}{':T' if exp_x == 'Date' else ':N'}"), y=alt.Y(f"{exp_y}:Q"), tooltip=[exp_x, exp_y])
                    st.altair_chart(chart.interactive(), use_container_width=True)
                else:
                    st.info("No active historical registry values detected to perform correlation modeling.")

        with ro_tabs[5]:
            st.subheader("Prediction Model Setup")
            if not SKLEARN_INSTALLED:
                st.error("The scikit-learn package is not installed.")
            else:
                st.markdown("### Baseline Coefficients")
                st.markdown(f"**Active Model:** `{ro_model_type}`")
                c_reset, _ = st.columns([1, 1])
                with c_reset:
                    if st.button("Reset to Default Coefficients", use_container_width=True, key="ro_factory_reset"):
                        st.session_state.ro_mra_coef = RO_MRA_COEF_BASE.copy()
                        save_config(db_conn, st.session_state.ro_mra_coef, RO_LOCAL_CONFIG_FILE)
                        st.success("Coefficients reset to defaults.")
                        time.sleep(1.5)
                        st.rerun()

                st.divider()
                st.markdown("### Multi-Variable Predictive Optimization Logic Model Builder")
                
                req_cols = ["Date", "Permeate Flow", "Feed Flow", "Feed TDS", "Coagulant PPM", "SMBS PPM"]
                template_df = pd.DataFrame(columns=req_cols)
                st.download_button(label="Download Training Template", data=template_df.to_csv(index=False).encode('utf-8'), file_name='RO_ML_CalibrationTemplate.csv', mime='text/csv', key="ro_train_dl")
                
                st.divider()
                uploaded_file = st.file_uploader("Upload Training Data", type=["csv"], key="ro_mra_trainer")
                
                if uploaded_file is not None:
                    try:
                        df_train = pd.read_csv(uploaded_file)
                        if not all(col in df_train.columns for col in req_cols): st.error(f"Training file is missing required columns.")
                        else:
                            for col in req_cols:
                                if col != "Date":
                                    if df_train[col].dtype == object: df_train[col] = pd.to_numeric(df_train[col].astype(str).str.replace(',', '', regex=False), errors='coerce')
                            
                            df_train = df_train.dropna(subset=[c for c in req_cols if c != "Date"])
                            st.success(f"Training complete using {len(df_train)} rows.")
                            
                            if len(df_train) > 0:
                                X = df_train[["Feed Flow", "Feed TDS", "Coagulant PPM", "SMBS PPM"]]
                                Y = df_train["Permeate Flow"]
                                
                                model_ols = LinearRegression(fit_intercept=True).fit(X, Y)
                                r2_ols = r2_score(Y, model_ols.predict(X))
                                
                                model_rf = RandomForestRegressor(n_estimators=100, random_state=42).fit(X, Y)
                                r2_rf = r2_score(Y, model_rf.predict(X))
                                
                                if XGB_INSTALLED:
                                    model_xgb = xgb.XGBRegressor(n_estimators=100, random_state=42).fit(X, Y)
                                    r2_xgb = r2_score(Y, model_xgb.predict(X))
                                
                                st.markdown("### Model Accuracy")
                                m1, m2, m3 = st.columns(3)
                                m1.metric("1. Linear OLS Fit (R² Coefficient)", f"{r2_ols * 100:.2f}%")
                                m2.metric("2. Random Forest Tree Logic (R²)", f"{r2_rf * 100:.2f}%")
                                if XGB_INSTALLED: m3.metric("3. Extreme Gradient Boost XGB (R²)", f"{r2_xgb * 100:.2f}%")
                                else: m3.warning("Advanced Gradient boosting library dependency not activated.")
                                
                                st.markdown("#### Dynamic Feature Sensitivity Weights / Scaling Coefficients")
                                comp_dict = {
                                    "Parameter": ["Feed_Flow", "Feed_TDS", "Coag_PPM", "SMBS_PPM"],
                                    "OLS (Coefficients)": np.round(model_ols.coef_, 4),
                                    "Random Forest (Importance %)": np.round(model_rf.feature_importances_ * 100, 2)
                                }
                                if XGB_INSTALLED: comp_dict["XGBoost (Importance %)"] = np.round(model_xgb.feature_importances_ * 100, 2)
                                st.dataframe(pd.DataFrame(comp_dict).style.format(precision=4), use_container_width=True, hide_index=True)
                                
                                st.markdown("### Select Active Model")
                                opts = ["OLS (Linear)", "Random Forest"]
                                if XGB_INSTALLED: opts.append("XGBoost")
                                selected_model = st.radio("Configure Active Live Prediction Logic Block:", opts, key="ro_model_radio")
                                
                                if st.button("Confirm & Activate Model", type="primary", use_container_width=True, key="ro_model_lock"):
                                    if selected_model == "OLS (Linear)":
                                        new_coefs = {
                                            "model_type": "OLS", "Intercept": float(model_ols.intercept_),
                                            "Feed_Flow": float(model_ols.coef_[0]), "Feed_TDS": float(model_ols.coef_[1]), 
                                            "Coag_PPM": float(model_ols.coef_[2]), "SMBS_PPM": float(model_ols.coef_[3])
                                        }
                                        st.session_state.ro_mra_coef = new_coefs
                                        save_config(db_conn, new_coefs, RO_LOCAL_CONFIG_FILE)
                                    else:
                                        target_m = model_rf if selected_model == "Random Forest" else model_xgb
                                        joblib.dump(target_m, RO_AI_MODEL_FILE)
                                        ai_coefs = {
                                            "model_type": selected_model,
                                            "Feed_Flow": float(target_m.feature_importances_[0]), "Feed_TDS": float(target_m.feature_importances_[1]), 
                                            "Coag_PPM": float(target_m.feature_importances_[2]), "SMBS_PPM": float(target_m.feature_importances_[3])
                                        }
                                        st.session_state.ro_mra_coef = ai_coefs
                                        save_config(db_conn, ai_coefs, RO_LOCAL_CONFIG_FILE)
                                        
                                    st.success(f"{selected_model} model activated.")
                                    time.sleep(1.5)
                                    st.rerun()
                            else: st.error("Uploaded data produced no valid values.")
                    except Exception as e: st.error(f"Error processing data: {e}")

        with ro_tabs[6]:
            st.subheader("Bulk Data Upload")
            st.markdown("Download the target spreadsheet schema file. Copy/pasting raw values in the exact historical Excel configuration is fully supported.")
            
            bulk_template = pd.DataFrame(columns=RO_EXACT_DB_COLUMNS)
            st.download_button(label="Download Template", data=bulk_template.to_csv(index=False).encode('utf-8'), file_name='RO_BulkMatrixInletSchema.csv', mime='text/csv', key="ro_bulk_dl")
            
            st.divider()
            bulk_file = st.file_uploader("Upload Batch File (.csv)", type=["csv"], key="ro_bulk_uploader")
            
            if bulk_file is not None:
                try:
                    df_bulk = pd.read_csv(bulk_file)
                    
                    missing = [c for c in RO_EXACT_DB_COLUMNS if c not in df_bulk.columns]
                    if missing:
                        st.warning(f"Some columns are missing and will be filled with baseline values: {', '.join(missing)}")
                        for c in missing: df_bulk[c] = np.nan
                    
                    num_cols = [c for c in RO_EXACT_DB_COLUMNS if c not in ["Date", "Remarks"]]
                    for col in num_cols:
                        if col in df_bulk.columns:
                            if df_bulk[col].dtype == object: df_bulk[col] = pd.to_numeric(df_bulk[col].astype(str).str.replace(',', '', regex=False), errors='coerce')
                    
                    df_bulk = df_bulk.dropna(subset=["Date"])
                    
                    if len(df_bulk) > 0:
                        df_bulk['Feed Flow'] = df_bulk['Feed Flow'].fillna(450.0)
                        df_bulk['Feed TDS'] = df_bulk['Feed TDS'].fillna(2000.0)
                        df_bulk['Coagulant PPM'] = df_bulk['Coagulant PPM'].fillna(2.0)
                        df_bulk['SMBS PPM'] = df_bulk['SMBS PPM'].fillna(3.0)
                        df_bulk['Permeate Flow'] = df_bulk['Permeate Flow'].fillna(0.0)
                        df_bulk['Permeate TDS'] = df_bulk['Permeate TDS'].fillna(90.0)
                        
                        df_bulk['Date_Clean'] = pd.to_datetime(df_bulk['Date'], dayfirst=True, errors='coerce').dt.strftime('%Y-%m-%d')
                        df_bulk['Recovery'] = np.where(df_bulk['Feed Flow'] > 0, df_bulk['Permeate Flow'] / df_bulk['Feed Flow'] * 100, 0)
                        df_bulk['Rejection'] = np.where(df_bulk['Feed TDS'] > 0, (df_bulk['Feed TDS'] - df_bulk['Permeate TDS']) / df_bulk['Feed TDS'] * 100, 0)

                        if ro_model_type == "OLS":
                            df_bulk['Predicted'] = (
                                ro_coefs["Intercept"] + 
                                (ro_coefs["Feed_Flow"] * df_bulk['Feed Flow']) + 
                                (ro_coefs["Feed_TDS"] * df_bulk['Feed TDS']) + 
                                (ro_coefs["Coag_PPM"] * df_bulk['Coagulant PPM']) + 
                                (ro_coefs["SMBS_PPM"] * df_bulk['SMBS PPM'])
                            )
                        else:
                            try:
                                active_model = joblib.load(RO_AI_MODEL_FILE)
                                bulk_input_df = df_bulk[['Feed Flow', 'Feed TDS', 'Coagulant PPM', 'SMBS PPM']].copy()
                                bulk_input_df.columns = ["Feed_Flow", "Feed_TDS", "Coag_PPM", "SMBS_PPM"]
                                df_bulk['Predicted'] = active_model.predict(bulk_input_df)
                            except: df_bulk['Predicted'] = 0.0
                                
                        df_bulk['Residual'] = df_bulk['Permeate Flow'] - df_bulk['Predicted']
                        
                        db_ready_dict = {
                            "Date": df_bulk['Date_Clean'], 
                            "Feed Flow": df_bulk['Feed Flow'], "Permeate Flow": df_bulk['Permeate Flow'],
                            "Feed TDS": df_bulk['Feed TDS'], "Permeate TDS": df_bulk['Permeate TDS'],
                            "Clarifier TSS": df_bulk['Clarifier TSS'].fillna(8.0), "PDMF TSS": df_bulk['PDMF TSS'].fillna(2.0),
                            "SDMF TSS": df_bulk['SDMF TSS'].fillna(0.5), "Softener Hardness": df_bulk['Softener Hardness'].fillna(4.0),
                            "HRU Hardness": df_bulk['HRU Hardness'].fillna(0.5), "Cartridge SDI": df_bulk['Cartridge SDI'].fillna(2.5),
                            "Permeate pH": df_bulk['Permeate pH'].fillna(7.2), "Permeate COD": df_bulk['Permeate COD'].fillna(8.0),
                            "Coagulant PPM": df_bulk['Coagulant PPM'], "Flocculant PPM": df_bulk['Flocculant PPM'].fillna(1.0),
                            "SMBS PPM": df_bulk['SMBS PPM'], "Recovery": df_bulk['Recovery'].round(2),
                            "Rejection": df_bulk['Rejection'].round(2), "Residual": df_bulk['Residual'].round(1),
                            "Remarks": df_bulk['Remarks'].fillna("")
                        }
                                
                        db_ready_df = pd.DataFrame(db_ready_dict)
                        
                        st.success(f"Processed {len(db_ready_df)} rows.")
                        st.dataframe(db_ready_df.style.format(precision=2), use_container_width=True, hide_index=True)
                        
                        st.markdown("### Append Transferred Batch Elements")
                        c_pwd, c_save = st.columns([2, 2])
                        with c_pwd: pwd_bulk = st.text_input("Master Password", type="password", key="ro_pwd_bulk", label_visibility="collapsed", placeholder="Enter master password to save")
                        with c_save:
                            if st.button("Save Records to Database", use_container_width=True, key="ro_bulk_save"):
                                if pwd_bulk == "12345678":
                                    st.session_state.ro_daily_logs = pd.concat([st.session_state.ro_daily_logs, db_ready_df], ignore_index=True)
                                    st.session_state.ro_daily_logs = st.session_state.ro_daily_logs.drop_duplicates(subset=['Date'], keep='last').reset_index(drop=True)
                                    save_database(db_conn, st.session_state.ro_daily_logs, RO_LOCAL_DB_FILE)
                                    st.success("Records saved to database.")
                                    time.sleep(1.5)
                                    st.rerun()
                                elif pwd_bulk != "": st.error("Incorrect password.")
                    else: st.error("No valid rows found in file.")
                except Exception as e: st.error(f"Error processing upload: {e}")
        
        render_chatbot()
        return

if __name__ == "__main__":
    main()
